from PyQt5.QtCore import *
from PyQt5.QtWidgets import *
from PyQt5.QtGui import QFont
from PyQt5 import QtCore, QtWidgets
from PyQt5.QtGui import *
from PyQt5.QtCore import *
from PyQt5.QtWidgets import QMessageBox
from docx import Document
from docx.shared import Inches
import io
import csv
from PyQt5.Qt import *
import sys
import xlrd
import xlwt 
from xlwt import Workbook 

class MainWindow(QtWidgets.QWidget):

    switch_window = QtCore.pyqtSignal()
    switch_window1 = QtCore.pyqtSignal()


    def __init__(self):
        QtWidgets.QWidget.__init__(self)
        self.setWindowTitle('Mark Easy(40%)')
        #self.resize(1300, 700)

        layout = QtWidgets.QGridLayout()

        label = QLabel(" STAFF NAME :")
        label.setFixedWidth(100)

        layout.addWidget(label, 0, 0)
        label = QLabel(" DESIGNATION :")
        label.setFixedWidth(100)

        layout.addWidget(label, 1, 0)
        label = QLabel("BRANCH :")
        label.setFixedWidth(100)
        label.setWordWrap(True)
        layout.addWidget(label, 5,0)
        dept=['MECH','EEE','CSE','IT']
        self.combobox1 = QComboBox()
        for i in  (dept):
            self.combobox1.addItem(i)
        '''self.combobox.addItem("EEE")
        self.combobox.addItem("ECE")
        self.combobox.addItem("CSE")
        self.combobox.addItem("IT")'''
        self.combobox1.setFixedWidth(100)
        self.combobox1.currentTextChanged.connect(self.combobox_changed)
        layout.addWidget(self.combobox1, 5,1)   

        label = QLabel("YEAR     :")
        label.setWordWrap(True)
        label.setFixedWidth(100)
        layout.addWidget(label, 4,0)
        year=['1 year','2 year','3 year','4 year']
        self.combobox2 = QComboBox()
        for i in  (year):
            self.combobox2.addItem(i)
        self.combobox2.setFixedWidth(100)
        self.combobox2.currentTextChanged.connect(self.combobox_changed)
        layout.addWidget(self.combobox2, 4,1)   

        label = QLabel("SUBJECT NAME :")
        label.setWordWrap(True)
        label.setFixedWidth(100)
        layout.addWidget(label, 2,0)

        self.lineedit2 = QLineEdit()
        self.lineedit2.returnPressed.connect(self.return_pressed)
        layout.addWidget(self.lineedit2, 2, 1) 

        label = QLabel("SUBJECT CODE :")
        label.setWordWrap(True)
        label.setFixedWidth(100)
        layout.addWidget(label, 3,0)

        self.lineedit3 = QLineEdit()
        self.lineedit3.returnPressed.connect(self.return_pressed)
        self.lineedit3.setFixedWidth(200)
        layout.addWidget(self.lineedit3, 3, 1)

        label = QLabel("SECTION :")
        label.setWordWrap(True)
        label.setFixedWidth(100)
        layout.addWidget(label, 6,0)

        self.combobox3 = QComboBox()
        self.combobox3.addItem("A")
        self.combobox3.addItem("B")
        self.combobox3.addItem("C")
        self.combobox3.setFixedWidth(100)
        self.combobox3.currentTextChanged.connect(self.combobox_changed)
        layout.addWidget(self.combobox3, 6,1) 


        self.lineedit = QLineEdit()
        self.lineedit.returnPressed.connect(self.return_pressed)
        self.lineedit.setFixedWidth(200)
        layout.addWidget(self.lineedit, 0, 1)
        
        self.lineedit1 = QLineEdit()
        self.lineedit1.returnPressed.connect(self.return_pressed)
        self.lineedit1.setFixedWidth(200)
        layout.addWidget(self.lineedit1, 1, 1)
        self.button = QtWidgets.QPushButton('Next')
        self.button.setFixedWidth(100)
        self.button.clicked.connect(self.switches)
        layout.addWidget(self.button,7,5)
        self.button1 = QtWidgets.QPushButton('Back')
        self.button1.setFixedWidth(100)
        self.button1.clicked.connect(self.switch1)
        layout.addWidget(self.button1,7,0)

        self.setLayout(layout)
    def return_pressed(self):
        #print(type(self.lineedit.text()))
        #print(type(self.lineedit1.text()))
        #print(type(self.lineedit2.text()))
        #print(type(self.lineedit3.text()))
        self.staffname=self.lineedit.text()
        self.designation=self.lineedit1.text()
        self.subname=self.lineedit2.text()
        self.subcode=self.lineedit3.text()
        #print("start",self.lineedit.text(),"end")
        #print("start",self.lineedit1.text(),"end")
        #print("start",self.lineedit2.text(),"end")
        #print("start",self.lineedit3.text(),"end")
    def check(self):
        self.msg=[]
        if  self.staffname:
            print(self.staffname)
        else:
            self.msg.append("staffname")
        if  self.designation:
            print(self.designation)
        else:
            self.msg.append("Designation")
        if  self.subname:
            print(self.subname)
        else:
            self.msg.append("Subject Name")
        if  self.subcode:
            print(self.subcode)
        else:
            self.msg.append("Subject Code")
        
        print(len(self.msg))
    def combobox_changed(self):
        self.depart = self.combobox1.currentText()
        #print(self.depart)
        self.year = self.combobox2.currentText()
        #print(self.year)
        self.section = self.combobox3.currentText()
        #print(self.section)
   
    def aboutmsg(self):
        if  self.staffname:
            QMessageBox.about(self, "Please Listen "+str(self.staffname),"enter the details  "+str(self.msg))
        else:
            QMessageBox.about(self, "Please Listen ","enter the details  "+str(self.msg))

    def finish(self):
        fileName, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Open Excel File "+str(staffname),
               (QtCore.QDir.homePath()), "Excel (*.xlsx *.xls)")
        if fileName:
            global loc
            loc = fileName 

            # To open Workbook 
            wb = xlrd.open_workbook(loc) 
            sheet = wb.sheet_by_index(0) 
            print( sheet.name)
            print (sheet.nrows)
            print (sheet.ncols)
            # For row 0 and column 0

            
            global trows
            trows=sheet.nrows
            cols=sheet.ncols

            studentdata=[]
            for row in range(trows):
                datainrow=[]
                for col in range(cols):
                        text=sheet.cell_value(row,col)
                        datainrow.append((text))
                studentdata.append(datainrow)
                
            #print(studentdata)
            global studentdata1
            studentdata1=studentdata
            self.switch_window.emit()

    def switch1(self):
        self.switch_window1.emit()
    def switches(self):
        self.return_pressed()
        self.check()
        self.combobox_changed()
        if (len(self.msg)>=1):
            self.aboutmsg()
        elif(len(self.msg)==0):
            global staffname
            global designation
            global subjectname
            global subjectcode
            global year
            global department
            global section

            staffname=self.staffname
            designation=self.designation
            subjectname=self.subname
            subjectcode=self.subcode

            year=self.year
            department=self.depart
            section=self.section
            self.finish()
class MainWindow1(QtWidgets.QWidget):

    switch_window = QtCore.pyqtSignal()
    switch_window1 = QtCore.pyqtSignal()


    def __init__(self):
        QtWidgets.QWidget.__init__(self)
        self.setWindowTitle('Mark Easy(60%)')

        self.layout = QGridLayout()
        self.setLayout(self.layout)
        self.com1=QLabel("Number of 'PART A' Questions ")
        self.com1.setFixedWidth(150)
        self.layout.addWidget(self.com1, 0, 0)
        self.radiobuttona1 = QSpinBox()
        self.radiobuttona1.setSuffix("Questions")
        self.radiobuttona1.setMinimum(0)
        self.radiobuttona1.setMaximum(30)
        self.radiobuttona1.setFixedWidth(150)
        self.layout.addWidget(self.radiobuttona1,0,1)
        self.radiobuttona1.valueChanged.connect(self.spinbox_toggled2)
        

        self.com1=QLabel("Number of 'PART B' Questions ")
        self.com1.setFixedWidth(150)
        self.layout.addWidget(self.com1, 0, 2)
        self.radiobuttonb1 = QSpinBox()
        self.radiobuttonb1.setSuffix("Questions")
        self.radiobuttonb1.setMinimum(0)
        self.radiobuttonb1.setMaximum(20)
        self.radiobuttonb1.setSingleStep(2)
        self.radiobuttonb1.setFixedWidth(150)
        self.layout.addWidget(self.radiobuttonb1,0,3)
        self.radiobuttonb1.valueChanged.connect(self.spinbox_toggledb)

        self.com1=QLabel("Number of 'PART C' Questions ")
        self.com1.setFixedWidth(150)
        self.layout.addWidget(self.com1, 0, 4)
        self.radiobuttonc1 = QSpinBox()
        self.radiobuttonc1.setSuffix("Questions")
        self.radiobuttonc1.setMinimum(0)
        self.radiobuttonc1.setMaximum(10)
        self.radiobuttonc1.setSingleStep(2)
        self.radiobuttonc1.setFixedWidth(150)
        self.layout.addWidget(self.radiobuttonc1,0,5)
        
        self.radiobuttonc1.valueChanged.connect(self.spinbox_toggledc)
        self.com1=QLabel("Marks For Part A")
        self.com1.setFixedWidth(150)
        self.layout.addWidget(self.com1, 2, 0)


        self.radiobuttonma1 = QSpinBox()
        self.radiobuttonma1.setSuffix("Marks")
        self.radiobuttonma1.setMinimum(1)
        self.radiobuttonma1.setMaximum(2)
        self.radiobuttonma1.setFixedWidth(150)
        self.layout.addWidget(self.radiobuttonma1,2,1)
        self.radiobuttonma1.valueChanged.connect(self.marks)


        self.com1=QLabel("Marks For Part B")
        self.com1.setFixedWidth(150)
        self.layout.addWidget(self.com1, 2, 2)
        self.radiobuttonmb1 = QSpinBox()
        self.radiobuttonmb1.setSuffix("Marks")
        self.radiobuttonmb1.setMinimum(1)
        self.radiobuttonmb1.setMaximum(13)
        self.radiobuttonmb1.setFixedWidth(150)
        self.layout.addWidget(self.radiobuttonmb1,2,3)
        self.radiobuttonmb1.valueChanged.connect(self.marks)

        self.com1=QLabel("Marks For  Part C")
        self.com1.setFixedWidth(150)
        self.layout.addWidget(self.com1, 2, 4)
        self.radiobuttonmc1 = QSpinBox()
        self.radiobuttonmc1.setSuffix("Marks")
        self.radiobuttonmc1.setMinimum(1)
        self.radiobuttonmc1.setMaximum(15)
        self.radiobuttonmc1.setFixedWidth(150)
        self.layout.addWidget(self.radiobuttonmc1,2,5)
        self.radiobuttonmc1.valueChanged.connect(self.marks)
       
        self.button = QtWidgets.QPushButton('Next')
        self.button.setFixedWidth(100)
        self.button.clicked.connect(self.switch)
        self.layout.addWidget(self.button,4,4)
    def marks(self):
        global marka
        global markb
        global markc
        marka=self.radiobuttonma1.value()
        markb=self.radiobuttonmb1.value()
        markc=self.radiobuttonmc1.value()
    def spinbox_toggled2(self):
        global partasp
        partasp=self.radiobuttona1.value()
    def spinbox_toggledb(self):
        global partbsp
        partbsp=self.radiobuttonb1.value()
    def spinbox_toggledc(self):
        global partcsp
        partcsp=self.radiobuttonc1.value()
    def partcmsg(self):
        QMessageBox.about(self, "Please Listen "+str(staffname), "you entered '0' questions for part C")
    def partbmsg(self):
        QMessageBox.about(self, "Please Listen "+str(staffname), "you entered '0' questions for part B")
    def partcbmsg(self):
        QMessageBox.about(self, "Please Listen "+str(staffname), "you entered '0' questions for part B and part C")
    def partabcmsg(self):
        QMessageBox.about(self, "Please Listen "+str(staffname), "Nothing to next step")

    def switch(self):
        self.marks()
        global partasp
        partasp=self.radiobuttona1.value()
        global partbsp
        partbsp=self.radiobuttonb1.value()
        global partcsp
        partcsp=self.radiobuttonc1.value()
        partaspcorrection=self.radiobuttona1.value()
        partbspcorrection=self.radiobuttonb1.value()
        partcspcorrection=self.radiobuttonc1.value()
        print(partaspcorrection,partbspcorrection,partcspcorrection)
        if (partaspcorrection>0 and partbspcorrection>0 and partcspcorrection==0):
            self.partcmsg()
            self.switch_window.emit()

        elif (partaspcorrection>0 and partbspcorrection==0 and partcspcorrection>0):
            self.partbmsg()
            self.switch_window.emit()

        elif (partaspcorrection>0 and partbspcorrection==0 and partcspcorrection==0):
            self.partcbmsg()
            self.switch_window.emit()

        elif (partaspcorrection==0 and partbspcorrection==0 and partcspcorrection==0):
            self.partabcmsg()
        else:
            self.switch_window.emit()
class Windowthree(QtWidgets.QWidget):
    switch_window = QtCore.pyqtSignal()

    def __init__(self):
        QtWidgets.QWidget.__init__(self)
        self.setWindowTitle('Mark Easy(80%)')


        self.layout = QGridLayout()
        self.setLayout(self.layout)
        
        self.parta=partasp

        self.partb=partbsp
        self.partc=partcsp
        if (self.parta>0):
            label = QLabel("Part A")
            label.setFixedWidth(100)
        
            self.layout.addWidget(label, 0, 0)
        if (self.partb>0):
            label = QLabel("Part B")
            label.setFixedWidth(100)
            self.layout.addWidget(label, 0, 2)
        if (self.partc>0):
            label = QLabel("Part C")
            label.setFixedWidth(100)
            self.layout.addWidget(label, 0, 5)
        que=self.parta+1
        qre=1
        for i in range(0,self.partb):
           dr=" A "
           de=" B "
           df=str(que)
           mm=i%2
           nn=i%2
           
           if mm==0:
               
               dd=" Question  "+df+dr
           else:
                
               dd=" Question  "+df+de 
           com=QLabel(dd)
           com.setFont(QFont("Stencil",10))
           com.setFixedWidth(100)
           self.layout.addWidget(com,qre, 2)
           if i==1 or i==3 or i==5 or i==7 or i==9 or i==11 or i==13 or i==15 or i==17 or i==19:
              que +=1
           qre+=1
        que=1
        for i in range(0,self.parta):
           df=str(que)
           dd="Question  "+df 
           com=QLabel(dd)
           com.setFont(QFont("Stencil",10))
           com.setFixedWidth(100)
           self.layout.addWidget(com,que, 0)

           que +=1
        inter=self.partb/2
        chinter=int(inter)
        que=self.parta+chinter+1
        qre=1
        for i in range(0,self.partc):
           dr=" A "
           de=" B "
           df=str(que)
           mm=i%2
           nn=i%2
           
           if mm==0:
               
               dd=" Question  "+df+dr
           else:
                
               dd=" Question  "+df+de 
           com=QLabel(dd)
           com.setFont(QFont("Stencil",10))
           com.setFixedWidth(100)
           self.layout.addWidget(com,qre, 5)
           if i==1 or i==3 or i==5 or i==7 or i==9 :
              que +=1

           qre +=1
        
        partbsub=[]
    
      
        self.subdivisions()
        self.courseb()
        self.coursea()
        self.coursec()
        
        self.subdivisionsforc()
        self.button = QtWidgets.QPushButton('Next')
        self.button.setFixedWidth(100)
        self.button.clicked.connect(self.switch)
        self.layout.addWidget(self.button,10,8)
    def subdivisions(self):
        
        if (self.partb==1):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

        elif (self.partb==2):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

        elif (self.partb==3):
            self.a1  = QSpinBox()
            #self.a1.setSuffix(" subdivions ")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

        elif (self.partb==4):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

        elif (self.partb==5):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)
        elif (self.partb==6):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)


        elif (self.partb==7):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

           
        elif (self.partb==8):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

           

        elif (self.partb==9):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

           

        elif (self.partb==10):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

            self.a10  = QSpinBox()
            #self.a10.setSuffix("  subdivions")
            self.a10.setMinimum(1)
            self.a10.setMaximum(5)
            self.a10.setFixedWidth(100)
            self.layout.addWidget(self.a10,10,3)
            self.a10.valueChanged.connect(self.spinbox_toggled2)

        elif (self.partb==11):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

            self.a10  = QSpinBox()
            #self.a10.setSuffix("  subdivions")
            self.a10.setMinimum(1)
            self.a10.setMaximum(5)
            self.a10.setFixedWidth(100)
            self.layout.addWidget(self.a10,10,3)
            self.a10.valueChanged.connect(self.spinbox_toggled2)

            self.a11  = QSpinBox()
            #self.a11.setSuffix("  subdivions")
            self.a11.setMinimum(1)
            self.a11.setMaximum(5)
            self.a11.setFixedWidth(100)
            self.layout.addWidget(self.a11,11,3)
            self.a11.valueChanged.connect(self.spinbox_toggled2)


        elif (self.partb==12):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

            self.a10  = QSpinBox()
            #self.a10.setSuffix("  subdivions")
            self.a10.setMinimum(1)
            self.a10.setMaximum(5)
            self.a10.setFixedWidth(100)
            self.layout.addWidget(self.a10,10,3)
            self.a10.valueChanged.connect(self.spinbox_toggled2)

            self.a11  = QSpinBox()
            #self.a11.setSuffix("  subdivions")
            self.a11.setMinimum(1)
            self.a11.setMaximum(5)
            self.a11.setFixedWidth(100)
            self.layout.addWidget(self.a11,11,3)
            self.a11.valueChanged.connect(self.spinbox_toggled2)

            self.a12  = QSpinBox()
            #self.a12.setSuffix("  subdivions")
            self.a12.setMinimum(1)
            self.a12.setMaximum(5)
            self.a12.setFixedWidth(100)
            self.layout.addWidget(self.a12,12,3)
            self.a12.valueChanged.connect(self.spinbox_toggled2)

           
        elif (self.partb==13):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

            self.a10  = QSpinBox()
            #self.a10.setSuffix("  subdivions")
            self.a10.setMinimum(1)
            self.a10.setMaximum(5)
            self.a10.setFixedWidth(100)
            self.layout.addWidget(self.a10,10,3)
            self.a10.valueChanged.connect(self.spinbox_toggled2)

            self.a11  = QSpinBox()
            #self.a11.setSuffix("  subdivions")
            self.a11.setMinimum(1)
            self.a11.setMaximum(5)
            self.a11.setFixedWidth(100)
            self.layout.addWidget(self.a11,11,3)
            self.a11.valueChanged.connect(self.spinbox_toggled2)

            self.a12  = QSpinBox()
            #self.a12.setSuffix("  subdivions")
            self.a12.setMinimum(1)
            self.a12.setMaximum(5)
            self.a12.setFixedWidth(100)
            self.layout.addWidget(self.a12,12,3)
            self.a12.valueChanged.connect(self.spinbox_toggled2)

            self.a13  = QSpinBox()
            #self.a13.setSuffix("  subdivions")
            self.a13.setMinimum(1)
            self.a13.setMaximum(5)
            self.a13.setFixedWidth(100)
            self.layout.addWidget(self.a13,13,3)
            self.a13.valueChanged.connect(self.spinbox_toggled2)

           

        elif (self.partb==14):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

            self.a10  = QSpinBox()
            #self.a10.setSuffix("  subdivions")
            self.a10.setMinimum(1)
            self.a10.setMaximum(5)
            self.a10.setFixedWidth(100)
            self.layout.addWidget(self.a10,10,3)
            self.a10.valueChanged.connect(self.spinbox_toggled2)

            self.a11  = QSpinBox()
            #self.a11.setSuffix("  subdivions")
            self.a11.setMinimum(1)
            self.a11.setMaximum(5)
            self.a11.setFixedWidth(100)
            self.layout.addWidget(self.a11,11,3)
            self.a11.valueChanged.connect(self.spinbox_toggled2)

            self.a12  = QSpinBox()
            #self.a12.setSuffix("  subdivions")
            self.a12.setMinimum(1)
            self.a12.setMaximum(5)
            self.a12.setFixedWidth(100)
            self.layout.addWidget(self.a12,12,3)
            self.a12.valueChanged.connect(self.spinbox_toggled2)

            self.a13  = QSpinBox()
            #self.a13.setSuffix("  subdivions")
            self.a13.setMinimum(1)
            self.a13.setMaximum(5)
            self.a13.setFixedWidth(100)
            self.layout.addWidget(self.a13,13,3)
            self.a13.valueChanged.connect(self.spinbox_toggled2)

            self.a14  = QSpinBox()
            #self.a14.setSuffix("  subdivions")
            self.a14.setMinimum(1)
            self.a14.setMaximum(5)
            self.a14.setFixedWidth(100)
            self.layout.addWidget(self.a14,14,3)
            self.a14.valueChanged.connect(self.spinbox_toggled2)

           
        elif (self.partb==15):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

            self.a10  = QSpinBox()
            #self.a10.setSuffix("  subdivions")
            self.a10.setMinimum(1)
            self.a10.setMaximum(5)
            self.a10.setFixedWidth(100)
            self.layout.addWidget(self.a10,10,3)
            self.a10.valueChanged.connect(self.spinbox_toggled2)

            self.a11  = QSpinBox()
            #self.a11.setSuffix("  subdivions")
            self.a11.setMinimum(1)
            self.a11.setMaximum(5)
            self.a11.setFixedWidth(100)
            self.layout.addWidget(self.a11,11,3)
            self.a11.valueChanged.connect(self.spinbox_toggled2)

            self.a12  = QSpinBox()
            #self.a12.setSuffix("  subdivions")
            self.a12.setMinimum(1)
            self.a12.setMaximum(5)
            self.a12.setFixedWidth(100)
            self.layout.addWidget(self.a12,12,3)
            self.a12.valueChanged.connect(self.spinbox_toggled2)

            self.a13  = QSpinBox()
            #self.a13.setSuffix("  subdivions")
            self.a13.setMinimum(1)
            self.a13.setMaximum(5)
            self.a13.setFixedWidth(100)
            self.layout.addWidget(self.a13,13,3)
            self.a13.valueChanged.connect(self.spinbox_toggled2)

            self.a14  = QSpinBox()
            #self.a14.setSuffix("  subdivions")
            self.a14.setMinimum(1)
            self.a14.setMaximum(5)
            self.a14.setFixedWidth(100)
            self.layout.addWidget(self.a14,14,3)
            self.a14.valueChanged.connect(self.spinbox_toggled2)

            self.a15  = QSpinBox()
            #self.a15.setSuffix("  subdivions")
            self.a15.setMinimum(1)
            self.a15.setMaximum(5)
            self.a15.setFixedWidth(100)
            self.layout.addWidget(self.a15,15,3)
            self.a15.valueChanged.connect(self.spinbox_toggled2)

        elif (self.partb==16):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

            self.a10  = QSpinBox()
            #self.a10.setSuffix("  subdivions")
            self.a10.setMinimum(1)
            self.a10.setMaximum(5)
            self.a10.setFixedWidth(100)
            self.layout.addWidget(self.a10,10,3)
            self.a10.valueChanged.connect(self.spinbox_toggled2)

            self.a11  = QSpinBox()
            #self.a11.setSuffix("  subdivions")
            self.a11.setMinimum(1)
            self.a11.setMaximum(5)
            self.a11.setFixedWidth(100)
            self.layout.addWidget(self.a11,11,3)
            self.a11.valueChanged.connect(self.spinbox_toggled2)

            self.a12  = QSpinBox()
            #self.a12.setSuffix("  subdivions")
            self.a12.setMinimum(1)
            self.a12.setMaximum(5)
            self.a12.setFixedWidth(100)
            self.layout.addWidget(self.a12,12,3)
            self.a12.valueChanged.connect(self.spinbox_toggled2)

            self.a13  = QSpinBox()
            #self.a13.setSuffix("  subdivions")
            self.a13.setMinimum(1)
            self.a13.setMaximum(5)
            self.a13.setFixedWidth(100)
            self.layout.addWidget(self.a13,13,3)
            self.a13.valueChanged.connect(self.spinbox_toggled2)

            self.a14  = QSpinBox()
            #self.a14.setSuffix("  subdivions")
            self.a14.setMinimum(1)
            self.a14.setMaximum(5)
            self.a14.setFixedWidth(100)
            self.layout.addWidget(self.a14,14,3)
            self.a14.valueChanged.connect(self.spinbox_toggled2)
            self.a15  = QSpinBox()
            #self.a15.setSuffix("  subdivions")
            self.a15.setMinimum(1)
            self.a15.setMaximum(5)
            self.a15.setFixedWidth(100)
            self.layout.addWidget(self.a15,15,3)
            self.a15.valueChanged.connect(self.spinbox_toggled2)

            self.a16  = QSpinBox()
            #self.a16.setSuffix("  subdivions")
            self.a16.setMinimum(1)
            self.a16.setMaximum(5)
            self.a16.setFixedWidth(100)
            self.layout.addWidget(self.a16,16,3)
            self.a16.valueChanged.connect(self.spinbox_toggled2)

        elif (self.partb==17):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

            self.a10  = QSpinBox()
            #self.a10.setSuffix("  subdivions")
            self.a10.setMinimum(1)
            self.a10.setMaximum(5)
            self.a10.setFixedWidth(100)
            self.layout.addWidget(self.a10,10,3)
            self.a10.valueChanged.connect(self.spinbox_toggled2)

            self.a11  = QSpinBox()
            #self.a11.setSuffix("  subdivions")
            self.a11.setMinimum(1)
            self.a11.setMaximum(5)
            self.a11.setFixedWidth(100)
            self.layout.addWidget(self.a11,11,3)
            self.a11.valueChanged.connect(self.spinbox_toggled2)

            self.a12  = QSpinBox()
            #self.a12.setSuffix("  subdivions")
            self.a12.setMinimum(1)
            self.a12.setMaximum(5)
            self.a12.setFixedWidth(100)
            self.layout.addWidget(self.a12,12,3)
            self.a12.valueChanged.connect(self.spinbox_toggled2)

            self.a13  = QSpinBox()
            #self.a13.setSuffix("  subdivions")
            self.a13.setMinimum(1)
            self.a13.setMaximum(5)
            self.a13.setFixedWidth(100)
            self.layout.addWidget(self.a13,13,3)
            self.a13.valueChanged.connect(self.spinbox_toggled2)

            self.a14  = QSpinBox()
            #self.a14.setSuffix("  subdivions")
            self.a14.setMinimum(1)
            self.a14.setMaximum(5)
            self.a14.setFixedWidth(100)
            self.layout.addWidget(self.a14,14,3)
            self.a14.valueChanged.connect(self.spinbox_toggled2)
            self.a15  = QSpinBox()
            #self.a15.setSuffix("  subdivions")
            self.a15.setMinimum(1)
            self.a15.setMaximum(5)
            self.a15.setFixedWidth(100)
            self.layout.addWidget(self.a15,15,3)
            self.a15.valueChanged.connect(self.spinbox_toggled2)

            self.a16  = QSpinBox()
            #self.a16.setSuffix("  subdivions")
            self.a16.setMinimum(1)
            self.a16.setMaximum(5)
            self.a16.setFixedWidth(100)
            self.layout.addWidget(self.a16,16,3)
            self.a16.valueChanged.connect(self.spinbox_toggled2)

            self.a17  = QSpinBox()
            #self.a17.setSuffix("  subdivions")
            self.a17.setMinimum(1)
            self.a17.setMaximum(5)
            self.a17.setFixedWidth(100)
            self.layout.addWidget(self.a17,17,3)
            self.a17.valueChanged.connect(self.spinbox_toggled2)

           
        elif (self.partb==18):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

            self.a10  = QSpinBox()
            #self.a10.setSuffix("  subdivions")
            self.a10.setMinimum(1)
            self.a10.setMaximum(5)
            self.a10.setFixedWidth(100)
            self.layout.addWidget(self.a10,10,3)
            self.a10.valueChanged.connect(self.spinbox_toggled2)

            self.a11  = QSpinBox()
            #self.a11.setSuffix("  subdivions")
            self.a11.setMinimum(1)
            self.a11.setMaximum(5)
            self.a11.setFixedWidth(100)
            self.layout.addWidget(self.a11,11,3)
            self.a11.valueChanged.connect(self.spinbox_toggled2)

            self.a12  = QSpinBox()
            #self.a12.setSuffix("  subdivions")
            self.a12.setMinimum(1)
            self.a12.setMaximum(5)
            self.a12.setFixedWidth(100)
            self.layout.addWidget(self.a12,12,3)
            self.a12.valueChanged.connect(self.spinbox_toggled2)

            self.a13  = QSpinBox()
            #self.a13.setSuffix("  subdivions")
            self.a13.setMinimum(1)
            self.a13.setMaximum(5)
            self.a13.setFixedWidth(100)
            self.layout.addWidget(self.a13,13,3)
            self.a13.valueChanged.connect(self.spinbox_toggled2)

            self.a14  = QSpinBox()
            #self.a14.setSuffix("  subdivions")
            self.a14.setMinimum(1)
            self.a14.setMaximum(5)
            self.a14.setFixedWidth(100)
            self.layout.addWidget(self.a14,14,3)
            self.a14.valueChanged.connect(self.spinbox_toggled2)
            self.a15  = QSpinBox()
            #self.a15.setSuffix("  subdivions")
            self.a15.setMinimum(1)
            self.a15.setMaximum(5)
            self.a15.setFixedWidth(100)
            self.layout.addWidget(self.a15,15,3)
            self.a15.valueChanged.connect(self.spinbox_toggled2)

            self.a16  = QSpinBox()
            #self.a16.setSuffix("  subdivions")
            self.a16.setMinimum(1)
            self.a16.setMaximum(5)
            self.a16.setFixedWidth(100)
            self.layout.addWidget(self.a16,16,3)
            self.a16.valueChanged.connect(self.spinbox_toggled2)

            self.a17  = QSpinBox()
            #self.a17.setSuffix("  subdivions")
            self.a17.setMinimum(1)
            self.a17.setMaximum(5)
            self.a17.setFixedWidth(100)
            self.layout.addWidget(self.a17,17,3)
            self.a17.valueChanged.connect(self.spinbox_toggled2)

            self.a18  = QSpinBox()
            #self.a18.setSuffix("  subdivions")
            self.a18.setMinimum(1)
            self.a18.setMaximum(5)
            self.a18.setFixedWidth(100)
            self.layout.addWidget(self.a18,18,3)
            self.a18.valueChanged.connect(self.spinbox_toggled2)

           
        elif (self.partb==19):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

            self.a10  = QSpinBox()
            #self.a10.setSuffix("  subdivions")
            self.a10.setMinimum(1)
            self.a10.setMaximum(5)
            self.a10.setFixedWidth(100)
            self.layout.addWidget(self.a10,10,3)
            self.a10.valueChanged.connect(self.spinbox_toggled2)

            self.a11  = QSpinBox()
            #self.a11.setSuffix("  subdivions")
            self.a11.setMinimum(1)
            self.a11.setMaximum(5)
            self.a11.setFixedWidth(100)
            self.layout.addWidget(self.a11,11,3)
            self.a11.valueChanged.connect(self.spinbox_toggled2)

            self.a12  = QSpinBox()
            #self.a12.setSuffix("  subdivions")
            self.a12.setMinimum(1)
            self.a12.setMaximum(5)
            self.a12.setFixedWidth(100)
            self.layout.addWidget(self.a12,12,3)
            self.a12.valueChanged.connect(self.spinbox_toggled2)

            self.a13  = QSpinBox()
            #self.a13.setSuffix("  subdivions")
            self.a13.setMinimum(1)
            self.a13.setMaximum(5)
            self.a13.setFixedWidth(100)
            self.layout.addWidget(self.a13,13,3)
            self.a13.valueChanged.connect(self.spinbox_toggled2)

            self.a14  = QSpinBox()
            #self.a14.setSuffix("  subdivions")
            self.a14.setMinimum(1)
            self.a14.setMaximum(5)
            self.a14.setFixedWidth(100)
            self.layout.addWidget(self.a14,14,3)
            self.a14.valueChanged.connect(self.spinbox_toggled2)
            self.a15  = QSpinBox()
            #self.a15.setSuffix("  subdivions")
            self.a15.setMinimum(1)
            self.a15.setMaximum(5)
            self.a15.setFixedWidth(100)
            self.layout.addWidget(self.a15,15,3)
            self.a15.valueChanged.connect(self.spinbox_toggled2)

            self.a16  = QSpinBox()
            #self.a16.setSuffix("  subdivions")
            self.a16.setMinimum(1)
            self.a16.setMaximum(5)
            self.a16.setFixedWidth(100)
            self.layout.addWidget(self.a16,16,3)
            self.a16.valueChanged.connect(self.spinbox_toggled2)

            self.a17  = QSpinBox()
            #self.a17.setSuffix("  subdivions")
            self.a17.setMinimum(1)
            self.a17.setMaximum(5)
            self.a17.setFixedWidth(100)
            self.layout.addWidget(self.a17,17,3)
            self.a17.valueChanged.connect(self.spinbox_toggled2)

            self.a18  = QSpinBox()
            #self.a18.setSuffix("  subdivions")
            self.a18.setMinimum(1)
            self.a18.setMaximum(5)
            self.a18.setFixedWidth(100)
            self.layout.addWidget(self.a18,18,3)
            self.a18.valueChanged.connect(self.spinbox_toggled2)

            self.a19  = QSpinBox()
            #self.a19.setSuffix("  subdivions")
            self.a19.setMinimum(1)
            self.a19.setMaximum(5)
            self.a19.setFixedWidth(100)
            self.layout.addWidget(self.a19,19,3)
            self.a19.valueChanged.connect(self.spinbox_toggled2)

            
        elif (self.partb==20):
            self.a1  = QSpinBox()
            #self.a1.setSuffix("  subdivions")
            self.a1.setMinimum(1)
            self.a1.setMaximum(5)
            self.a1.setFixedWidth(100)
            self.layout.addWidget(self.a1,1,3)
            self.a1.valueChanged.connect(self.spinbox_toggled2)

            self.a2  = QSpinBox()
            #self.a2.setSuffix("  subdivions")
            self.a2.setMinimum(1)
            self.a2.setMaximum(5)
            self.a2.setFixedWidth(100)
            self.layout.addWidget(self.a2,2,3)
            self.a2.valueChanged.connect(self.spinbox_toggled2)

            self.a3  = QSpinBox()
            #self.a3.setSuffix("  subdivions")
            self.a3.setMinimum(1)
            self.a3.setMaximum(5)
            self.a3.setFixedWidth(100)
            self.layout.addWidget(self.a3,3,3)
            self.a3.valueChanged.connect(self.spinbox_toggled2)

            self.a4  = QSpinBox()
            #self.a4.setSuffix("  subdivions")
            self.a4.setMinimum(1)
            self.a4.setMaximum(5)
            self.a4.setFixedWidth(100)
            self.layout.addWidget(self.a4,4,3)
            self.a4.valueChanged.connect(self.spinbox_toggled2)

            self.a5  = QSpinBox()
            #self.a5.setSuffix("  subdivions")
            self.a5.setMinimum(1)
            self.a5.setMaximum(5)
            self.a5.setFixedWidth(100)
            self.layout.addWidget(self.a5,5,3)
            self.a5.valueChanged.connect(self.spinbox_toggled2)

            self.a6  = QSpinBox()
            #self.a6.setSuffix("  subdivions")
            self.a6.setMinimum(1)
            self.a6.setMaximum(5)
            self.a6.setFixedWidth(100)
            self.layout.addWidget(self.a6,6,3)
            self.a6.valueChanged.connect(self.spinbox_toggled2)

            self.a7  = QSpinBox()
            #self.a7.setSuffix("  subdivions")
            self.a7.setMinimum(1)
            self.a7.setMaximum(5)
            self.a7.setFixedWidth(100)
            self.layout.addWidget(self.a7,7,3)
            self.a7.valueChanged.connect(self.spinbox_toggled2)

            self.a8  = QSpinBox()
            #self.a8.setSuffix("  subdivions")
            self.a8.setMinimum(1)
            self.a8.setMaximum(5)
            self.a8.setFixedWidth(100)
            self.layout.addWidget(self.a8,8,3)
            self.a8.valueChanged.connect(self.spinbox_toggled2)

            self.a9  = QSpinBox()
            #self.a9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.a9.setMaximum(5)
            self.a9.setFixedWidth(100)
            self.layout.addWidget(self.a9,9,3)
            self.a9.valueChanged.connect(self.spinbox_toggled2)

            self.a10  = QSpinBox()
            #self.a10.setSuffix("  subdivions")
            self.a10.setMinimum(1)
            self.a10.setMaximum(5)
            self.a10.setFixedWidth(100)
            self.layout.addWidget(self.a10,10,3)
            self.a10.valueChanged.connect(self.spinbox_toggled2)

            self.a11  = QSpinBox()
            #self.a11.setSuffix("  subdivions")
            self.a11.setMinimum(1)
            self.a11.setMaximum(5)
            self.a11.setFixedWidth(100)
            self.layout.addWidget(self.a11,11,3)
            self.a11.valueChanged.connect(self.spinbox_toggled2)

            self.a12  = QSpinBox()
            #self.a12.setSuffix("  subdivions")
            self.a12.setMinimum(1)
            self.a12.setMaximum(5)
            self.a12.setFixedWidth(100)
            self.layout.addWidget(self.a12,12,3)
            self.a12.valueChanged.connect(self.spinbox_toggled2)

            self.a13  = QSpinBox()
            #self.a13.setSuffix("  subdivions")
            self.a13.setMinimum(1)
            self.a13.setMaximum(5)
            self.a13.setFixedWidth(100)
            self.layout.addWidget(self.a13,13,3)
            self.a13.valueChanged.connect(self.spinbox_toggled2)

            self.a14  = QSpinBox()
            #self.a14.setSuffix("  subdivions")
            self.a14.setMinimum(1)
            self.a14.setMaximum(5)
            self.a14.setFixedWidth(100)
            self.layout.addWidget(self.a14,14,3)
            self.a14.valueChanged.connect(self.spinbox_toggled2)
            self.a15  = QSpinBox()
            #self.a15.setSuffix("  subdivions")
            self.a15.setMinimum(1)
            self.a15.setMaximum(5)
            self.a15.setFixedWidth(100)
            self.layout.addWidget(self.a15,15,3)
            self.a15.valueChanged.connect(self.spinbox_toggled2)

            self.a16  = QSpinBox()
            #self.a16.setSuffix("  subdivions")
            self.a16.setMinimum(1)
            self.a16.setMaximum(5)
            self.a16.setFixedWidth(100)
            self.layout.addWidget(self.a16,16,3)
            self.a16.valueChanged.connect(self.spinbox_toggled2)

            self.a17  = QSpinBox()
            #self.a17.setSuffix("  subdivions")
            self.a17.setMinimum(1)
            self.a17.setMaximum(5)
            self.a17.setFixedWidth(100)
            self.layout.addWidget(self.a17,17,3)
            self.a17.valueChanged.connect(self.spinbox_toggled2)

            self.a18  = QSpinBox()
            #self.a18.setSuffix("  subdivions")
            self.a18.setMinimum(1)
            self.a18.setMaximum(5)
            self.a18.setFixedWidth(100)
            self.layout.addWidget(self.a18,18,3)
            self.a18.valueChanged.connect(self.spinbox_toggled2)

            self.a19  = QSpinBox()
            #self.a19.setSuffix("  subdivions")
            self.a19.setMinimum(1)
            self.a19.setMaximum(5)
            self.a19.setFixedWidth(100)
            self.layout.addWidget(self.a19,19,3)
            self.a19.valueChanged.connect(self.spinbox_toggled2)

            self.a20  = QSpinBox()
            #self.a20.setSuffix("  subdivions")
            self.a20.setMinimum(1)
            self.a20.setMaximum(5)
            self.a20.setFixedWidth(100)
            self.layout.addWidget(self.a20,20,3)
            self.a20.valueChanged.connect(self.spinbox_toggled2)

            
                
            
        
    
    def spinbox_toggled2(self):
        
        self.partbsub=[]
        if (self.partb==0):
            self.partbsub=[0]
        elif (self.partb==1):
            self.qsb1=self.a1.value()
            self.partbsub.append(self.qsb1)
            print(self.partbsub)
        elif (self.partb==2):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)

            print(self.partbsub)
        elif (self.partb==3):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            print(self.partbsub)
            
        elif (self.partb==4):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            print(self.partbsub)
            
        elif (self.partb==5):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            print(self.partbsub)
            
        elif (self.partb==6):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            print(self.partbsub)
            
        elif (self.partb==7):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            print(self.partbsub)
           
        elif (self.partb==8):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            print(self.partbsub)
            
        elif (self.partb==9):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            print(self.partbsub)
            
        elif (self.partb==10):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.qsb10=self.a10.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            self.partbsub.append(self.qsb10)
            print(self.partbsub)

        elif (self.partb==11):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.qsb10=self.a10.value()
            self.qsb11=self.a11.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            self.partbsub.append(self.qsb10)
            self.partbsub.append(self.qsb11)
            print(self.partbsub)
            
        elif (self.partb==12):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.qsb10=self.a10.value()
            self.qsb11=self.a11.value()
            self.qsb12=self.a12.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            self.partbsub.append(self.qsb10)
            self.partbsub.append(self.qsb11)
            self.partbsub.append(self.qsb12)
            print(self.partbsub)
            
        elif (self.partb==13):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.qsb10=self.a10.value()
            self.qsb11=self.a11.value()
            self.qsb12=self.a12.value()
            self.qsb13=self.a13.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            self.partbsub.append(self.qsb10)
            self.partbsub.append(self.qsb11)
            self.partbsub.append(self.qsb12)
            self.partbsub.append(self.qsb13)
            print(self.partbsub)
            
        elif (self.partb==14):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.qsb10=self.a10.value()
            self.qsb11=self.a11.value()
            self.qsb12=self.a12.value()
            self.qsb13=self.a13.value()
            self.qsb14=self.a14.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            self.partbsub.append(self.qsb10)
            self.partbsub.append(self.qsb11)
            self.partbsub.append(self.qsb12)
            self.partbsub.append(self.qsb13)
            self.partbsub.append(self.qsb14)
            print(self.partbsub)
            
        elif (self.partb==15):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.qsb10=self.a10.value()
            self.qsb11=self.a11.value()
            self.qsb12=self.a12.value()
            self.qsb13=self.a13.value()
            self.qsb14=self.a14.value()
            self.qsb15=self.a15.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            self.partbsub.append(self.qsb10)
            self.partbsub.append(self.qsb11)
            self.partbsub.append(self.qsb12)
            self.partbsub.append(self.qsb13)
            self.partbsub.append(self.qsb14)
            self.partbsub.append(self.qsb15)
            print(self.partbsub)
            
        elif (self.partb==16):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.qsb10=self.a10.value()
            self.qsb11=self.a11.value()
            self.qsb12=self.a12.value()
            self.qsb13=self.a13.value()
            self.qsb14=self.a14.value()
            self.qsb15=self.a15.value()
            self.qsb16=self.a16.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            self.partbsub.append(self.qsb10)
            self.partbsub.append(self.qsb11)
            self.partbsub.append(self.qsb12)
            self.partbsub.append(self.qsb13)
            self.partbsub.append(self.qsb14)
            self.partbsub.append(self.qsb15)
            self.partbsub.append(self.qsb16)
            print(self.partbsub)
            
        elif (self.partb==17):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.qsb10=self.a10.value()
            self.qsb11=self.a11.value()
            self.qsb12=self.a12.value()
            self.qsb13=self.a13.value()
            self.qsb14=self.a14.value()
            self.qsb15=self.a15.value()
            self.qsb16=self.a16.value()
            self.qsb17=self.a17.value()
            self.qsb18=self.a18.value()
            self.qsb19=self.a19.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            self.partbsub.append(self.qsb10)
            self.partbsub.append(self.qsb11)
            self.partbsub.append(self.qsb12)
            self.partbsub.append(self.qsb13)
            self.partbsub.append(self.qsb14)
            self.partbsub.append(self.qsb15)
            self.partbsub.append(self.qsb16)
            self.partbsub.append(self.qsb17)
            print(self.partbsub)

        elif (self.partb==18):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.qsb10=self.a10.value()
            self.qsb11=self.a11.value()
            self.qsb12=self.a12.value()
            self.qsb13=self.a13.value()
            self.qsb14=self.a14.value()
            self.qsb15=self.a15.value()
            self.qsb16=self.a16.value()
            self.qsb17=self.a17.value()
            self.qsb18=self.a18.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            self.partbsub.append(self.qsb10)
            self.partbsub.append(self.qsb11)
            self.partbsub.append(self.qsb12)
            self.partbsub.append(self.qsb13)
            self.partbsub.append(self.qsb14)
            self.partbsub.append(self.qsb15)
            self.partbsub.append(self.qsb16)
            self.partbsub.append(self.qsb17)
            self.partbsub.append(self.qsb18)
            print(self.partbsub)

        elif (self.partb==19):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.qsb10=self.a10.value()
            self.qsb11=self.a11.value()
            self.qsb12=self.a12.value()
            self.qsb13=self.a13.value()
            self.qsb14=self.a14.value()
            self.qsb15=self.a15.value()
            self.qsb16=self.a16.value()
            self.qsb17=self.a17.value()
            self.qsb18=self.a18.value()
            self.qsb19=self.a19.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            self.partbsub.append(self.qsb10)
            self.partbsub.append(self.qsb11)
            self.partbsub.append(self.qsb12)
            self.partbsub.append(self.qsb13)
            self.partbsub.append(self.qsb14)
            self.partbsub.append(self.qsb15)
            self.partbsub.append(self.qsb16)
            self.partbsub.append(self.qsb17)
            self.partbsub.append(self.qsb18)
            self.partbsub.append(self.qsb19)
            print(self.partbsub)

        elif (self.partb==20):
            self.qsb1=self.a1.value()
            self.qsb2=self.a2.value()
            self.qsb3=self.a3.value()
            self.qsb4=self.a4.value()
            self.qsb5=self.a5.value()
            self.qsb6=self.a6.value()
            self.qsb7=self.a7.value()
            self.qsb8=self.a8.value()
            self.qsb9=self.a9.value()
            self.qsb10=self.a10.value()
            self.qsb11=self.a11.value()
            self.qsb12=self.a12.value()
            self.qsb13=self.a13.value()
            self.qsb14=self.a14.value()
            self.qsb15=self.a15.value()
            self.qsb16=self.a16.value()
            self.qsb17=self.a17.value()
            self.qsb18=self.a18.value()
            self.qsb19=self.a19.value()
            self.qsb20=self.a20.value()
            self.partbsub.append(self.qsb1)
            self.partbsub.append(self.qsb2)
            self.partbsub.append(self.qsb3)
            self.partbsub.append(self.qsb4)
            self.partbsub.append(self.qsb5)
            self.partbsub.append(self.qsb6)
            self.partbsub.append(self.qsb7)
            self.partbsub.append(self.qsb8)
            self.partbsub.append(self.qsb9)
            self.partbsub.append(self.qsb10)
            self.partbsub.append(self.qsb11)
            self.partbsub.append(self.qsb12)
            self.partbsub.append(self.qsb13)
            self.partbsub.append(self.qsb14)
            self.partbsub.append(self.qsb15)
            self.partbsub.append(self.qsb16)
            self.partbsub.append(self.qsb17)
            self.partbsub.append(self.qsb18)
            self.partbsub.append(self.qsb19)
            self.partbsub.append(self.qsb20)
            
            print(self.partbsub)
        global partbsub
        partbsub=self.partbsub
    def courseb(self):
        if (self.partb==1):
            
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            
 
            
        elif (self.partb==2):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            
        elif (self.partb==3):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            
        elif (self.partb==4):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            
        elif (self.partb==5):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
        elif (self.partb==6):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            
        elif (self.partb==7):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
        elif (self.partb==8):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

        
        elif (self.partb==9):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,4)
            self.copbq9.currentTextChanged.connect(self.coseeb)

        elif (self.partb==10):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,4)
            self.copbq9.currentTextChanged.connect(self.coseeb)

            self.copbq10= QComboBox()
            #self.copbq10.addItem("CO")
            self.copbq10.addItem("CO1")
            self.copbq10.addItem("CO2")
            self.copbq10.addItem("CO3")
            self.copbq10.addItem("CO4")
            self.copbq10.addItem("CO5")
            self.layout.addWidget(self.copbq10,10,4)
            self.copbq10.currentTextChanged.connect(self.coseeb)

        elif (self.partb==11):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,3)
            self.copbq9.currentTextChanged.connect(self.coseeb)

            self.copbq10= QComboBox()
            #self.copbq10.addItem("CO")
            self.copbq10.addItem("CO1")
            self.copbq10.addItem("CO2")
            self.copbq10.addItem("CO3")
            self.copbq10.addItem("CO4")
            self.copbq10.addItem("CO5")
            self.layout.addWidget(self.copbq10,10,4)
            self.copbq10.currentTextChanged.connect(self.coseeb)

            self.copbq11= QComboBox()
            #self.copbq11.addItem("CO")
            self.copbq11.addItem("CO1")
            self.copbq11.addItem("CO2")
            self.copbq11.addItem("CO3")
            self.copbq11.addItem("CO4")
            self.copbq11.addItem("CO5")
            self.layout.addWidget(self.copbq11,11,4)
            self.copbq11.currentTextChanged.connect(self.coseeb)

        elif (self.partb==12):
            
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,4)
            self.copbq9.currentTextChanged.connect(self.coseeb)

            self.copbq10= QComboBox()
            #self.copbq10.addItem("CO")
            self.copbq10.addItem("CO1")
            self.copbq10.addItem("CO2")
            self.copbq10.addItem("CO3")
            self.copbq10.addItem("CO4")
            self.copbq10.addItem("CO5")
            self.layout.addWidget(self.copbq10,10,4)
            self.copbq10.currentTextChanged.connect(self.coseeb)

            self.copbq11= QComboBox()
            #self.copbq11.addItem("CO")
            self.copbq11.addItem("CO1")
            self.copbq11.addItem("CO2")
            self.copbq11.addItem("CO3")
            self.copbq11.addItem("CO4")
            self.copbq11.addItem("CO5")
            self.layout.addWidget(self.copbq11,11,4)
            self.copbq11.currentTextChanged.connect(self.coseeb)

            self.copbq12= QComboBox()
            #self.copbq12.addItem("CO")
            self.copbq12.addItem("CO1")
            self.copbq12.addItem("CO2")
            self.copbq12.addItem("CO3")
            self.copbq12.addItem("CO4")
            self.copbq12.addItem("CO5")
            self.layout.addWidget(self.copbq12,12,4)
            self.copbq12.currentTextChanged.connect(self.coseeb)

        elif (self.partb==13):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,4)
            self.copbq9.currentTextChanged.connect(self.coseeb)

            self.copbq10= QComboBox()
            #self.copbq10.addItem("CO")
            self.copbq10.addItem("CO1")
            self.copbq10.addItem("CO2")
            self.copbq10.addItem("CO3")
            self.copbq10.addItem("CO4")
            self.copbq10.addItem("CO5")
            self.layout.addWidget(self.copbq10,10,4)
            self.copbq10.currentTextChanged.connect(self.coseeb)

            self.copbq11= QComboBox()
            #self.copbq11.addItem("CO")
            self.copbq11.addItem("CO1")
            self.copbq11.addItem("CO2")
            self.copbq11.addItem("CO3")
            self.copbq11.addItem("CO4")
            self.copbq11.addItem("CO5")
            self.layout.addWidget(self.copbq11,11,4)
            self.copbq11.currentTextChanged.connect(self.coseeb)

            self.copbq12= QComboBox()
            #self.copbq12.addItem("CO")
            self.copbq12.addItem("CO1")
            self.copbq12.addItem("CO2")
            self.copbq12.addItem("CO3")
            self.copbq12.addItem("CO4")
            self.copbq12.addItem("CO5")
            self.layout.addWidget(self.copbq12,12,4)
            self.copbq12.currentTextChanged.connect(self.coseeb)

            self.copbq13= QComboBox()
            #self.copbq13.addItem("CO")
            self.copbq13.addItem("CO1")
            self.copbq13.addItem("CO2")
            self.copbq13.addItem("CO3")
            self.copbq13.addItem("CO4")
            self.copbq13.addItem("CO5")
            self.layout.addWidget(self.copbq13,13,4)
            self.copbq13.currentTextChanged.connect(self.coseeb)

           

            
        elif (self.partb==14):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,3)
            self.copbq9.currentTextChanged.connect(self.coseeb)

            self.copbq10= QComboBox()
            #self.copbq10.addItem("CO")
            self.copbq10.addItem("CO1")
            self.copbq10.addItem("CO2")
            self.copbq10.addItem("CO3")
            self.copbq10.addItem("CO4")
            self.copbq10.addItem("CO5")
            self.layout.addWidget(self.copbq10,10,4)
            self.copbq10.currentTextChanged.connect(self.coseeb)

            self.copbq11= QComboBox()
            #self.copbq11.addItem("CO")
            self.copbq11.addItem("CO1")
            self.copbq11.addItem("CO2")
            self.copbq11.addItem("CO3")
            self.copbq11.addItem("CO4")
            self.copbq11.addItem("CO5")
            self.layout.addWidget(self.copbq11,11,4)
            self.copbq11.currentTextChanged.connect(self.coseeb)

            self.copbq12= QComboBox()
            #self.copbq12.addItem("CO")
            self.copbq12.addItem("CO1")
            self.copbq12.addItem("CO2")
            self.copbq12.addItem("CO3")
            self.copbq12.addItem("CO4")
            self.copbq12.addItem("CO5")
            self.layout.addWidget(self.copbq12,12,4)
            self.copbq12.currentTextChanged.connect(self.coseeb)

            self.copbq13= QComboBox()
            #self.copbq13.addItem("CO")
            self.copbq13.addItem("CO1")
            self.copbq13.addItem("CO2")
            self.copbq13.addItem("CO3")
            self.copbq13.addItem("CO4")
            self.copbq13.addItem("CO5")
            self.layout.addWidget(self.copbq13,13,4)
            self.copbq13.currentTextChanged.connect(self.coseeb)

            self.copbq14= QComboBox()
            #self.copbq14.addItem("CO")
            self.copbq14.addItem("CO1")
            self.copbq14.addItem("CO2")
            self.copbq14.addItem("CO3")
            self.copbq14.addItem("CO4")
            self.copbq14.addItem("CO5")
            self.layout.addWidget(self.copbq14,14,4)
            self.copbq14.currentTextChanged.connect(self.coseeb)

           
        elif (self.partb==15):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,4)
            self.copbq9.currentTextChanged.connect(self.coseeb)

            self.copbq10= QComboBox()
            #self.copbq10.addItem("CO")
            self.copbq10.addItem("CO1")
            self.copbq10.addItem("CO2")
            self.copbq10.addItem("CO3")
            self.copbq10.addItem("CO4")
            self.copbq10.addItem("CO5")
            self.layout.addWidget(self.copbq10,10,4)
            self.copbq10.currentTextChanged.connect(self.coseeb)

            self.copbq11= QComboBox()
            #self.copbq11.addItem("CO")
            self.copbq11.addItem("CO1")
            self.copbq11.addItem("CO2")
            self.copbq11.addItem("CO3")
            self.copbq11.addItem("CO4")
            self.copbq11.addItem("CO5")
            self.layout.addWidget(self.copbq11,11,4)
            self.copbq11.currentTextChanged.connect(self.coseeb)

            self.copbq12= QComboBox()
            #self.copbq12.addItem("CO")
            self.copbq12.addItem("CO1")
            self.copbq12.addItem("CO2")
            self.copbq12.addItem("CO3")
            self.copbq12.addItem("CO4")
            self.copbq12.addItem("CO5")
            self.layout.addWidget(self.copbq12,12,4)
            self.copbq12.currentTextChanged.connect(self.coseeb)

            self.copbq13= QComboBox()
            #self.copbq13.addItem("CO")
            self.copbq13.addItem("CO1")
            self.copbq13.addItem("CO2")
            self.copbq13.addItem("CO3")
            self.copbq13.addItem("CO4")
            self.copbq13.addItem("CO5")
            self.layout.addWidget(self.copbq13,13,4)
            self.copbq13.currentTextChanged.connect(self.coseeb)

            self.copbq14= QComboBox()
            #self.copbq14.addItem("CO")
            self.copbq14.addItem("CO1")
            self.copbq14.addItem("CO2")
            self.copbq14.addItem("CO3")
            self.copbq14.addItem("CO4")
            self.copbq14.addItem("CO5")
            self.layout.addWidget(self.copbq14,14,4)
            self.copbq14.currentTextChanged.connect(self.coseeb)

            self.copbq15= QComboBox()
            #self.copbq15.addItem("CO")
            self.copbq15.addItem("CO1")
            self.copbq15.addItem("CO2")
            self.copbq15.addItem("CO3")
            self.copbq15.addItem("CO4")
            self.copbq15.addItem("CO5")
            self.layout.addWidget(self.copbq15,15,4)
            self.copbq15.currentTextChanged.connect(self.coseeb)

            
        elif (self.partb==16):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,4)
            self.copbq9.currentTextChanged.connect(self.coseeb)

            self.copbq10= QComboBox()
            #self.copbq10.addItem("CO")
            self.copbq10.addItem("CO1")
            self.copbq10.addItem("CO2")
            self.copbq10.addItem("CO3")
            self.copbq10.addItem("CO4")
            self.copbq10.addItem("CO5")
            self.layout.addWidget(self.copbq10,10,4)
            self.copbq10.currentTextChanged.connect(self.coseeb)

            self.copbq11= QComboBox()
            #self.copbq11.addItem("CO")
            self.copbq11.addItem("CO1")
            self.copbq11.addItem("CO2")
            self.copbq11.addItem("CO3")
            self.copbq11.addItem("CO4")
            self.copbq11.addItem("CO5")
            self.layout.addWidget(self.copbq11,11,4)
            self.copbq11.currentTextChanged.connect(self.coseeb)

            self.copbq12= QComboBox()
            #self.copbq12.addItem("CO")
            self.copbq12.addItem("CO1")
            self.copbq12.addItem("CO2")
            self.copbq12.addItem("CO3")
            self.copbq12.addItem("CO4")
            self.copbq12.addItem("CO5")
            self.layout.addWidget(self.copbq12,12,4)
            self.copbq12.currentTextChanged.connect(self.coseeb)

            self.copbq13= QComboBox()
            #self.copbq13.addItem("CO")
            self.copbq13.addItem("CO1")
            self.copbq13.addItem("CO2")
            self.copbq13.addItem("CO3")
            self.copbq13.addItem("CO4")
            self.copbq13.addItem("CO5")
            self.layout.addWidget(self.copbq13,13,4)
            self.copbq13.currentTextChanged.connect(self.coseeb)

            self.copbq14= QComboBox()
            #self.copbq14.addItem("CO")
            self.copbq14.addItem("CO1")
            self.copbq14.addItem("CO2")
            self.copbq14.addItem("CO3")
            self.copbq14.addItem("CO4")
            self.copbq14.addItem("CO5")
            self.layout.addWidget(self.copbq14,14,3)
            self.copbq14.currentTextChanged.connect(self.coseeb)

            self.copbq15= QComboBox()
            #self.copbq15.addItem("CO")
            self.copbq15.addItem("CO1")
            self.copbq15.addItem("CO2")
            self.copbq15.addItem("CO3")
            self.copbq15.addItem("CO4")
            self.copbq15.addItem("CO5")
            self.layout.addWidget(self.copbq15,15,4)
            self.copbq15.currentTextChanged.connect(self.coseeb)

            self.copbq16= QComboBox()
            #self.copbq16.addItem("CO")
            self.copbq16.addItem("CO1")
            self.copbq16.addItem("CO2")
            self.copbq16.addItem("CO3")
            self.copbq16.addItem("CO4")
            self.copbq16.addItem("CO5")
            self.layout.addWidget(self.copbq16,16,4)
            self.copbq16.currentTextChanged.connect(self.coseeb)

            
        elif (self.partb==17):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,4)
            self.copbq9.currentTextChanged.connect(self.coseeb)

            self.copbq10= QComboBox()
            #self.copbq10.addItem("CO")
            self.copbq10.addItem("CO1")
            self.copbq10.addItem("CO2")
            self.copbq10.addItem("CO3")
            self.copbq10.addItem("CO4")
            self.copbq10.addItem("CO5")
            self.layout.addWidget(self.copbq10,10,4)
            self.copbq10.currentTextChanged.connect(self.coseeb)

            self.copbq11= QComboBox()
            #self.copbq11.addItem("CO")
            self.copbq11.addItem("CO1")
            self.copbq11.addItem("CO2")
            self.copbq11.addItem("CO3")
            self.copbq11.addItem("CO4")
            self.copbq11.addItem("CO5")
            self.layout.addWidget(self.copbq11,11,4)
            self.copbq11.currentTextChanged.connect(self.coseeb)

            self.copbq12= QComboBox()
            #self.copbq12.addItem("CO")
            self.copbq12.addItem("CO1")
            self.copbq12.addItem("CO2")
            self.copbq12.addItem("CO3")
            self.copbq12.addItem("CO4")
            self.copbq12.addItem("CO5")
            self.layout.addWidget(self.copbq12,12,4)
            self.copbq12.currentTextChanged.connect(self.coseeb)

            self.copbq13= QComboBox()
            #self.copbq13.addItem("CO")
            self.copbq13.addItem("CO1")
            self.copbq13.addItem("CO2")
            self.copbq13.addItem("CO3")
            self.copbq13.addItem("CO4")
            self.copbq13.addItem("CO5")
            self.layout.addWidget(self.copbq13,13,4)
            self.copbq13.currentTextChanged.connect(self.coseeb)

            self.copbq14= QComboBox()
            #self.copbq14.addItem("CO")
            self.copbq14.addItem("CO1")
            self.copbq14.addItem("CO2")
            self.copbq14.addItem("CO3")
            self.copbq14.addItem("CO4")
            self.copbq14.addItem("CO5")
            self.layout.addWidget(self.copbq14,14,4)
            self.copbq14.currentTextChanged.connect(self.coseeb)

            self.copbq15= QComboBox()
            #self.copbq15.addItem("CO")
            self.copbq15.addItem("CO1")
            self.copbq15.addItem("CO2")
            self.copbq15.addItem("CO3")
            self.copbq15.addItem("CO4")
            self.copbq15.addItem("CO5")
            self.layout.addWidget(self.copbq15,15,4)
            self.copbq15.currentTextChanged.connect(self.coseeb)

            self.copbq16= QComboBox()
            #self.copbq16.addItem("CO")
            self.copbq16.addItem("CO1")
            self.copbq16.addItem("CO2")
            self.copbq16.addItem("CO3")
            self.copbq16.addItem("CO4")
            self.copbq16.addItem("CO5")
            self.layout.addWidget(self.copbq16,16,4)
            self.copbq16.currentTextChanged.connect(self.coseeb)

            self.copbq17= QComboBox()
            #self.copbq17.addItem("CO")
            self.copbq17.addItem("CO1")
            self.copbq17.addItem("CO2")
            self.copbq17.addItem("CO3")
            self.copbq17.addItem("CO4")
            self.copbq17.addItem("CO5")
            self.layout.addWidget(self.copbq17,17,4)
            self.copbq17.currentTextChanged.connect(self.coseeb)

            
        elif (self.partb==18):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,4)
            self.copbq9.currentTextChanged.connect(self.coseeb)

            self.copbq10= QComboBox()
            #self.copbq10.addItem("CO")
            self.copbq10.addItem("CO1")
            self.copbq10.addItem("CO2")
            self.copbq10.addItem("CO3")
            self.copbq10.addItem("CO4")
            self.copbq10.addItem("CO5")
            self.layout.addWidget(self.copbq10,10,4)
            self.copbq10.currentTextChanged.connect(self.coseeb)

            self.copbq11= QComboBox()
            #self.copbq11.addItem("CO")
            self.copbq11.addItem("CO1")
            self.copbq11.addItem("CO2")
            self.copbq11.addItem("CO3")
            self.copbq11.addItem("CO4")
            self.copbq11.addItem("CO5")
            self.layout.addWidget(self.copbq11,11,4)
            self.copbq11.currentTextChanged.connect(self.coseeb)

            self.copbq12= QComboBox()
            #self.copbq12.addItem("CO")
            self.copbq12.addItem("CO1")
            self.copbq12.addItem("CO2")
            self.copbq12.addItem("CO3")
            self.copbq12.addItem("CO4")
            self.copbq12.addItem("CO5")
            self.layout.addWidget(self.copbq12,12,4)
            self.copbq12.currentTextChanged.connect(self.coseeb)

            self.copbq13= QComboBox()
            #self.copbq13.addItem("CO")
            self.copbq13.addItem("CO1")
            self.copbq13.addItem("CO2")
            self.copbq13.addItem("CO3")
            self.copbq13.addItem("CO4")
            self.copbq13.addItem("CO5")
            self.layout.addWidget(self.copbq13,13,4)
            self.copbq13.currentTextChanged.connect(self.coseeb)

            self.copbq14= QComboBox()
            #self.copbq14.addItem("CO")
            self.copbq14.addItem("CO1")
            self.copbq14.addItem("CO2")
            self.copbq14.addItem("CO3")
            self.copbq14.addItem("CO4")
            self.copbq14.addItem("CO5")
            self.layout.addWidget(self.copbq14,14,4)
            self.copbq14.currentTextChanged.connect(self.coseeb)

            self.copbq15= QComboBox()
            #self.copbq15.addItem("CO")
            self.copbq15.addItem("CO1")
            self.copbq15.addItem("CO2")
            self.copbq15.addItem("CO3")
            self.copbq15.addItem("CO4")
            self.copbq15.addItem("CO5")
            self.layout.addWidget(self.copbq15,15,4)
            self.copbq15.currentTextChanged.connect(self.coseeb)

            self.copbq16= QComboBox()
            #self.copbq16.addItem("CO")
            self.copbq16.addItem("CO1")
            self.copbq16.addItem("CO2")
            self.copbq16.addItem("CO3")
            self.copbq16.addItem("CO4")
            self.copbq16.addItem("CO5")
            self.layout.addWidget(self.copbq16,16,4)
            self.copbq16.currentTextChanged.connect(self.coseeb)

            self.copbq17= QComboBox()
            #self.copbq17.addItem("CO")
            self.copbq17.addItem("CO1")
            self.copbq17.addItem("CO2")
            self.copbq17.addItem("CO3")
            self.copbq17.addItem("CO4")
            self.copbq17.addItem("CO5")
            self.layout.addWidget(self.copbq17,17,4)
            self.copbq17.currentTextChanged.connect(self.coseeb)

            self.copbq18= QComboBox()
            #self.copbq18.addItem("CO")
            self.copbq18.addItem("CO1")
            self.copbq18.addItem("CO2")
            self.copbq18.addItem("CO3")
            self.copbq18.addItem("CO4")
            self.copbq18.addItem("CO5")
            self.layout.addWidget(self.copbq18,18,4)
            self.copbq18.currentTextChanged.connect(self.coseeb)

        elif (self.partb==19):
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,4)
            self.copbq9.currentTextChanged.connect(self.coseeb)

            self.copbq10= QComboBox()
            #self.copbq10.addItem("CO")
            self.copbq10.addItem("CO1")
            self.copbq10.addItem("CO2")
            self.copbq10.addItem("CO3")
            self.copbq10.addItem("CO4")
            self.copbq10.addItem("CO5")
            self.layout.addWidget(self.copbq10,10,4)
            self.copbq10.currentTextChanged.connect(self.coseeb)

            self.copbq11= QComboBox()
            #self.copbq11.addItem("CO")
            self.copbq11.addItem("CO1")
            self.copbq11.addItem("CO2")
            self.copbq11.addItem("CO3")
            self.copbq11.addItem("CO4")
            self.copbq11.addItem("CO5")
            self.layout.addWidget(self.copbq11,11,4)
            self.copbq11.currentTextChanged.connect(self.coseeb)

            self.copbq12= QComboBox()
            #self.copbq12.addItem("CO")
            self.copbq12.addItem("CO1")
            self.copbq12.addItem("CO2")
            self.copbq12.addItem("CO3")
            self.copbq12.addItem("CO4")
            self.copbq12.addItem("CO5")
            self.layout.addWidget(self.copbq12,12,4)
            self.copbq12.currentTextChanged.connect(self.coseeb)

            self.copbq13= QComboBox()
            #self.copbq13.addItem("CO")
            self.copbq13.addItem("CO1")
            self.copbq13.addItem("CO2")
            self.copbq13.addItem("CO3")
            self.copbq13.addItem("CO4")
            self.copbq13.addItem("CO5")
            self.layout.addWidget(self.copbq13,13,4)
            self.copbq13.currentTextChanged.connect(self.coseeb)

            self.copbq14= QComboBox()
            #self.copbq14.addItem("CO")
            self.copbq14.addItem("CO1")
            self.copbq14.addItem("CO2")
            self.copbq14.addItem("CO3")
            self.copbq14.addItem("CO4")
            self.copbq14.addItem("CO5")
            self.layout.addWidget(self.copbq14,14,4)
            self.copbq14.currentTextChanged.connect(self.coseeb)

            self.copbq15= QComboBox()
            #self.copbq15.addItem("CO")
            self.copbq15.addItem("CO1")
            self.copbq15.addItem("CO2")
            self.copbq15.addItem("CO3")
            self.copbq15.addItem("CO4")
            self.copbq15.addItem("CO5")
            self.layout.addWidget(self.copbq15,15,4)
            self.copbq15.currentTextChanged.connect(self.coseeb)

            self.copbq16= QComboBox()
            #self.copbq16.addItem("CO")
            self.copbq16.addItem("CO1")
            self.copbq16.addItem("CO2")
            self.copbq16.addItem("CO3")
            self.copbq16.addItem("CO4")
            self.copbq16.addItem("CO5")
            self.layout.addWidget(self.copbq16,16,4)
            self.copbq16.currentTextChanged.connect(self.coseeb)

            self.copbq17= QComboBox()
            #self.copbq17.addItem("CO")
            self.copbq17.addItem("CO1")
            self.copbq17.addItem("CO2")
            self.copbq17.addItem("CO3")
            self.copbq17.addItem("CO4")
            self.copbq17.addItem("CO5")
            self.layout.addWidget(self.copbq17,17,4)
            self.copbq17.currentTextChanged.connect(self.coseeb)

            self.copbq18= QComboBox()
            #self.copbq18.addItem("CO")
            self.copbq18.addItem("CO1")
            self.copbq18.addItem("CO2")
            self.copbq18.addItem("CO3")
            self.copbq18.addItem("CO4")
            self.copbq18.addItem("CO5")
            self.layout.addWidget(self.copbq18,18,4)
            self.copbq18.currentTextChanged.connect(self.coseeb)

            self.copbq19= QComboBox()
            #self.copbq19.addItem("CO")
            self.copbq19.addItem("CO1")
            self.copbq19.addItem("CO2")
            self.copbq19.addItem("CO3")
            self.copbq19.addItem("CO4")
            self.copbq19.addItem("CO5")
            self.layout.addWidget(self.copbq19,19,4)
            self.copbq19.currentTextChanged.connect(self.coseeb)

            



        elif (self.partb==20):
        
            self.copbq1= QComboBox()
            #self.copbq1.addItem("CO")

            self.copbq1.addItem("CO1")
            self.copbq1.addItem("CO2")
            self.copbq1.addItem("CO3")
            self.copbq1.addItem("CO4")
            self.copbq1.addItem("CO5")
            self.layout.addWidget(self.copbq1,1,4)
            self.copbq1.currentTextChanged.connect(self.coseeb)
            self.copbq2= QComboBox()
            #self.copbq2.addItem("CO")

            self.copbq2.addItem("CO1")
            self.copbq2.addItem("CO2")
            self.copbq2.addItem("CO3")
            self.copbq2.addItem("CO4")
            self.copbq2.addItem("CO5")
            self.layout.addWidget(self.copbq2,2,4)
            self.copbq2.currentTextChanged.connect(self.coseeb)
            
            self.copbq3= QComboBox()
            #self.copbq3.addItem("CO")

            self.copbq3.addItem("CO1")
            self.copbq3.addItem("CO2")
            self.copbq3.addItem("CO3")
            self.copbq3.addItem("CO4")
            self.copbq3.addItem("CO5")
            self.layout.addWidget(self.copbq3,3,4)
            
            self.copbq3.currentTextChanged.connect(self.coseeb)
            
            self.copbq4= QComboBox()
            #self.copbq4.addItem("CO")
            self.copbq4.addItem("CO1")
            self.copbq4.addItem("CO2")
            self.copbq4.addItem("CO3")
            self.copbq4.addItem("CO4")
            self.copbq4.addItem("CO5")
            self.layout.addWidget(self.copbq4,4,4)
            self.copbq4.currentTextChanged.connect(self.coseeb)
            
            self.copbq5= QComboBox()
            #self.copbq5.addItem("CO")
            self.copbq5.addItem("CO1")
            self.copbq5.addItem("CO2")
            self.copbq5.addItem("CO3")
            self.copbq5.addItem("CO4")
            self.copbq5.addItem("CO5")
            self.layout.addWidget(self.copbq5,5,4)
            self.copbq5.currentTextChanged.connect(self.coseeb)
            
            self.copbq6= QComboBox()
            #self.copbq6.addItem("CO")
            self.copbq6.addItem("CO1")
            self.copbq6.addItem("CO2")
            self.copbq6.addItem("CO3")
            self.copbq6.addItem("CO4")
            self.copbq6.addItem("CO5")
            self.layout.addWidget(self.copbq6,6,4)
            self.copbq6.currentTextChanged.connect(self.coseeb)
            
            self.copbq7= QComboBox()
            #self.copbq7.addItem("CO")
            self.copbq7.addItem("CO1")
            self.copbq7.addItem("CO2")
            self.copbq7.addItem("CO3")
            self.copbq7.addItem("CO4")
            self.copbq7.addItem("CO5")
            self.layout.addWidget(self.copbq7,7,4)
            self.copbq7.currentTextChanged.connect(self.coseeb)
            
            self.copbq8= QComboBox()
            #self.copbq8.addItem("CO")
            self.copbq8.addItem("CO1")
            self.copbq8.addItem("CO2")
            self.copbq8.addItem("CO3")
            self.copbq8.addItem("CO4")
            self.copbq8.addItem("CO5")
            self.layout.addWidget(self.copbq8,8,4)
            self.copbq8.currentTextChanged.connect(self.coseeb)

            self.copbq9= QComboBox()
            #self.copbq9.addItem("CO")
            self.copbq9.addItem("CO1")
            self.copbq9.addItem("CO2")
            self.copbq9.addItem("CO3")
            self.copbq9.addItem("CO4")
            self.copbq9.addItem("CO5")
            self.layout.addWidget(self.copbq9,9,4)
            self.copbq9.currentTextChanged.connect(self.coseeb)

            self.copbq10= QComboBox()
            #self.copbq10.addItem("CO")
            self.copbq10.addItem("CO1")
            self.copbq10.addItem("CO2")
            self.copbq10.addItem("CO3")
            self.copbq10.addItem("CO4")
            self.copbq10.addItem("CO5")
            self.layout.addWidget(self.copbq10,10,4)
            self.copbq10.currentTextChanged.connect(self.coseeb)

            self.copbq11= QComboBox()
            #self.copbq11.addItem("CO")
            self.copbq11.addItem("CO1")
            self.copbq11.addItem("CO2")
            self.copbq11.addItem("CO3")
            self.copbq11.addItem("CO4")
            self.copbq11.addItem("CO5")
            self.layout.addWidget(self.copbq11,11,4)
            self.copbq11.currentTextChanged.connect(self.coseeb)

            self.copbq12= QComboBox()
            #self.copbq12.addItem("CO")
            self.copbq12.addItem("CO1")
            self.copbq12.addItem("CO2")
            self.copbq12.addItem("CO3")
            self.copbq12.addItem("CO4")
            self.copbq12.addItem("CO5")
            self.layout.addWidget(self.copbq12,12,4)
            self.copbq12.currentTextChanged.connect(self.coseeb)

            self.copbq13= QComboBox()
            #self.copbq13.addItem("CO")
            self.copbq13.addItem("CO1")
            self.copbq13.addItem("CO2")
            self.copbq13.addItem("CO3")
            self.copbq13.addItem("CO4")
            self.copbq13.addItem("CO5")
            self.layout.addWidget(self.copbq13,13,4)
            self.copbq13.currentTextChanged.connect(self.coseeb)

            self.copbq14= QComboBox()
            #self.copbq14.addItem("CO")
            self.copbq14.addItem("CO1")
            self.copbq14.addItem("CO2")
            self.copbq14.addItem("CO3")
            self.copbq14.addItem("CO4")
            self.copbq14.addItem("CO5")
            self.layout.addWidget(self.copbq14,14,4)
            self.copbq14.currentTextChanged.connect(self.coseeb)

            self.copbq15= QComboBox()
            #self.copbq15.addItem("CO")
            self.copbq15.addItem("CO1")
            self.copbq15.addItem("CO2")
            self.copbq15.addItem("CO3")
            self.copbq15.addItem("CO4")
            self.copbq15.addItem("CO5")
            self.layout.addWidget(self.copbq15,15,4)
            self.copbq15.currentTextChanged.connect(self.coseeb)

            self.copbq16= QComboBox()
            #self.copbq16.addItem("CO")
            self.copbq16.addItem("CO1")
            self.copbq16.addItem("CO2")
            self.copbq16.addItem("CO3")
            self.copbq16.addItem("CO4")
            self.copbq16.addItem("CO5")
            self.layout.addWidget(self.copbq16,16,4)
            self.copbq16.currentTextChanged.connect(self.coseeb)

            self.copbq17= QComboBox()
            #self.copbq17.addItem("CO")
            self.copbq17.addItem("CO1")
            self.copbq17.addItem("CO2")
            self.copbq17.addItem("CO3")
            self.copbq17.addItem("CO4")
            self.copbq17.addItem("CO5")
            self.layout.addWidget(self.copbq17,17,4)
            self.copbq17.currentTextChanged.connect(self.coseeb)
            
            self.copbq18= QComboBox()
            #self.copbq18.addItem("CO")
            self.copbq18.addItem("CO1")
            self.copbq18.addItem("CO2")
            self.copbq18.addItem("CO3")
            self.copbq18.addItem("CO4")
            self.copbq18.addItem("CO5")
            self.layout.addWidget(self.copbq18,18,4)
            self.copbq18.currentTextChanged.connect(self.coseeb)

            self.copbq19= QComboBox()
            #self.copbq19.addItem("CO")
            self.copbq19.addItem("CO1")
            self.copbq19.addItem("CO2")
            self.copbq19.addItem("CO3")
            self.copbq19.addItem("CO4")
            self.copbq19.addItem("CO5")
            self.layout.addWidget(self.copbq19,19,4)
            self.copbq19.currentTextChanged.connect(self.coseeb)

            self.copbq20= QComboBox()
            #self.copbq20.addItem("CO")
            self.copbq20.addItem("CO1")
            self.copbq20.addItem("CO2")
            self.copbq20.addItem("CO3")
            self.copbq20.addItem("CO4")
            self.copbq20.addItem("CO5")
            self.layout.addWidget(self.copbq20,20,4)
            self.copbq20.currentTextChanged.connect(self.coseeb)
            
        
    def cosee(self):

        self.partacos=[]
        if (self.parta==1):
            q1 = self.copaq1.currentText()
            
            
            self.partacos.append(q1)
           
         
            print(self.partacos)
        elif (self.parta==2):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            
         
            print(self.partacos)
        elif (self.parta==3):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
         
            print(self.partacos)
        elif (self.parta==4):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
         
            print(self.partacos)
        elif (self.parta==5):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
         
            print(self.partacos)
        elif (self.parta==6):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
           
            
            print(self.partacos)
        elif (self.parta==7):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            
            print(self.partacos)
        elif (self.parta==8):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            
            
            print(self.partacos)
        elif (self.parta==9):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            
            print(self.partacos)
        elif (self.parta==10):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            
            print(self.partacos)
        elif (self.parta==11):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            
           
        
            
            print(self.partacos)
        elif (self.parta==12):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
                        
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            
           
        
            
            print(self.partacos)
        elif (self.parta==13):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            
           
        
            
            print(self.partacos)
        elif (self.parta==14):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
           
           
        
            
            print(self.partacos)
        
        elif (self.parta==15):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
           
        
            
            print(self.partacos)
        elif (self.parta==16):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
           
        
            
            print(self.partacos)
        elif (self.parta==17):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
          
        
            
            print(self.partacos)
        elif (self.parta==18):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            
            
            print(self.partacos)
        elif (self.parta==19):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
          
        
            
            print(self.partacos)
        elif (self.parta==20):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            q20 = self.copaq20.currentText()
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
            self.partacos.append(q20)
        
            
            print(self.partacos)
            
        elif (self.parta==21):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            q20 = self.copaq20.currentText()
            q21 = self.copaq21.currentText()
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
            self.partacos.append(q20)
            self.partacos.append(q21)
        
            
            print(self.partacos)
            
        elif (self.parta==22):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            q20 = self.copaq20.currentText()
            q21 = self.copaq21.currentText()
            q22= self.copaq22.currentText()
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
            self.partacos.append(q20)
            self.partacos.append(q21)
            self.partacos.append(q22)
            
            
            print(self.partacos)
        elif (self.parta==23):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            q20 = self.copaq20.currentText()
            q21 = self.copaq21.currentText()
            q22= self.copaq22.currentText()
            q23 = self.copaq23.currentText()
            
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
            self.partacos.append(q20)
            self.partacos.append(q21)
            self.partacos.append(q22)
            self.partacos.append(q23)
            
            print(self.partacos)
        elif (self.parta==24):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            q20 = self.copaq20.currentText()
            q21 = self.copaq21.currentText()
            q22= self.copaq22.currentText()
            q23 = self.copaq23.currentText()
            q24 = self.copaq24.currentText()
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
            self.partacos.append(q20)
            self.partacos.append(q21)
            self.partacos.append(q22)
            self.partacos.append(q23)
            self.partacos.append(q24)
            
            
            print(self.partacos)
        elif (self.parta==25):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            q20 = self.copaq20.currentText()
            q21 = self.copaq21.currentText()
            q22= self.copaq22.currentText()
            q23 = self.copaq23.currentText()
            q24 = self.copaq24.currentText()
            q25 = self.copaq25.currentText()
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
            self.partacos.append(q20)
            self.partacos.append(q21)
            self.partacos.append(q22)
            self.partacos.append(q23)
            self.partacos.append(q24)
            self.partacos.append(q25)
            
            print(self.partacos)
        elif (self.parta==26):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            q20 = self.copaq20.currentText()
            q21 = self.copaq21.currentText()
            q22= self.copaq22.currentText()
            q23 = self.copaq23.currentText()
            q24 = self.copaq24.currentText()
            q25 = self.copaq25.currentText()
            q26 = self.copaq26.currentText()
            
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
            self.partacos.append(q20)
            self.partacos.append(q21)
            self.partacos.append(q22)
            self.partacos.append(q23)
            self.partacos.append(q24)
            self.partacos.append(q25)
            self.partacos.append(q26)
            
            print(self.partacos)
        elif (self.parta==27):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            q20 = self.copaq20.currentText()
            q21 = self.copaq21.currentText()
            q22= self.copaq22.currentText()
            q23 = self.copaq23.currentText()
            q24 = self.copaq24.currentText()
            q25 = self.copaq25.currentText()
            q26 = self.copaq26.currentText()
            q27 = self.copaq27.currentText()
          
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
            self.partacos.append(q20)
            self.partacos.append(q21)
            self.partacos.append(q22)
            self.partacos.append(q23)
            self.partacos.append(q24)
            self.partacos.append(q25)
            self.partacos.append(q26)
            self.partacos.append(q27)
            
            print(self.partacos)
        elif (self.parta==28):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            q20 = self.copaq20.currentText()
            q21 = self.copaq21.currentText()
            q22= self.copaq22.currentText()
            q23 = self.copaq23.currentText()
            q24 = self.copaq24.currentText()
            q25 = self.copaq25.currentText()
            q26 = self.copaq26.currentText()
            q27 = self.copaq27.currentText()
            q28 = self.copaq28.currentText()
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
            self.partacos.append(q20)
            self.partacos.append(q21)
            self.partacos.append(q22)
            self.partacos.append(q23)
            self.partacos.append(q24)
            self.partacos.append(q25)
            self.partacos.append(q26)
            self.partacos.append(q27)
            self.partacos.append(q28)
            print(self.partacos)
        elif (self.parta==29):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            q20 = self.copaq20.currentText()
            q21 = self.copaq21.currentText()
            q22= self.copaq22.currentText()
            q23 = self.copaq23.currentText()
            q24 = self.copaq24.currentText()
            q25 = self.copaq25.currentText()
            q26 = self.copaq26.currentText()
            q27 = self.copaq27.currentText()
            q28 = self.copaq28.currentText()
            q29 = self.copaq29.currentText()
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
            self.partacos.append(q20)
            self.partacos.append(q21)
            self.partacos.append(q22)
            self.partacos.append(q23)
            self.partacos.append(q24)
            self.partacos.append(q25)
            self.partacos.append(q26)
            self.partacos.append(q27)
            self.partacos.append(q28)
            self.partacos.append(q29)
            print(self.partacos)
        elif (self.parta==30):
            q1 = self.copaq1.currentText()
            q2 = self.copaq2.currentText()
            q3 = self.copaq3.currentText()
            q4 = self.copaq4.currentText()
            q5 = self.copaq5.currentText()
            q6 = self.copaq6.currentText()
            q7 = self.copaq7.currentText()
            q8 = self.copaq8.currentText()
            q9 = self.copaq9.currentText()
            q10 = self.copaq10.currentText()
            q11= self.copaq11.currentText()
            q12 = self.copaq12.currentText()
            q13 = self.copaq13.currentText()
            q14 = self.copaq14.currentText()
            q15 = self.copaq15.currentText()
            q16 = self.copaq16.currentText()
            q17 = self.copaq17.currentText()
            q18 = self.copaq18.currentText()
            q19= self.copaq19.currentText()
            q20 = self.copaq20.currentText()
            q21 = self.copaq21.currentText()
            q22= self.copaq22.currentText()
            q23 = self.copaq23.currentText()
            q24 = self.copaq24.currentText()
            q25 = self.copaq25.currentText()
            q26 = self.copaq26.currentText()
            q27 = self.copaq27.currentText()
            q28 = self.copaq28.currentText()
            q29 = self.copaq29.currentText()
            q30 = self.copaq30.currentText()
            self.partacos.append(q1)
            self.partacos.append(q2)
            self.partacos.append(q3)
            self.partacos.append(q4)
            self.partacos.append(q5)
            self.partacos.append(q6)
            self.partacos.append(q7)
            self.partacos.append(q8)
            self.partacos.append(q9)
            self.partacos.append(q10)
            self.partacos.append(q11)
            self.partacos.append(q12)
            self.partacos.append(q13)
            self.partacos.append(q14)
            self.partacos.append(q15)
            self.partacos.append(q16)
            self.partacos.append(q17)
            self.partacos.append(q18)
            self.partacos.append(q19)
            self.partacos.append(q20)
            self.partacos.append(q21)
            self.partacos.append(q22)
            self.partacos.append(q23)
            self.partacos.append(q24)
            self.partacos.append(q25)
            self.partacos.append(q26)
            self.partacos.append(q27)
            self.partacos.append(q28)
            self.partacos.append(q29)
            self.partacos.append(q30)
            print(self.partacos)
        global partacos
        partacos=self.partacos
    def coseeb(self):
        self.partbcos=[]
        if (self.partb==1):
            q1 = self.copbq1.currentText()
            
            self.partbcos.append(q1)
             
            print(self.partbcos)
        elif (self.partb==2):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
             
            print(self.partbcos)
        elif (self.partb==3):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
             
            print(self.partbcos)
        elif (self.partb==4):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            
            print(self.partbcos)
        elif (self.partb==5):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()

            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
             
            print(self.partbcos)
        elif (self.partb==6):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
             
            print(self.partbcos)
        elif (self.partb==7):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
             
            print(self.partbcos)
        elif (self.partb==8):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            
            print(self.partbcos)
        elif (self.partb==9):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
             
            print(self.partbcos)
        elif (self.partb==10):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            q10 = self.copbq10.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
            self.partbcos.append(q10)
             
            print(self.partbcos)
        elif (self.partb==11):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            q10 = self.copbq10.currentText()
            q11= self.copbq11.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
            self.partbcos.append(q10)
            self.partbcos.append(q11)
             
            print(self.partbcos)
        elif (self.partb==12):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            q10 = self.copbq10.currentText()
            q11= self.copbq11.currentText()
            q12 = self.copbq12.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
            self.partbcos.append(q10)
            self.partbcos.append(q11)
            self.partbcos.append(q12)
            
            print(self.partbcos)
        elif (self.partb==13):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            q10 = self.copbq10.currentText()
            q11= self.copbq11.currentText()
            q12 = self.copbq12.currentText()
            q13 = self.copbq13.currentText()
            
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
            self.partbcos.append(q10)
            self.partbcos.append(q11)
            self.partbcos.append(q12)
            self.partbcos.append(q13)
            
            
            print(self.partbcos)
        elif (self.partb==14):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            q10 = self.copbq10.currentText()
            q11= self.copbq11.currentText()
            q12 = self.copbq12.currentText()
            q13 = self.copbq13.currentText()
            q14 = self.copbq14.currentText()
            
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
            self.partbcos.append(q10)
            self.partbcos.append(q11)
            self.partbcos.append(q12)
            self.partbcos.append(q13)
            self.partbcos.append(q14)
          
            
            print(self.partbcos)
        elif (self.partb==15):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            q10 = self.copbq10.currentText()
            q11= self.copbq11.currentText()
            q12 = self.copbq12.currentText()
            q13 = self.copbq13.currentText()
            q14 = self.copbq14.currentText()
            q15 = self.copbq15.currentText()
            
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
            self.partbcos.append(q10)
            self.partbcos.append(q11)
            self.partbcos.append(q12)
            self.partbcos.append(q13)
            self.partbcos.append(q14)
            self.partbcos.append(q15)
            
            
            print(self.partbcos)
        elif (self.partb==16):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            q10 = self.copbq10.currentText()
            q11= self.copbq11.currentText()
            q12 = self.copbq12.currentText()
            q13 = self.copbq13.currentText()
            q14 = self.copbq14.currentText()
            q15 = self.copbq15.currentText()
            q16 = self.copbq16.currentText()
            
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
            self.partbcos.append(q10)
            self.partbcos.append(q11)
            self.partbcos.append(q12)
            self.partbcos.append(q13)
            self.partbcos.append(q14)
            self.partbcos.append(q15)
            self.partbcos.append(q16)
            
            
            print(self.partbcos)
        elif (self.partb==17):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            q10 = self.copbq10.currentText()
            q11= self.copbq11.currentText()
            q12 = self.copbq12.currentText()
            q13 = self.copbq13.currentText()
            q14 = self.copbq14.currentText()
            q15 = self.copbq15.currentText()
            q16 = self.copbq16.currentText()
            q17 = self.copbq17.currentText()
            
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
            self.partbcos.append(q10)
            self.partbcos.append(q11)
            self.partbcos.append(q12)
            self.partbcos.append(q13)
            self.partbcos.append(q14)
            self.partbcos.append(q15)
            self.partbcos.append(q16)
            self.partbcos.append(q17)
            
            
            print(self.partbcos)
        elif (self.partb==18):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            q10 = self.copbq10.currentText()
            q11= self.copbq11.currentText()
            q12 = self.copbq12.currentText()
            q13 = self.copbq13.currentText()
            q14 = self.copbq14.currentText()
            q15 = self.copbq15.currentText()
            q16 = self.copbq16.currentText()
            q17 = self.copbq17.currentText()
            q18 = self.copbq18.currentText()
            
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
            self.partbcos.append(q10)
            self.partbcos.append(q11)
            self.partbcos.append(q12)
            self.partbcos.append(q13)
            self.partbcos.append(q14)
            self.partbcos.append(q15)
            self.partbcos.append(q16)
            self.partbcos.append(q17)
            self.partbcos.append(q18)
            
            
            print(self.partbcos)
        elif (self.partb==19):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            q10 = self.copbq10.currentText()
            q11= self.copbq11.currentText()
            q12 = self.copbq12.currentText()
            q13 = self.copbq13.currentText()
            q14 = self.copbq14.currentText()
            q15 = self.copbq15.currentText()
            q16 = self.copbq16.currentText()
            q17 = self.copbq17.currentText()
            q18 = self.copbq18.currentText()
            q19= self.copbq19.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
            self.partbcos.append(q10)
            self.partbcos.append(q11)
            self.partbcos.append(q12)
            self.partbcos.append(q13)
            self.partbcos.append(q14)
            self.partbcos.append(q15)
            self.partbcos.append(q16)
            self.partbcos.append(q17)
            self.partbcos.append(q18)
            self.partbcos.append(q19)
            
            print(self.partbcos)
        
        elif (self.partb==20):
            q1 = self.copbq1.currentText()
            q2 = self.copbq2.currentText()
            q3 = self.copbq3.currentText()
            q4 = self.copbq4.currentText()
            q5 = self.copbq5.currentText()
            q6 = self.copbq6.currentText()
            q7 = self.copbq7.currentText()
            q8 = self.copbq8.currentText()
            q9 = self.copbq9.currentText()
            q10 = self.copbq10.currentText()
            q11= self.copbq11.currentText()
            q12 = self.copbq12.currentText()
            q13 = self.copbq13.currentText()
            q14 = self.copbq14.currentText()
            q15 = self.copbq15.currentText()
            q16 = self.copbq16.currentText()
            q17 = self.copbq17.currentText()
            q18 = self.copbq18.currentText()
            q19= self.copbq19.currentText()
            q20 = self.copbq20.currentText()
            
            self.partbcos.append(q1)
            self.partbcos.append(q2)
            self.partbcos.append(q3)
            self.partbcos.append(q4)
            self.partbcos.append(q5)
            self.partbcos.append(q6)
            self.partbcos.append(q7)
            self.partbcos.append(q8)
            self.partbcos.append(q9)
            self.partbcos.append(q10)
            self.partbcos.append(q11)
            self.partbcos.append(q12)
            self.partbcos.append(q13)
            self.partbcos.append(q14)
            self.partbcos.append(q15)
            self.partbcos.append(q16)
            self.partbcos.append(q17)
            self.partbcos.append(q18)
            self.partbcos.append(q19)
            self.partbcos.append(q20)
            
            print(self.partbcos)
        global partbcos
        partbcos=self.partbcos    
    def coursea(self):
        if (self.parta==1):
            
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            
            

 
            
        elif (self.parta==2):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            
            
        elif (self.parta==3):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            
            
        elif (self.parta==4):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
                        
        elif (self.parta==5):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            


            
        elif (self.parta==6):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            

            
            
        elif (self.parta==7):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            
        elif (self.parta==8):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)




        
        elif (self.parta==9):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)




        elif (self.parta==10):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)




        elif (self.parta==11):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

        elif (self.parta==12):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)


                    
 

        elif (self.parta==13):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)


                    

           

            
        elif (self.parta==14):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            

                    

           
        elif (self.parta==15):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)


                    

            
        elif (self.parta==16):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)


                    

            
        elif (self.parta==17):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

                    

            
        elif (self.parta==18):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copq18.currentTextChanged.connect(self.cosee)


                    

        elif (self.parta==19):
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            

            

            



        elif (self.parta==20):
        
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copaq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            self.copaq20= QComboBox()
            #self.copaq20.addItem("CO")
            self.copaq20.addItem("CO1")
            self.copaq20.addItem("CO2")
            self.copaq20.addItem("CO3")
            self.copaq20.addItem("CO4")
            self.copaq20.addItem("CO5")
            self.layout.addWidget(self.copaq20,20,1)
            self.copaq20.currentTextChanged.connect(self.cosee)
        
        elif (self.parta==21):
        
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copaq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            self.copaq20= QComboBox()
            #self.copaq20.addItem("CO")
            self.copaq20.addItem("CO1")
            self.copaq20.addItem("CO2")
            self.copaq20.addItem("CO3")
            self.copaq20.addItem("CO4")
            self.copaq20.addItem("CO5")
            self.layout.addWidget(self.copaq20,20,1)
            self.copaq20.currentTextChanged.connect(self.cosee)

            self.copaq21= QComboBox()
            #self.copaq21.addItem("CO")
            self.copaq21.addItem("CO1")
            self.copaq21.addItem("CO2")
            self.copaq21.addItem("CO3")
            self.copaq21.addItem("CO4")
            self.copaq21.addItem("CO5")
            self.layout.addWidget(self.copaq21,21,1)
            self.copaq21.currentTextChanged.connect(self.cosee)


        elif (self.parta==22):
        
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copaq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            self.copaq20= QComboBox()
            #self.copaq20.addItem("CO")
            self.copaq20.addItem("CO1")
            self.copaq20.addItem("CO2")
            self.copaq20.addItem("CO3")
            self.copaq20.addItem("CO4")
            self.copaq20.addItem("CO5")
            self.layout.addWidget(self.copaq20,20,1)
            self.copaq20.currentTextChanged.connect(self.cosee)

            self.copaq21= QComboBox()
            #self.copaq21.addItem("CO")
            self.copaq21.addItem("CO1")
            self.copaq21.addItem("CO2")
            self.copaq21.addItem("CO3")
            self.copaq21.addItem("CO4")
            self.copaq21.addItem("CO5")
            self.layout.addWidget(self.copaq21,21,1)
            self.copaq21.currentTextChanged.connect(self.cosee)

            self.copaq22= QComboBox()
            #self.copaq22.addItem("CO")
            self.copaq22.addItem("CO1")
            self.copaq22.addItem("CO2")
            self.copaq22.addItem("CO3")
            self.copaq22.addItem("CO4")
            self.copaq22.addItem("CO5")
            self.layout.addWidget(self.copaq22,22,1)
            self.copaq22.currentTextChanged.connect(self.cosee)


        elif (self.parta==23):
        
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copaq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            self.copaq20= QComboBox()
            #self.copaq20.addItem("CO")
            self.copaq20.addItem("CO1")
            self.copaq20.addItem("CO2")
            self.copaq20.addItem("CO3")
            self.copaq20.addItem("CO4")
            self.copaq20.addItem("CO5")
            self.layout.addWidget(self.copaq20,20,1)
            self.copaq20.currentTextChanged.connect(self.cosee)

            self.copaq21= QComboBox()
            #self.copaq21.addItem("CO")
            self.copaq21.addItem("CO1")
            self.copaq21.addItem("CO2")
            self.copaq21.addItem("CO3")
            self.copaq21.addItem("CO4")
            self.copaq21.addItem("CO5")
            self.layout.addWidget(self.copaq21,21,1)
            self.copaq21.currentTextChanged.connect(self.cosee)

            self.copaq22= QComboBox()
            #self.copaq22.addItem("CO")
            self.copaq22.addItem("CO1")
            self.copaq22.addItem("CO2")
            self.copaq22.addItem("CO3")
            self.copaq22.addItem("CO4")
            self.copaq22.addItem("CO5")
            self.layout.addWidget(self.copaq22,22,1)
            self.copaq22.currentTextChanged.connect(self.cosee)

            self.copaq23= QComboBox()
            #self.copaq23.addItem("CO")
            self.copaq23.addItem("CO1")
            self.copaq23.addItem("CO2")
            self.copaq23.addItem("CO3")
            self.copaq23.addItem("CO4")
            self.copaq23.addItem("CO5")
            self.layout.addWidget(self.copaq23,23,1)
            self.copaq23.currentTextChanged.connect(self.cosee)

        elif (self.parta==24):
        
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copaq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            self.copaq20= QComboBox()
            #self.copaq20.addItem("CO")
            self.copaq20.addItem("CO1")
            self.copaq20.addItem("CO2")
            self.copaq20.addItem("CO3")
            self.copaq20.addItem("CO4")
            self.copaq20.addItem("CO5")
            self.layout.addWidget(self.copaq20,20,1)
            self.copaq20.currentTextChanged.connect(self.cosee)

            self.copaq21= QComboBox()
            #self.copaq21.addItem("CO")
            self.copaq21.addItem("CO1")
            self.copaq21.addItem("CO2")
            self.copaq21.addItem("CO3")
            self.copaq21.addItem("CO4")
            self.copaq21.addItem("CO5")
            self.layout.addWidget(self.copaq21,21,1)
            self.copaq21.currentTextChanged.connect(self.cosee)

            self.copaq22= QComboBox()
            #self.copaq22.addItem("CO")
            self.copaq22.addItem("CO1")
            self.copaq22.addItem("CO2")
            self.copaq22.addItem("CO3")
            self.copaq22.addItem("CO4")
            self.copaq22.addItem("CO5")
            self.layout.addWidget(self.copaq22,22,1)
            self.copaq22.currentTextChanged.connect(self.cosee)

            self.copaq23= QComboBox()
            #self.copaq23.addItem("CO")
            self.copaq23.addItem("CO1")
            self.copaq23.addItem("CO2")
            self.copaq23.addItem("CO3")
            self.copaq23.addItem("CO4")
            self.copaq23.addItem("CO5")
            self.layout.addWidget(self.copaq23,23,1)
            self.copaq23.currentTextChanged.connect(self.cosee)

            self.copaq24= QComboBox()
            #self.copaq24.addItem("CO")
            self.copaq24.addItem("CO1")
            self.copaq24.addItem("CO2")
            self.copaq24.addItem("CO3")
            self.copaq24.addItem("CO4")
            self.copaq24.addItem("CO5")
            self.layout.addWidget(self.copaq24,24,1)
            self.copaq24.currentTextChanged.connect(self.cosee)


        elif (self.parta==25):
        
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copaq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            self.copaq20= QComboBox()
            #self.copaq20.addItem("CO")
            self.copaq20.addItem("CO1")
            self.copaq20.addItem("CO2")
            self.copaq20.addItem("CO3")
            self.copaq20.addItem("CO4")
            self.copaq20.addItem("CO5")
            self.layout.addWidget(self.copaq20,20,1)
            self.copaq20.currentTextChanged.connect(self.cosee)

            self.copaq21= QComboBox()
            #self.copaq21.addItem("CO")
            self.copaq21.addItem("CO1")
            self.copaq21.addItem("CO2")
            self.copaq21.addItem("CO3")
            self.copaq21.addItem("CO4")
            self.copaq21.addItem("CO5")
            self.layout.addWidget(self.copaq21,21,1)
            self.copaq21.currentTextChanged.connect(self.cosee)

            self.copaq22= QComboBox()
            #self.copaq22.addItem("CO")
            self.copaq22.addItem("CO1")
            self.copaq22.addItem("CO2")
            self.copaq22.addItem("CO3")
            self.copaq22.addItem("CO4")
            self.copaq22.addItem("CO5")
            self.layout.addWidget(self.copaq22,22,1)
            self.copaq22.currentTextChanged.connect(self.cosee)

            self.copaq23= QComboBox()
            #self.copaq23.addItem("CO")
            self.copaq23.addItem("CO1")
            self.copaq23.addItem("CO2")
            self.copaq23.addItem("CO3")
            self.copaq23.addItem("CO4")
            self.copaq23.addItem("CO5")
            self.layout.addWidget(self.copaq23,23,1)
            self.copaq23.currentTextChanged.connect(self.cosee)

            self.copaq24= QComboBox()
            #self.copaq24.addItem("CO")
            self.copaq24.addItem("CO1")
            self.copaq24.addItem("CO2")
            self.copaq24.addItem("CO3")
            self.copaq24.addItem("CO4")
            self.copaq24.addItem("CO5")
            self.layout.addWidget(self.copaq24,24,1)
            self.copaq24.currentTextChanged.connect(self.cosee)

            self.copaq25= QComboBox()
            #self.copaq25.addItem("CO")
            self.copaq25.addItem("CO1")
            self.copaq25.addItem("CO2")
            self.copaq25.addItem("CO3")
            self.copaq25.addItem("CO4")
            self.copaq25.addItem("CO5")
            self.layout.addWidget(self.copaq25,25,1)
            self.copaq25.currentTextChanged.connect(self.cosee)


        elif (self.parta==26):
        
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copaq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            self.copaq20= QComboBox()
            #self.copaq20.addItem("CO")
            self.copaq20.addItem("CO1")
            self.copaq20.addItem("CO2")
            self.copaq20.addItem("CO3")
            self.copaq20.addItem("CO4")
            self.copaq20.addItem("CO5")
            self.layout.addWidget(self.copaq20,20,1)
            self.copaq20.currentTextChanged.connect(self.cosee)

            self.copaq21= QComboBox()
            #self.copaq21.addItem("CO")
            self.copaq21.addItem("CO1")
            self.copaq21.addItem("CO2")
            self.copaq21.addItem("CO3")
            self.copaq21.addItem("CO4")
            self.copaq21.addItem("CO5")
            self.layout.addWidget(self.copaq21,21,1)
            self.copaq21.currentTextChanged.connect(self.cosee)

            self.copaq22= QComboBox()
            #self.copaq22.addItem("CO")
            self.copaq22.addItem("CO1")
            self.copaq22.addItem("CO2")
            self.copaq22.addItem("CO3")
            self.copaq22.addItem("CO4")
            self.copaq22.addItem("CO5")
            self.layout.addWidget(self.copaq22,22,1)
            self.copaq22.currentTextChanged.connect(self.cosee)

            self.copaq23= QComboBox()
            #self.copaq23.addItem("CO")
            self.copaq23.addItem("CO1")
            self.copaq23.addItem("CO2")
            self.copaq23.addItem("CO3")
            self.copaq23.addItem("CO4")
            self.copaq23.addItem("CO5")
            self.layout.addWidget(self.copaq23,23,1)
            self.copaq23.currentTextChanged.connect(self.cosee)

            self.copaq24= QComboBox()
            #self.copaq24.addItem("CO")
            self.copaq24.addItem("CO1")
            self.copaq24.addItem("CO2")
            self.copaq24.addItem("CO3")
            self.copaq24.addItem("CO4")
            self.copaq24.addItem("CO5")
            self.layout.addWidget(self.copaq24,24,1)
            self.copaq24.currentTextChanged.connect(self.cosee)

            self.copaq25= QComboBox()
            #self.copaq25.addItem("CO")
            self.copaq25.addItem("CO1")
            self.copaq25.addItem("CO2")
            self.copaq25.addItem("CO3")
            self.copaq25.addItem("CO4")
            self.copaq25.addItem("CO5")
            self.layout.addWidget(self.copaq25,25,1)
            self.copaq25.currentTextChanged.connect(self.cosee)

            self.copaq26= QComboBox()
            #self.copaq26.addItem("CO")
            self.copaq26.addItem("CO1")
            self.copaq26.addItem("CO2")
            self.copaq26.addItem("CO3")
            self.copaq26.addItem("CO4")
            self.copaq26.addItem("CO5")
            self.layout.addWidget(self.copaq26,26,1)
            self.copaq26.currentTextChanged.connect(self.cosee)

        elif (self.parta==27):
        
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copaq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            self.copaq20= QComboBox()
            #self.copaq20.addItem("CO")
            self.copaq20.addItem("CO1")
            self.copaq20.addItem("CO2")
            self.copaq20.addItem("CO3")
            self.copaq20.addItem("CO4")
            self.copaq20.addItem("CO5")
            self.layout.addWidget(self.copaq20,20,1)
            self.copaq20.currentTextChanged.connect(self.cosee)

            self.copaq21= QComboBox()
            #self.copaq21.addItem("CO")
            self.copaq21.addItem("CO1")
            self.copaq21.addItem("CO2")
            self.copaq21.addItem("CO3")
            self.copaq21.addItem("CO4")
            self.copaq21.addItem("CO5")
            self.layout.addWidget(self.copaq21,21,1)
            self.copaq21.currentTextChanged.connect(self.cosee)

            self.copaq22= QComboBox()
            #self.copaq22.addItem("CO")
            self.copaq22.addItem("CO1")
            self.copaq22.addItem("CO2")
            self.copaq22.addItem("CO3")
            self.copaq22.addItem("CO4")
            self.copaq22.addItem("CO5")
            self.layout.addWidget(self.copaq22,22,1)
            self.copaq22.currentTextChanged.connect(self.cosee)

            self.copaq23= QComboBox()
            #self.copaq23.addItem("CO")
            self.copaq23.addItem("CO1")
            self.copaq23.addItem("CO2")
            self.copaq23.addItem("CO3")
            self.copaq23.addItem("CO4")
            self.copaq23.addItem("CO5")
            self.layout.addWidget(self.copaq23,23,1)
            self.copaq23.currentTextChanged.connect(self.cosee)

            self.copaq24= QComboBox()
            #self.copaq24.addItem("CO")
            self.copaq24.addItem("CO1")
            self.copaq24.addItem("CO2")
            self.copaq24.addItem("CO3")
            self.copaq24.addItem("CO4")
            self.copaq24.addItem("CO5")
            self.layout.addWidget(self.copaq24,24,1)
            self.copaq24.currentTextChanged.connect(self.cosee)

            self.copaq25= QComboBox()
            #self.copaq25.addItem("CO")
            self.copaq25.addItem("CO1")
            self.copaq25.addItem("CO2")
            self.copaq25.addItem("CO3")
            self.copaq25.addItem("CO4")
            self.copaq25.addItem("CO5")
            self.layout.addWidget(self.copaq25,25,1)
            self.copaq25.currentTextChanged.connect(self.cosee)

            self.copaq26= QComboBox()
            #self.copaq26.addItem("CO")
            self.copaq26.addItem("CO1")
            self.copaq26.addItem("CO2")
            self.copaq26.addItem("CO3")
            self.copaq26.addItem("CO4")
            self.copaq26.addItem("CO5")
            self.layout.addWidget(self.copaq26,26,1)
            self.copaq26.currentTextChanged.connect(self.cosee)

            self.copaq27= QComboBox()
            #self.copaq27.addItem("CO")
            self.copaq27.addItem("CO1")
            self.copaq27.addItem("CO2")
            self.copaq27.addItem("CO3")
            self.copaq27.addItem("CO4")
            self.copaq27.addItem("CO5")
            self.layout.addWidget(self.copaq27,27,1)
            self.copaq27.currentTextChanged.connect(self.cosee)

           
        elif (self.parta==28):
        
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copaq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            self.copaq20= QComboBox()
            #self.copaq20.addItem("CO")
            self.copaq20.addItem("CO1")
            self.copaq20.addItem("CO2")
            self.copaq20.addItem("CO3")
            self.copaq20.addItem("CO4")
            self.copaq20.addItem("CO5")
            self.layout.addWidget(self.copaq20,20,1)
            self.copaq20.currentTextChanged.connect(self.cosee)

            self.copaq21= QComboBox()
            #self.copaq21.addItem("CO")
            self.copaq21.addItem("CO1")
            self.copaq21.addItem("CO2")
            self.copaq21.addItem("CO3")
            self.copaq21.addItem("CO4")
            self.copaq21.addItem("CO5")
            self.layout.addWidget(self.copaq21,21,1)
            self.copaq21.currentTextChanged.connect(self.cosee)

            self.copaq22= QComboBox()
            #self.copaq22.addItem("CO")
            self.copaq22.addItem("CO1")
            self.copaq22.addItem("CO2")
            self.copaq22.addItem("CO3")
            self.copaq22.addItem("CO4")
            self.copaq22.addItem("CO5")
            self.layout.addWidget(self.copaq22,22,1)
            self.copaq22.currentTextChanged.connect(self.cosee)

            self.copaq23= QComboBox()
            #self.copaq23.addItem("CO")
            self.copaq23.addItem("CO1")
            self.copaq23.addItem("CO2")
            self.copaq23.addItem("CO3")
            self.copaq23.addItem("CO4")
            self.copaq23.addItem("CO5")
            self.layout.addWidget(self.copaq23,23,1)
            self.copaq23.currentTextChanged.connect(self.cosee)

            self.copaq24= QComboBox()
            #self.copaq24.addItem("CO")
            self.copaq24.addItem("CO1")
            self.copaq24.addItem("CO2")
            self.copaq24.addItem("CO3")
            self.copaq24.addItem("CO4")
            self.copaq24.addItem("CO5")
            self.layout.addWidget(self.copaq24,24,1)
            self.copaq24.currentTextChanged.connect(self.cosee)

            self.copaq25= QComboBox()
            #self.copaq25.addItem("CO")
            self.copaq25.addItem("CO1")
            self.copaq25.addItem("CO2")
            self.copaq25.addItem("CO3")
            self.copaq25.addItem("CO4")
            self.copaq25.addItem("CO5")
            self.layout.addWidget(self.copaq25,25,1)
            self.copaq25.currentTextChanged.connect(self.cosee)

            self.copaq26= QComboBox()
            #self.copaq26.addItem("CO")
            self.copaq26.addItem("CO1")
            self.copaq26.addItem("CO2")
            self.copaq26.addItem("CO3")
            self.copaq26.addItem("CO4")
            self.copaq26.addItem("CO5")
            self.layout.addWidget(self.copaq26,26,1)
            self.copaq26.currentTextChanged.connect(self.cosee)

            self.copaq27= QComboBox()
            #self.copaq27.addItem("CO")
            self.copaq27.addItem("CO1")
            self.copaq27.addItem("CO2")
            self.copaq27.addItem("CO3")
            self.copaq27.addItem("CO4")
            self.copaq27.addItem("CO5")
            self.layout.addWidget(self.copaq27,27,1)
            self.copaq27.currentTextChanged.connect(self.cosee)

            self.copaq28= QComboBox()
            #self.copaq28.addItem("CO")
            self.copaq28.addItem("CO1")
            self.copaq28.addItem("CO2")
            self.copaq28.addItem("CO3")
            self.copaq28.addItem("CO4")
            self.copaq28.addItem("CO5")
            self.layout.addWidget(self.copaq28,28,1)
            self.copaq28.currentTextChanged.connect(self.cosee)
        elif (self.parta==29):
        
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copaq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            self.copaq20= QComboBox()
            #self.copaq20.addItem("CO")
            self.copaq20.addItem("CO1")
            self.copaq20.addItem("CO2")
            self.copaq20.addItem("CO3")
            self.copaq20.addItem("CO4")
            self.copaq20.addItem("CO5")
            self.layout.addWidget(self.copaq20,20,1)
            self.copaq20.currentTextChanged.connect(self.cosee)

            self.copaq21= QComboBox()
            #self.copaq21.addItem("CO")
            self.copaq21.addItem("CO1")
            self.copaq21.addItem("CO2")
            self.copaq21.addItem("CO3")
            self.copaq21.addItem("CO4")
            self.copaq21.addItem("CO5")
            self.layout.addWidget(self.copaq21,21,1)
            self.copaq21.currentTextChanged.connect(self.cosee)

            self.copaq22= QComboBox()
            #self.copaq22.addItem("CO")
            self.copaq22.addItem("CO1")
            self.copaq22.addItem("CO2")
            self.copaq22.addItem("CO3")
            self.copaq22.addItem("CO4")
            self.copaq22.addItem("CO5")
            self.layout.addWidget(self.copaq22,22,1)
            self.copaq22.currentTextChanged.connect(self.cosee)

            self.copaq23= QComboBox()
            #self.copaq23.addItem("CO")
            self.copaq23.addItem("CO1")
            self.copaq23.addItem("CO2")
            self.copaq23.addItem("CO3")
            self.copaq23.addItem("CO4")
            self.copaq23.addItem("CO5")
            self.layout.addWidget(self.copaq23,23,1)
            self.copaq23.currentTextChanged.connect(self.cosee)

            self.copaq24= QComboBox()
            #self.copaq24.addItem("CO")
            self.copaq24.addItem("CO1")
            self.copaq24.addItem("CO2")
            self.copaq24.addItem("CO3")
            self.copaq24.addItem("CO4")
            self.copaq24.addItem("CO5")
            self.layout.addWidget(self.copaq24,24,1)
            self.copaq24.currentTextChanged.connect(self.cosee)

            self.copaq25= QComboBox()
            #self.copaq25.addItem("CO")
            self.copaq25.addItem("CO1")
            self.copaq25.addItem("CO2")
            self.copaq25.addItem("CO3")
            self.copaq25.addItem("CO4")
            self.copaq25.addItem("CO5")
            self.layout.addWidget(self.copaq25,25,1)
            self.copaq25.currentTextChanged.connect(self.cosee)

            self.copaq26= QComboBox()
            #self.copaq26.addItem("CO")
            self.copaq26.addItem("CO1")
            self.copaq26.addItem("CO2")
            self.copaq26.addItem("CO3")
            self.copaq26.addItem("CO4")
            self.copaq26.addItem("CO5")
            self.layout.addWidget(self.copaq26,26,1)
            self.copaq26.currentTextChanged.connect(self.cosee)

            self.copaq27= QComboBox()
            #self.copaq27.addItem("CO")
            self.copaq27.addItem("CO1")
            self.copaq27.addItem("CO2")
            self.copaq27.addItem("CO3")
            self.copaq27.addItem("CO4")
            self.copaq27.addItem("CO5")
            self.layout.addWidget(self.copaq27,27,1)
            self.copaq27.currentTextChanged.connect(self.cosee)

            self.copaq28= QComboBox()
            #self.copaq28.addItem("CO")
            self.copaq28.addItem("CO1")
            self.copaq28.addItem("CO2")
            self.copaq28.addItem("CO3")
            self.copaq28.addItem("CO4")
            self.copaq28.addItem("CO5")
            self.layout.addWidget(self.copaq28,28,1)
            self.copaq28.currentTextChanged.connect(self.cosee)

            self.copaq29= QComboBox()
            #self.copaq29.addItem("CO")
            self.copaq29.addItem("CO1")
            self.copaq29.addItem("CO2")
            self.copaq29.addItem("CO3")
            self.copaq29.addItem("CO4")
            self.copaq29.addItem("CO5")
            self.layout.addWidget(self.copaq29,29,1)
            self.copaq29.currentTextChanged.connect(self.cosee)

        elif (self.parta==30):
        
            self.copaq1= QComboBox()
            #self.copaq1.addItem("CO")

            self.copaq1.addItem("CO1")
            self.copaq1.addItem("CO2")
            self.copaq1.addItem("CO3")
            self.copaq1.addItem("CO4")
            self.copaq1.addItem("CO5")
            self.layout.addWidget(self.copaq1,1,1)
            self.copaq1.currentTextChanged.connect(self.cosee)
            self.copaq2= QComboBox()
            #self.copaq2.addItem("CO")

            self.copaq2.addItem("CO1")
            self.copaq2.addItem("CO2")
            self.copaq2.addItem("CO3")
            self.copaq2.addItem("CO4")
            self.copaq2.addItem("CO5")
            self.layout.addWidget(self.copaq2,2,1)
            self.copaq2.currentTextChanged.connect(self.cosee)
            
            self.copaq3= QComboBox()
            #self.copaq3.addItem("CO")

            self.copaq3.addItem("CO1")
            self.copaq3.addItem("CO2")
            self.copaq3.addItem("CO3")
            self.copaq3.addItem("CO4")
            self.copaq3.addItem("CO5")
            self.layout.addWidget(self.copaq3,3,1)
            
            self.copaq3.currentTextChanged.connect(self.cosee)
            
            self.copaq4= QComboBox()
            #self.copaq4.addItem("CO")
            self.copaq4.addItem("CO1")
            self.copaq4.addItem("CO2")
            self.copaq4.addItem("CO3")
            self.copaq4.addItem("CO4")
            self.copaq4.addItem("CO5")
            self.layout.addWidget(self.copaq4,4,1)
            self.copaq4.currentTextChanged.connect(self.cosee)
            
            self.copaq5= QComboBox()
            #self.copaq5.addItem("CO")
            self.copaq5.addItem("CO1")
            self.copaq5.addItem("CO2")
            self.copaq5.addItem("CO3")
            self.copaq5.addItem("CO4")
            self.copaq5.addItem("CO5")
            self.layout.addWidget(self.copaq5,5,1)
            self.copaq5.currentTextChanged.connect(self.cosee)
            
            self.copaq6= QComboBox()
            #self.copaq6.addItem("CO")
            self.copaq6.addItem("CO1")
            self.copaq6.addItem("CO2")
            self.copaq6.addItem("CO3")
            self.copaq6.addItem("CO4")
            self.copaq6.addItem("CO5")
            self.layout.addWidget(self.copaq6,6,1)
            self.copaq6.currentTextChanged.connect(self.cosee)
            
            self.copaq7= QComboBox()
            #self.copaq7.addItem("CO")
            self.copaq7.addItem("CO1")
            self.copaq7.addItem("CO2")
            self.copaq7.addItem("CO3")
            self.copaq7.addItem("CO4")
            self.copaq7.addItem("CO5")
            self.layout.addWidget(self.copaq7,7,1)
            self.copaq7.currentTextChanged.connect(self.cosee)
            
            self.copaq8= QComboBox()
            #self.copaq8.addItem("CO")
            self.copaq8.addItem("CO1")
            self.copaq8.addItem("CO2")
            self.copaq8.addItem("CO3")
            self.copaq8.addItem("CO4")
            self.copaq8.addItem("CO5")
            self.layout.addWidget(self.copaq8,8,1)
            self.copaq8.currentTextChanged.connect(self.cosee)

            self.copaq9= QComboBox()
            #self.copaq9.addItem("CO")
            self.copaq9.addItem("CO1")
            self.copaq9.addItem("CO2")
            self.copaq9.addItem("CO3")
            self.copaq9.addItem("CO4")
            self.copaq9.addItem("CO5")
            self.layout.addWidget(self.copaq9,9,1)
            self.copaq9.currentTextChanged.connect(self.cosee)

            self.copaq10= QComboBox()
            #self.copaq10.addItem("CO")
            self.copaq10.addItem("CO1")
            self.copaq10.addItem("CO2")
            self.copaq10.addItem("CO3")
            self.copaq10.addItem("CO4")
            self.copaq10.addItem("CO5")
            self.layout.addWidget(self.copaq10,10,1)
            self.copaq10.currentTextChanged.connect(self.cosee)

            self.copaq11= QComboBox()
            #self.copaq11.addItem("CO")
            self.copaq11.addItem("CO1")
            self.copaq11.addItem("CO2")
            self.copaq11.addItem("CO3")
            self.copaq11.addItem("CO4")
            self.copaq11.addItem("CO5")
            self.layout.addWidget(self.copaq11,11,1)
            self.copaq11.currentTextChanged.connect(self.cosee)

            self.copaq12= QComboBox()
            #self.copaq12.addItem("CO")
            self.copaq12.addItem("CO1")
            self.copaq12.addItem("CO2")
            self.copaq12.addItem("CO3")
            self.copaq12.addItem("CO4")
            self.copaq12.addItem("CO5")
            self.layout.addWidget(self.copaq12,12,1)
            self.copaq12.currentTextChanged.connect(self.cosee)

            self.copaq13= QComboBox()
            #self.copaq13.addItem("CO")
            self.copaq13.addItem("CO1")
            self.copaq13.addItem("CO2")
            self.copaq13.addItem("CO3")
            self.copaq13.addItem("CO4")
            self.copaq13.addItem("CO5")
            self.layout.addWidget(self.copaq13,13,1)
            self.copaq13.currentTextChanged.connect(self.cosee)

            self.copaq14= QComboBox()
            #self.copaq14.addItem("CO")
            self.copaq14.addItem("CO1")
            self.copaq14.addItem("CO2")
            self.copaq14.addItem("CO3")
            self.copaq14.addItem("CO4")
            self.copaq14.addItem("CO5")
            self.layout.addWidget(self.copaq14,14,1)
            self.copaq14.currentTextChanged.connect(self.cosee)
            
            self.copaq15= QComboBox()
            #self.copaq15.addItem("CO")
            self.copaq15.addItem("CO1")
            self.copaq15.addItem("CO2")
            self.copaq15.addItem("CO3")
            self.copaq15.addItem("CO4")
            self.copaq15.addItem("CO5")
            self.layout.addWidget(self.copaq15,15,1)
            self.copaq15.currentTextChanged.connect(self.cosee)

            self.copaq16= QComboBox()
            #self.copaq16.addItem("CO")
            self.copaq16.addItem("CO1")
            self.copaq16.addItem("CO2")
            self.copaq16.addItem("CO3")
            self.copaq16.addItem("CO4")
            self.copaq16.addItem("CO5")
            self.layout.addWidget(self.copaq16,16,1)
            self.copaq16.currentTextChanged.connect(self.cosee)

            self.copaq17= QComboBox()
            #self.copaq17.addItem("CO")
            self.copaq17.addItem("CO1")
            self.copaq17.addItem("CO2")
            self.copaq17.addItem("CO3")
            self.copaq17.addItem("CO4")
            self.copaq17.addItem("CO5")
            self.layout.addWidget(self.copaq17,17,1)
            self.copaq17.currentTextChanged.connect(self.cosee)

            self.copaq18= QComboBox()
            #self.copaq18.addItem("CO")
            self.copaq18.addItem("CO1")
            self.copaq18.addItem("CO2")
            self.copaq18.addItem("CO3")
            self.copaq18.addItem("CO4")
            self.copaq18.addItem("CO5")
            self.layout.addWidget(self.copaq18,18,1)
            self.copaq18.currentTextChanged.connect(self.cosee)

            self.copaq19= QComboBox()
            #self.copaq19.addItem("CO")
            self.copaq19.addItem("CO1")
            self.copaq19.addItem("CO2")
            self.copaq19.addItem("CO3")
            self.copaq19.addItem("CO4")
            self.copaq19.addItem("CO5")
            self.layout.addWidget(self.copaq19,19,1)
            self.copaq19.currentTextChanged.connect(self.cosee)

            self.copaq20= QComboBox()
            #self.copaq20.addItem("CO")
            self.copaq20.addItem("CO1")
            self.copaq20.addItem("CO2")
            self.copaq20.addItem("CO3")
            self.copaq20.addItem("CO4")
            self.copaq20.addItem("CO5")
            self.layout.addWidget(self.copaq20,20,1)
            self.copaq20.currentTextChanged.connect(self.cosee)

            self.copaq21= QComboBox()
            #self.copaq21.addItem("CO")
            self.copaq21.addItem("CO1")
            self.copaq21.addItem("CO2")
            self.copaq21.addItem("CO3")
            self.copaq21.addItem("CO4")
            self.copaq21.addItem("CO5")
            self.layout.addWidget(self.copaq21,21,1)
            self.copaq21.currentTextChanged.connect(self.cosee)

            self.copaq22= QComboBox()
            #self.copaq22.addItem("CO")
            self.copaq22.addItem("CO1")
            self.copaq22.addItem("CO2")
            self.copaq22.addItem("CO3")
            self.copaq22.addItem("CO4")
            self.copaq22.addItem("CO5")
            self.layout.addWidget(self.copaq22,22,1)
            self.copaq22.currentTextChanged.connect(self.cosee)

            self.copaq23= QComboBox()
            #self.copaq23.addItem("CO")
            self.copaq23.addItem("CO1")
            self.copaq23.addItem("CO2")
            self.copaq23.addItem("CO3")
            self.copaq23.addItem("CO4")
            self.copaq23.addItem("CO5")
            self.layout.addWidget(self.copaq23,23,1)
            self.copaq23.currentTextChanged.connect(self.cosee)

            self.copaq24= QComboBox()
            #self.copaq24.addItem("CO")
            self.copaq24.addItem("CO1")
            self.copaq24.addItem("CO2")
            self.copaq24.addItem("CO3")
            self.copaq24.addItem("CO4")
            self.copaq24.addItem("CO5")
            self.layout.addWidget(self.copaq24,24,1)
            self.copaq24.currentTextChanged.connect(self.cosee)

            self.copaq25= QComboBox()
            #self.copaq25.addItem("CO")
            self.copaq25.addItem("CO1")
            self.copaq25.addItem("CO2")
            self.copaq25.addItem("CO3")
            self.copaq25.addItem("CO4")
            self.copaq25.addItem("CO5")
            self.layout.addWidget(self.copaq25,25,1)
            self.copaq25.currentTextChanged.connect(self.cosee)

            self.copaq26= QComboBox()
            #self.copaq26.addItem("CO")
            self.copaq26.addItem("CO1")
            self.copaq26.addItem("CO2")
            self.copaq26.addItem("CO3")
            self.copaq26.addItem("CO4")
            self.copaq26.addItem("CO5")
            self.layout.addWidget(self.copaq26,26,1)
            self.copaq26.currentTextChanged.connect(self.cosee)

            self.copaq27= QComboBox()
            #self.copaq27.addItem("CO")
            self.copaq27.addItem("CO1")
            self.copaq27.addItem("CO2")
            self.copaq27.addItem("CO3")
            self.copaq27.addItem("CO4")
            self.copaq27.addItem("CO5")
            self.layout.addWidget(self.copaq27,27,1)
            self.copaq27.currentTextChanged.connect(self.cosee)

            self.copaq28= QComboBox()
            #self.copaq28.addItem("CO")
            self.copaq28.addItem("CO1")
            self.copaq28.addItem("CO2")
            self.copaq28.addItem("CO3")
            self.copaq28.addItem("CO4")
            self.copaq28.addItem("CO5")
            self.layout.addWidget(self.copaq28,28,1)
            self.copaq28.currentTextChanged.connect(self.cosee)

            self.copaq29= QComboBox()
            #self.copaq29.addItem("CO")
            self.copaq29.addItem("CO1")
            self.copaq29.addItem("CO2")
            self.copaq29.addItem("CO3")
            self.copaq29.addItem("CO4")
            self.copaq29.addItem("CO5")
            self.layout.addWidget(self.copaq29,29,1)
            self.copaq29.currentTextChanged.connect(self.cosee)

            self.copaq30= QComboBox()
            #self.copaq30.addItem("CO")
            self.copaq30.addItem("CO1")
            self.copaq30.addItem("CO2")
            self.copaq30.addItem("CO3")
            self.copaq30.addItem("CO4")
            self.copaq30.addItem("CO5")
            self.layout.addWidget(self.copaq30,30,1)
            self.copaq30.currentTextChanged.connect(self.cosee)


           

            
    def subdivisionsforc(self):
        
        if (self.partc==1):
            self.ac1  = QSpinBox()
            #self.ac1.setSuffix("  subdivions")
            self.ac1.setMinimum(1)
            self.ac1.setMaximum(6)
            self.layout.addWidget(self.ac1,1,6)
            self.ac1.valueChanged.connect(self.spinbox_toggledc2)

            
        elif (self.partc==2):
            self.ac1  = QSpinBox()
            #self.ac1.setSuffix("  subdivions")
            self.ac1.setMinimum(1)
            self.ac1.setMaximum(6)
            self.layout.addWidget(self.ac1,1,6)
            self.ac1.valueChanged.connect(self.spinbox_toggledc2)

            self.ac2  = QSpinBox()
            #self.ac2.setSuffix("  subdivions")
            self.ac2.setMinimum(1)
            self.ac2.setMaximum(6)
            self.layout.addWidget(self.ac2,2,6)
            self.ac2.valueChanged.connect(self.spinbox_toggledc2)


        elif (self.partc==3):
            self.ac1  = QSpinBox()
            #self.ac1.setSuffix("  subdivions")
            self.ac1.setMinimum(1)
            self.ac1.setMaximum(6)
            self.layout.addWidget(self.ac1,1,6)
            self.ac1.valueChanged.connect(self.spinbox_toggledc2)

            self.ac2  = QSpinBox()
            #self.ac2.setSuffix("  subdivions")
            self.ac2.setMinimum(1)
            self.ac2.setMaximum(6)
            self.layout.addWidget(self.ac2,2,6)
            self.ac2.valueChanged.connect(self.spinbox_toggledc2)

            self.ac3  = QSpinBox()
            #self.ac3.setSuffix("  subdivions")
            self.ac3.setMinimum(1)
            self.ac3.setMaximum(6)
            self.layout.addWidget(self.ac3,3,6)
            self.ac3.valueChanged.connect(self.spinbox_toggledc2)

        elif (self.partc==4):
            self.ac1  = QSpinBox()
            #self.ac1.setSuffix("  subdivions")
            self.ac1.setMinimum(1)
            self.ac1.setMaximum(6)
            self.layout.addWidget(self.ac1,1,6)
            self.ac1.valueChanged.connect(self.spinbox_toggledc2)

            self.ac2  = QSpinBox()
            #self.ac2.setSuffix("  subdivions")
            self.ac2.setMinimum(1)
            self.ac2.setMaximum(6)
            self.layout.addWidget(self.ac2,2,6)
            self.ac2.valueChanged.connect(self.spinbox_toggledc2)

            self.ac3  = QSpinBox()
            #self.ac3.setSuffix("  subdivions")
            self.ac3.setMinimum(1)
            self.ac3.setMaximum(6)
            self.layout.addWidget(self.ac3,3,6)
            self.ac3.valueChanged.connect(self.spinbox_toggledc2)

            self.ac4  = QSpinBox()
            #self.ac4.setSuffix("  subdivions")
            self.ac4.setMinimum(1)
            self.ac4.setMaximum(6)
            self.layout.addWidget(self.ac4,4,6)
            self.ac4.valueChanged.connect(self.spinbox_toggledc2)

        elif (self.partc==5):
            self.ac1  = QSpinBox()
            #self.ac1.setSuffix("  subdivions")
            self.ac1.setMinimum(1)
            self.ac1.setMaximum(6)
            self.layout.addWidget(self.ac1,1,6)
            self.ac1.valueChanged.connect(self.spinbox_toggledc2)

            self.ac2  = QSpinBox()
            #self.ac2.setSuffix("  subdivions")
            self.ac2.setMinimum(1)
            self.ac2.setMaximum(6)
            self.layout.addWidget(self.ac2,2,6)
            self.ac2.valueChanged.connect(self.spinbox_toggledc2)

            self.ac3  = QSpinBox()
            #self.ac3.setSuffix("  subdivions")
            self.ac3.setMinimum(1)
            self.ac3.setMaximum(6)
            self.layout.addWidget(self.ac3,3,6)
            self.ac3.valueChanged.connect(self.spinbox_toggledc2)

            self.ac4  = QSpinBox()
            #self.ac4.setSuffix("  subdivions")
            self.ac4.setMinimum(1)
            self.ac4.setMaximum(6)
            self.layout.addWidget(self.ac4,4,6)
            self.ac4.valueChanged.connect(self.spinbox_toggledc2)

            self.ac5  = QSpinBox()
            #self.ac5.setSuffix("  subdivions")
            self.ac5.setMinimum(1)
            self.ac5.setMaximum(6)
            self.layout.addWidget(self.ac5,5,6)
            self.ac5.valueChanged.connect(self.spinbox_toggledc2)

        elif (self.partc==6):
            self.ac1  = QSpinBox()
            #self.ac1.setSuffix("  subdivions")
            self.ac1.setMinimum(1)
            self.ac1.setMaximum(6)
            self.layout.addWidget(self.ac1,1,6)
            self.ac1.valueChanged.connect(self.spinbox_toggledc2)

            self.ac2  = QSpinBox()
            #self.ac2.setSuffix("  subdivions")
            self.ac2.setMinimum(1)
            self.ac2.setMaximum(6)
            self.layout.addWidget(self.ac2,2,6)
            self.ac2.valueChanged.connect(self.spinbox_toggledc2)

            self.ac3  = QSpinBox()
            #self.ac3.setSuffix("  subdivions")
            self.ac3.setMinimum(1)
            self.ac3.setMaximum(6)
            self.layout.addWidget(self.ac3,3,6)
            self.ac3.valueChanged.connect(self.spinbox_toggledc2)

            self.ac4  = QSpinBox()
            #self.ac4.setSuffix("  subdivions")
            self.ac4.setMinimum(1)
            self.ac4.setMaximum(6)
            self.layout.addWidget(self.ac4,4,6)
            self.ac4.valueChanged.connect(self.spinbox_toggledc2)

            self.ac5  = QSpinBox()
            #self.ac5.setSuffix("  subdivions")
            self.ac5.setMinimum(1)
            self.ac5.setMaximum(6)
            self.layout.addWidget(self.ac5,5,6)
            self.ac5.valueChanged.connect(self.spinbox_toggledc2)

            self.ac6  = QSpinBox()
            #self.ac6.setSuffix("  subdivions")
            self.ac6.setMinimum(1)
            self.ac6.setMaximum(6)
            self.layout.addWidget(self.ac6,6,6)
            self.ac6.valueChanged.connect(self.spinbox_toggledc2)


        elif (self.partc==7):
            self.ac1  = QSpinBox()
            #self.ac1.setSuffix("  subdivions")
            self.ac1.setMinimum(1)
            self.ac1.setMaximum(6)
            self.layout.addWidget(self.ac1,1,6)
            self.ac1.valueChanged.connect(self.spinbox_toggledc2)

            self.ac2  = QSpinBox()
            #self.ac2.setSuffix("  subdivions")
            self.ac2.setMinimum(1)
            self.ac2.setMaximum(6)
            self.layout.addWidget(self.ac2,2,6)
            self.ac2.valueChanged.connect(self.spinbox_toggledc2)

            self.ac3  = QSpinBox()
            #self.ac3.setSuffix("  subdivions")
            self.ac3.setMinimum(1)
            self.ac3.setMaximum(6)
            self.layout.addWidget(self.ac3,3,6)
            self.ac3.valueChanged.connect(self.spinbox_toggledc2)

            self.ac4  = QSpinBox()
            #self.ac4.setSuffix("  subdivions")
            self.ac4.setMinimum(1)
            self.ac4.setMaximum(6)
            self.layout.addWidget(self.ac4,4,6)
            self.ac4.valueChanged.connect(self.spinbox_toggledc2)

            self.ac5  = QSpinBox()
            #self.ac5.setSuffix("  subdivions")
            self.ac5.setMinimum(1)
            self.ac5.setMaximum(6)
            self.layout.addWidget(self.ac5,5,6)
            self.ac5.valueChanged.connect(self.spinbox_toggledc2)

            self.ac6  = QSpinBox()
            #self.ac6.setSuffix("  subdivions")
            self.ac6.setMinimum(1)
            self.ac6.setMaximum(6)
            self.layout.addWidget(self.ac6,6,6)
            self.ac6.valueChanged.connect(self.spinbox_toggledc2)

            self.ac7  = QSpinBox()
            #self.ac7.setSuffix("  subdivions")
            self.ac7.setMinimum(1)
            self.ac7.setMaximum(6)
            self.layout.addWidget(self.ac7,7,6)
            self.ac7.valueChanged.connect(self.spinbox_toggledc2)


        elif (self.partc==8):
            self.ac1  = QSpinBox()
            #self.ac1.setSuffix("  subdivions")
            self.ac1.setMinimum(1)
            self.ac1.setMaximum(6)
            self.layout.addWidget(self.ac1,1,6)
            self.ac1.valueChanged.connect(self.spinbox_toggledc2)

            self.ac2  = QSpinBox()
            #self.ac2.setSuffix("  subdivions")
            self.ac2.setMinimum(1)
            self.ac2.setMaximum(6)
            self.layout.addWidget(self.ac2,2,6)
            self.ac2.valueChanged.connect(self.spinbox_toggledc2)

            self.ac3  = QSpinBox()
            #self.ac3.setSuffix("  subdivions")
            self.ac3.setMinimum(1)
            self.ac3.setMaximum(6)
            self.layout.addWidget(self.ac3,3,6)
            self.ac3.valueChanged.connect(self.spinbox_toggledc2)

            self.ac4  = QSpinBox()
            #self.ac4.setSuffix("  subdivions")
            self.ac4.setMinimum(1)
            self.ac4.setMaximum(6)
            self.layout.addWidget(self.ac4,4,6)
            self.ac4.valueChanged.connect(self.spinbox_toggledc2)

            self.ac5  = QSpinBox()
            #self.ac5.setSuffix("  subdivions")
            self.ac5.setMinimum(1)
            self.ac5.setMaximum(6)
            self.layout.addWidget(self.ac5,5,6)
            self.ac5.valueChanged.connect(self.spinbox_toggledc2)

            self.ac6  = QSpinBox()
            #self.ac6.setSuffix("  subdivions")
            self.ac6.setMinimum(1)
            self.ac6.setMaximum(6)
            self.layout.addWidget(self.ac6,6,6)
            self.ac6.valueChanged.connect(self.spinbox_toggledc2)

            self.ac7  = QSpinBox()
            #self.ac7.setSuffix("  subdivions")
            self.ac7.setMinimum(1)
            self.ac7.setMaximum(6)
            self.layout.addWidget(self.ac7,7,6)
            self.ac7.valueChanged.connect(self.spinbox_toggledc2)

            self.ac8  = QSpinBox()
            #self.ac8.setSuffix("  subdivions")
            self.ac8.setMinimum(1)
            self.ac8.setMaximum(6)
            self.layout.addWidget(self.ac8,8,6)
            self.ac8.valueChanged.connect(self.spinbox_toggledc2)


        elif (self.partc==9):
            self.ac1  = QSpinBox()
            #self.ac1.setSuffix("  subdivions")
            self.ac1.setMinimum(1)
            self.ac1.setMaximum(6)
            self.layout.addWidget(self.ac1,1,6)
            self.ac1.valueChanged.connect(self.spinbox_toggledc2)

            self.ac2  = QSpinBox()
            #self.ac2.setSuffix("  subdivions")
            self.ac2.setMinimum(1)
            self.ac2.setMaximum(6)
            self.layout.addWidget(self.ac2,2,6)
            self.ac2.valueChanged.connect(self.spinbox_toggledc2)

            self.ac3  = QSpinBox()
            #self.ac3.setSuffix("  subdivions")
            self.ac3.setMinimum(1)
            self.ac3.setMaximum(6)
            self.layout.addWidget(self.ac3,3,6)
            self.ac3.valueChanged.connect(self.spinbox_toggledc2)

            self.ac4  = QSpinBox()
            #self.ac4.setSuffix("  subdivions")
            self.ac4.setMinimum(1)
            self.ac4.setMaximum(6)
            self.layout.addWidget(self.ac4,4,6)
            self.ac4.valueChanged.connect(self.spinbox_toggledc2)

            self.ac5  = QSpinBox()
            #self.ac5.setSuffix("  subdivions")
            self.ac5.setMinimum(1)
            self.ac5.setMaximum(6)
            self.layout.addWidget(self.ac5,5,6)
            self.ac5.valueChanged.connect(self.spinbox_toggledc2)

            self.ac6  = QSpinBox()
            #self.ac6.setSuffix("  subdivions")
            self.ac6.setMinimum(1)
            self.ac6.setMaximum(6)
            self.layout.addWidget(self.ac6,6,6)
            self.ac6.valueChanged.connect(self.spinbox_toggledc2)

            self.ac7  = QSpinBox()
            #self.ac7.setSuffix("  subdivions")
            self.ac7.setMinimum(1)
            self.ac7.setMaximum(6)
            self.layout.addWidget(self.ac7,7,6)
            self.ac7.valueChanged.connect(self.spinbox_toggledc2)

            self.ac8  = QSpinBox()
            #self.ac8.setSuffix("  subdivions")
            self.ac8.setMinimum(1)
            self.ac8.setMaximum(6)
            self.layout.addWidget(self.ac8,8,6)
            self.ac8.valueChanged.connect(self.spinbox_toggledc2)

            self.ac9  = QSpinBox()
            #self.ac9.setSuffix("  subdivions")
            self.a9.setMinimum(1)
            self.ac9.setMaximum(6)
            self.layout.addWidget(self.ac9,9,6)
            self.ac9.valueChanged.connect(self.spinbox_toggledc2)

           
           

        elif (self.partc==10):
            self.ac1  = QSpinBox()
            #self.ac1.setSuffix("  subdivions")
            self.ac1.setMinimum(1)
            self.ac1.setMaximum(6)
            self.layout.addWidget(self.ac1,1,6)
            self.ac1.valueChanged.connect(self.spinbox_toggledc2)

            self.ac2  = QSpinBox()
            #self.ac2.setSuffix("  subdivions")
            self.ac2.setMinimum(1)
            self.ac2.setMaximum(6)
            self.layout.addWidget(self.ac2,2,6)
            self.ac2.valueChanged.connect(self.spinbox_toggledc2)

            self.ac3  = QSpinBox()
            #self.ac3.setSuffix("  subdivions")
            self.ac3.setMinimum(1)
            self.ac3.setMaximum(6)
            self.layout.addWidget(self.ac3,3,6)
            self.ac3.valueChanged.connect(self.spinbox_toggledc2)

            self.ac4  = QSpinBox()
            #self.ac4.setSuffix("  subdivions")
            self.ac4.setMinimum(1)
            self.ac4.setMaximum(6)
            self.layout.addWidget(self.ac4,4,6)
            self.ac4.valueChanged.connect(self.spinbox_toggledc2)

            self.ac5  = QSpinBox()
            #self.ac5.setSuffix("  subdivions")
            self.ac5.setMinimum(1)
            self.ac5.setMaximum(6)
            self.layout.addWidget(self.ac5,5,6)
            self.ac5.valueChanged.connect(self.spinbox_toggledc2)

            self.ac6  = QSpinBox()
            #self.ac6.setSuffix("  subdivions")
            self.ac6.setMinimum(1)
            self.ac6.setMaximum(6)
            self.layout.addWidget(self.ac6,6,6)
            self.ac6.valueChanged.connect(self.spinbox_toggledc2)

            self.ac7  = QSpinBox()
            #self.ac7.setSuffix("  subdivions")
            self.ac7.setMinimum(1)
            self.ac7.setMaximum(6)
            self.layout.addWidget(self.ac7,7,6)
            self.ac7.valueChanged.connect(self.spinbox_toggledc2)

            self.ac8  = QSpinBox()
            #self.ac8.setSuffix("  subdivions")
            self.ac8.setMinimum(1)
            self.ac8.setMaximum(6)
            self.layout.addWidget(self.ac8,8,6)
            self.ac8.valueChanged.connect(self.spinbox_toggledc2)

            self.ac9  = QSpinBox()
            #self.ac9.setSuffix("  subdivions")
            self.ac9.setMinimum(1)
            self.ac9.setMaximum(6)
            self.layout.addWidget(self.ac9,9,6)
            self.ac9.valueChanged.connect(self.spinbox_toggledc2)

            self.ac10  = QSpinBox()
            #self.ac10.setSuffix("  subdivions")
            self.ac10.setMinimum(1)
            self.ac10.setMaximum(6)
            self.layout.addWidget(self.ac10,10,6)
            self.ac10.valueChanged.connect(self.spinbox_toggledc2)
    def spinbox_toggledc2(self):
       
        self.partcsub=[]
        if (self.partc==0):
            self.partcsub=[0]
            print(self.partcsub)
        elif (self.partc==1):
            self.qsc1=self.ac1.value()
            
            self.partcsub.append(self.qsc1)
            print(self.partcsub)
            
        elif (self.partc==2):
            self.qsc1=self.ac1.value()
            self.qsc2=self.ac2.value()
            
            self.partcsub.append(self.qsc1)
            self.partcsub.append(self.qsc2)
            print(self.partcsub)
            
        elif (self.partc==3):
            self.qsc1=self.ac1.value()
            self.qsc2=self.ac2.value()
            self.qsc3=self.ac3.value()
    
            self.partcsub.append(self.qsc1)
            self.partcsub.append(self.qsc2)
            self.partcsub.append(self.qsc3)
            print(self.partcsub)
           
        elif (self.partc==4):
            self.qsc1=self.ac1.value()
            self.qsc2=self.ac2.value()
            self.qsc3=self.ac3.value()
            self.qsc4=self.ac4.value()
            
            self.partcsub.append(self.qsc1)
            self.partcsub.append(self.qsc2)
            self.partcsub.append(self.qsc3)
            self.partcsub.append(self.qsc4)
            print(self.partcsub)
            
        elif (self.partc==5):
            self.qsc1=self.ac1.value()
            self.qsc2=self.ac2.value()
            self.qsc3=self.ac3.value()
            self.qsc4=self.ac4.value()
            self.qsc5=self.ac5.value()
            
            self.partcsub.append(self.qsc1)
            self.partcsub.append(self.qsc2)
            self.partcsub.append(self.qsc3)
            self.partcsub.append(self.qsc4)
            self.partcsub.append(self.qsc5)
            print(self.partcsub)
            
        elif (self.partc==6):
            self.qsc1=self.ac1.value()
            self.qsc2=self.ac2.value()
            self.qsc3=self.ac3.value()
            self.qsc4=self.ac4.value()
            self.qsc5=self.ac5.value()
            self.qsc6=self.ac6.value()
            
            self.partcsub.append(self.qsc1)
            self.partcsub.append(self.qsc2)
            self.partcsub.append(self.qsc3)
            self.partcsub.append(self.qsc4)
            self.partcsub.append(self.qsc5)
            self.partcsub.append(self.qsc6)
            print(self.partcsub)
            
        elif (self.partc==7):
            self.qsc1=self.ac1.value()
            self.qsc2=self.ac2.value()
            self.qsc3=self.ac3.value()
            self.qsc4=self.ac4.value()
            self.qsc5=self.ac5.value()
            self.qsc6=self.ac6.value()
            self.qsc7=self.ac7.value()
            
            self.partcsub.append(self.qsc1)
            self.partcsub.append(self.qsc2)
            self.partcsub.append(self.qsc3)
            self.partcsub.append(self.qsc4)
            self.partcsub.append(self.qsc5)
            self.partcsub.append(self.qsc6)
            self.partcsub.append(self.qsc7)
            print(self.partcsub)
            
        elif (self.partc==8):
            self.qsc1=self.ac1.value()
            self.qsc2=self.ac2.value()
            self.qsc3=self.ac3.value()
            self.qsc4=self.ac4.value()
            self.qsc5=self.ac5.value()
            self.qsc6=self.ac6.value()
            self.qsc7=self.ac7.value()
            self.qsc8=self.ac8.value()
            
            self.partcsub.append(self.qsc1)
            self.partcsub.append(self.qsc2)
            self.partcsub.append(self.qsc3)
            self.partcsub.append(self.qsc4)
            self.partcsub.append(self.qsc5)
            self.partcsub.append(self.qsc6)
            self.partcsub.append(self.qsc7)
            self.partcsub.append(self.qsc8)
            print(self.partcsub)

            
        elif (self.partc==9):
            self.qsc1=self.ac1.value()
            self.qsc2=self.ac2.value()
            self.qsc3=self.ac3.value()
            self.qsc4=self.ac4.value()
            self.qsc5=self.ac5.value()
            self.qsc6=self.ac6.value()
            self.qsc7=self.ac7.value()
            self.qsc8=self.ac8.value()
            self.qsc9=self.ac9.value()
            self.partcsub.append(self.qsc1)
            self.partcsub.append(self.qsc2)
            self.partcsub.append(self.qsc3)
            self.partcsub.append(self.qsc4)
            self.partcsub.append(self.qsc5)
            self.partcsub.append(self.qsc6)
            self.partcsub.append(self.qsc7)
            self.partcsub.append(self.qsc8)
            self.partcsub.append(self.qsc9)
            print(self.partcsub)
            
        elif (self.partc==10):
            #print("enter the 10")
            self.qsc1=self.ac1.value()
            self.qsc2=self.ac2.value()
            self.qsc3=self.ac3.value()
            self.qsc4=self.ac4.value()
            self.qsc5=self.ac5.value()
            self.qsc6=self.ac6.value()
            self.qsc7=self.ac7.value()
            self.qsc8=self.ac8.value()
            self.qsc9=self.ac9.value()
            self.qsc10=self.ac10.value()
            self.partcsub.append(self.qsc1)
            self.partcsub.append(self.qsc2)
            self.partcsub.append(self.qsc3)
            self.partcsub.append(self.qsc4)
            self.partcsub.append(self.qsc5)
            self.partcsub.append(self.qsc6)
            self.partcsub.append(self.qsc7)
            self.partcsub.append(self.qsc8)
            self.partcsub.append(self.qsc9)
            self.partcsub.append(self.qsc10)
            print(self.partcsub)
        global partcsub
        partcsub=self.partcsub

    def coursec(self):

        if (self.partc==1):
            self.copcq1= QComboBox()
            #self.copcq1.addItem("CO")

            self.copcq1.addItem("CO1")
            self.copcq1.addItem("CO2")
            self.copcq1.addItem("CO3")
            self.copcq1.addItem("CO4")
            self.copcq1.addItem("CO5")
            self.layout.addWidget(self.copcq1,1,7)
            self.copcq1.currentTextChanged.connect(self.coseec)
            
        elif (self.partc==2):
            self.copcq1= QComboBox()
            #self.copcq1.addItem("CO")

            self.copcq1.addItem("CO1")
            self.copcq1.addItem("CO2")
            self.copcq1.addItem("CO3")
            self.copcq1.addItem("CO4")
            self.copcq1.addItem("CO5")
            self.layout.addWidget(self.copcq1,1,7)
            self.copcq1.currentTextChanged.connect(self.coseec)
            self.copcq2= QComboBox()
            #self.copcq2.addItem("CO")

            self.copcq2.addItem("CO1")
            self.copcq2.addItem("CO2")
            self.copcq2.addItem("CO3")
            self.copcq2.addItem("CO4")
            self.copcq2.addItem("CO5")
            self.layout.addWidget(self.copcq2,2,7)
            self.copcq2.currentTextChanged.connect(self.coseec)
            

        elif (self.partc==3):
            self.copcq1= QComboBox()
            #self.copcq1.addItem("CO")

            self.copcq1.addItem("CO1")
            self.copcq1.addItem("CO2")
            self.copcq1.addItem("CO3")
            self.copcq1.addItem("CO4")
            self.copcq1.addItem("CO5")
            self.layout.addWidget(self.copcq1,1,7)
            self.copcq1.currentTextChanged.connect(self.coseec)
            self.copcq2= QComboBox()
            #self.copcq2.addItem("CO")

            self.copcq2.addItem("CO1")
            self.copcq2.addItem("CO2")
            self.copcq2.addItem("CO3")
            self.copcq2.addItem("CO4")
            self.copcq2.addItem("CO5")
            self.layout.addWidget(self.copcq2,2,7)
            self.copcq2.currentTextChanged.connect(self.coseec)
            
            self.copcq3= QComboBox()
            #self.copcq3.addItem("CO")

            self.copcq3.addItem("CO1")
            self.copcq3.addItem("CO2")
            self.copcq3.addItem("CO3")
            self.copcq3.addItem("CO4")
            self.copcq3.addItem("CO5")
            self.layout.addWidget(self.copcq3,3,7)
            
            self.copcq3.currentTextChanged.connect(self.coseec)
            

        elif (self.partc==4):
            self.copcq1= QComboBox()
            #self.copcq1.addItem("CO")

            self.copcq1.addItem("CO1")
            self.copcq1.addItem("CO2")
            self.copcq1.addItem("CO3")
            self.copcq1.addItem("CO4")
            self.copcq1.addItem("CO5")
            self.layout.addWidget(self.copcq1,1,7)
            self.copcq1.currentTextChanged.connect(self.coseec)
            self.copcq2= QComboBox()
            #self.copcq2.addItem("CO")

            self.copcq2.addItem("CO1")
            self.copcq2.addItem("CO2")
            self.copcq2.addItem("CO3")
            self.copcq2.addItem("CO4")
            self.copcq2.addItem("CO5")
            self.layout.addWidget(self.copcq2,2,7)
            self.copcq2.currentTextChanged.connect(self.coseec)
            
            self.copcq3= QComboBox()
            #self.copcq3.addItem("CO")

            self.copcq3.addItem("CO1")
            self.copcq3.addItem("CO2")
            self.copcq3.addItem("CO3")
            self.copcq3.addItem("CO4")
            self.copcq3.addItem("CO5")
            self.layout.addWidget(self.copcq3,3,7)
            
            self.copcq3.currentTextChanged.connect(self.coseec)
            
            self.copcq4= QComboBox()
            #self.copcq4.addItem("CO")
            self.copcq4.addItem("CO1")
            self.copcq4.addItem("CO2")
            self.copcq4.addItem("CO3")
            self.copcq4.addItem("CO4")
            self.copcq4.addItem("CO5")
            self.layout.addWidget(self.copcq4,4,7)
            self.copcq4.currentTextChanged.connect(self.coseec)
            

        elif (self.partc==5):
            self.copcq1= QComboBox()
            #self.copcq1.addItem("CO")

            self.copcq1.addItem("CO1")
            self.copcq1.addItem("CO2")
            self.copcq1.addItem("CO3")
            self.copcq1.addItem("CO4")
            self.copcq1.addItem("CO5")
            self.layout.addWidget(self.copcq1,1,7)
            self.copcq1.currentTextChanged.connect(self.coseec)
            self.copcq2= QComboBox()
            #self.copcq2.addItem("CO")

            self.copcq2.addItem("CO1")
            self.copcq2.addItem("CO2")
            self.copcq2.addItem("CO3")
            self.copcq2.addItem("CO4")
            self.copcq2.addItem("CO5")
            self.layout.addWidget(self.copcq2,2,7)
            self.copcq2.currentTextChanged.connect(self.coseec)
            
            self.copcq3= QComboBox()
            #self.copcq3.addItem("CO")

            self.copcq3.addItem("CO1")
            self.copcq3.addItem("CO2")
            self.copcq3.addItem("CO3")
            self.copcq3.addItem("CO4")
            self.copcq3.addItem("CO5")
            self.layout.addWidget(self.copcq3,3,7)
            
            self.copcq3.currentTextChanged.connect(self.coseec)
            
            self.copcq4= QComboBox()
            #self.copcq4.addItem("CO")
            self.copcq4.addItem("CO1")
            self.copcq4.addItem("CO2")
            self.copcq4.addItem("CO3")
            self.copcq4.addItem("CO4")
            self.copcq4.addItem("CO5")
            self.layout.addWidget(self.copcq4,4,7)
            self.copcq4.currentTextChanged.connect(self.coseec)
            
            self.copcq5= QComboBox()
            #self.copcq5.addItem("CO")
            self.copcq5.addItem("CO1")
            self.copcq5.addItem("CO2")
            self.copcq5.addItem("CO3")
            self.copcq5.addItem("CO4")
            self.copcq5.addItem("CO5")
            self.layout.addWidget(self.copcq5,5,7)
            self.copcq5.currentTextChanged.connect(self.coseec)
            

        elif (self.partc==6):
            self.copcq1= QComboBox()
            #self.copcq1.addItem("CO")

            self.copcq1.addItem("CO1")
            self.copcq1.addItem("CO2")
            self.copcq1.addItem("CO3")
            self.copcq1.addItem("CO4")
            self.copcq1.addItem("CO5")
            self.layout.addWidget(self.copcq1,1,7)
            self.copcq1.currentTextChanged.connect(self.coseec)
            self.copcq2= QComboBox()
            #self.copcq2.addItem("CO")

            self.copcq2.addItem("CO1")
            self.copcq2.addItem("CO2")
            self.copcq2.addItem("CO3")
            self.copcq2.addItem("CO4")
            self.copcq2.addItem("CO5")
            self.layout.addWidget(self.copcq2,2,7)
            self.copcq2.currentTextChanged.connect(self.coseec)
            
            self.copcq3= QComboBox()
            #self.copcq3.addItem("CO")

            self.copcq3.addItem("CO1")
            self.copcq3.addItem("CO2")
            self.copcq3.addItem("CO3")
            self.copcq3.addItem("CO4")
            self.copcq3.addItem("CO5")
            self.layout.addWidget(self.copcq3,3,7)
            
            self.copcq3.currentTextChanged.connect(self.coseec)
            
            self.copcq4= QComboBox()
            #self.copcq4.addItem("CO")
            self.copcq4.addItem("CO1")
            self.copcq4.addItem("CO2")
            self.copcq4.addItem("CO3")
            self.copcq4.addItem("CO4")
            self.copcq4.addItem("CO5")
            self.layout.addWidget(self.copcq4,4,7)
            self.copcq4.currentTextChanged.connect(self.coseec)
            
            self.copcq5= QComboBox()
            #self.copcq5.addItem("CO")
            self.copcq5.addItem("CO1")
            self.copcq5.addItem("CO2")
            self.copcq5.addItem("CO3")
            self.copcq5.addItem("CO4")
            self.copcq5.addItem("CO5")
            self.layout.addWidget(self.copcq5,5,7)
            self.copcq5.currentTextChanged.connect(self.coseec)
            
            self.copcq6= QComboBox()
            #self.copcq6.addItem("CO")
            self.copcq6.addItem("CO1")
            self.copcq6.addItem("CO2")
            self.copcq6.addItem("CO3")
            self.copcq6.addItem("CO4")
            self.copcq6.addItem("CO5")
            self.layout.addWidget(self.copcq6,6,7)
            self.copcq6.currentTextChanged.connect(self.coseec)
            

        elif (self.partc==7):
            self.copcq1= QComboBox()
            #self.copcq1.addItem("CO")

            self.copcq1.addItem("CO1")
            self.copcq1.addItem("CO2")
            self.copcq1.addItem("CO3")
            self.copcq1.addItem("CO4")
            self.copcq1.addItem("CO5")
            self.layout.addWidget(self.copcq1,1,7)
            self.copcq1.currentTextChanged.connect(self.coseec)
            self.copcq2= QComboBox()
            #self.copcq2.addItem("CO")

            self.copcq2.addItem("CO1")
            self.copcq2.addItem("CO2")
            self.copcq2.addItem("CO3")
            self.copcq2.addItem("CO4")
            self.copcq2.addItem("CO5")
            self.layout.addWidget(self.copcq2,2,7)
            self.copcq2.currentTextChanged.connect(self.coseec)
            
            self.copcq3= QComboBox()
            #self.copcq3.addItem("CO")

            self.copcq3.addItem("CO1")
            self.copcq3.addItem("CO2")
            self.copcq3.addItem("CO3")
            self.copcq3.addItem("CO4")
            self.copcq3.addItem("CO5")
            self.layout.addWidget(self.copcq3,3,7)
            
            self.copcq3.currentTextChanged.connect(self.coseec)
            
            self.copcq4= QComboBox()
            #self.copcq4.addItem("CO")
            self.copcq4.addItem("CO1")
            self.copcq4.addItem("CO2")
            self.copcq4.addItem("CO3")
            self.copcq4.addItem("CO4")
            self.copcq4.addItem("CO5")
            self.layout.addWidget(self.copcq4,4,7)
            self.copcq4.currentTextChanged.connect(self.coseec)
            
            self.copcq5= QComboBox()
            #self.copcq5.addItem("CO")
            self.copcq5.addItem("CO1")
            self.copcq5.addItem("CO2")
            self.copcq5.addItem("CO3")
            self.copcq5.addItem("CO4")
            self.copcq5.addItem("CO5")
            self.layout.addWidget(self.copcq5,5,7)
            self.copcq5.currentTextChanged.connect(self.coseec)
            
            self.copcq6= QComboBox()
            #self.copcq6.addItem("CO")
            self.copcq6.addItem("CO1")
            self.copcq6.addItem("CO2")
            self.copcq6.addItem("CO3")
            self.copcq6.addItem("CO4")
            self.copcq6.addItem("CO5")
            self.layout.addWidget(self.copcq6,6,7)
            self.copcq6.currentTextChanged.connect(self.coseec)
            
            self.copcq7= QComboBox()
            #self.copcq7.addItem("CO")
            self.copcq7.addItem("CO1")
            self.copcq7.addItem("CO2")
            self.copcq7.addItem("CO3")
            self.copcq7.addItem("CO4")
            self.copcq7.addItem("CO5")
            self.layout.addWidget(self.copcq7,7,7)
            self.copcq7.currentTextChanged.connect(self.coseec)
 
        elif (self.partc==8):
            self.copcq1= QComboBox()
            #self.copcq1.addItem("CO")

            self.copcq1.addItem("CO1")
            self.copcq1.addItem("CO2")
            self.copcq1.addItem("CO3")
            self.copcq1.addItem("CO4")
            self.copcq1.addItem("CO5")
            self.layout.addWidget(self.copcq1,1,7)
            self.copcq1.currentTextChanged.connect(self.coseec)
            self.copcq2= QComboBox()
            #self.copcq2.addItem("CO")

            self.copcq2.addItem("CO1")
            self.copcq2.addItem("CO2")
            self.copcq2.addItem("CO3")
            self.copcq2.addItem("CO4")
            self.copcq2.addItem("CO5")
            self.layout.addWidget(self.copcq2,2,7)
            self.copcq2.currentTextChanged.connect(self.coseec)
            
            self.copcq3= QComboBox()
            #self.copcq3.addItem("CO")

            self.copcq3.addItem("CO1")
            self.copcq3.addItem("CO2")
            self.copcq3.addItem("CO3")
            self.copcq3.addItem("CO4")
            self.copcq3.addItem("CO5")
            self.layout.addWidget(self.copcq3,3,7)
            
            self.copcq3.currentTextChanged.connect(self.coseec)
            
            self.copcq4= QComboBox()
            #self.copcq4.addItem("CO")
            self.copcq4.addItem("CO1")
            self.copcq4.addItem("CO2")
            self.copcq4.addItem("CO3")
            self.copcq4.addItem("CO4")
            self.copcq4.addItem("CO5")
            self.layout.addWidget(self.copcq4,4,7)
            self.copcq4.currentTextChanged.connect(self.coseec)
            
            self.copcq5= QComboBox()
            #self.copcq5.addItem("CO")
            self.copcq5.addItem("CO1")
            self.copcq5.addItem("CO2")
            self.copcq5.addItem("CO3")
            self.copcq5.addItem("CO4")
            self.copcq5.addItem("CO5")
            self.layout.addWidget(self.copcq5,5,7)
            self.copcq5.currentTextChanged.connect(self.coseec)
            
            self.copcq6= QComboBox()
            #self.copcq6.addItem("CO")
            self.copcq6.addItem("CO1")
            self.copcq6.addItem("CO2")
            self.copcq6.addItem("CO3")
            self.copcq6.addItem("CO4")
            self.copcq6.addItem("CO5")
            self.layout.addWidget(self.copcq6,6,7)
            self.copcq6.currentTextChanged.connect(self.coseec)
            
            self.copcq7= QComboBox()
            #self.copcq7.addItem("CO")
            self.copcq7.addItem("CO1")
            self.copcq7.addItem("CO2")
            self.copcq7.addItem("CO3")
            self.copcq7.addItem("CO4")
            self.copcq7.addItem("CO5")
            self.layout.addWidget(self.copcq7,7,7)
            self.copcq7.currentTextChanged.connect(self.coseec)
            
            self.copcq8= QComboBox()
            #self.copcq8.addItem("CO")
            self.copcq8.addItem("CO1")
            self.copcq8.addItem("CO2")
            self.copcq8.addItem("CO3")
            self.copcq8.addItem("CO4")
            self.copcq8.addItem("CO5")
            self.layout.addWidget(self.copcq8,8,7)
            self.copcq8.currentTextChanged.connect(self.coseec)


        elif (self.partc==9):
            self.copcq1= QComboBox()
            #self.copcq1.addItem("CO")

            self.copcq1.addItem("CO1")
            self.copcq1.addItem("CO2")
            self.copcq1.addItem("CO3")
            self.copcq1.addItem("CO4")
            self.copcq1.addItem("CO5")
            self.layout.addWidget(self.copcq1,1,7)
            self.copcq1.currentTextChanged.connect(self.coseec)
            self.copcq2= QComboBox()
            #self.copcq2.addItem("CO")

            self.copcq2.addItem("CO1")
            self.copcq2.addItem("CO2")
            self.copcq2.addItem("CO3")
            self.copcq2.addItem("CO4")
            self.copcq2.addItem("CO5")
            self.layout.addWidget(self.copcq2,2,7)
            self.copcq2.currentTextChanged.connect(self.coseec)
            
            self.copcq3= QComboBox()
            #self.copcq3.addItem("CO")

            self.copcq3.addItem("CO1")
            self.copcq3.addItem("CO2")
            self.copcq3.addItem("CO3")
            self.copcq3.addItem("CO4")
            self.copcq3.addItem("CO5")
            self.layout.addWidget(self.copcq3,3,7)
            
            self.copcq3.currentTextChanged.connect(self.coseec)
            
            self.copcq4= QComboBox()
            #self.copcq4.addItem("CO")
            self.copcq4.addItem("CO1")
            self.copcq4.addItem("CO2")
            self.copcq4.addItem("CO3")
            self.copcq4.addItem("CO4")
            self.copcq4.addItem("CO5")
            self.layout.addWidget(self.copcq4,4,7)
            self.copcq4.currentTextChanged.connect(self.coseec)
            
            self.copcq5= QComboBox()
            #self.copcq5.addItem("CO")
            self.copcq5.addItem("CO1")
            self.copcq5.addItem("CO2")
            self.copcq5.addItem("CO3")
            self.copcq5.addItem("CO4")
            self.copcq5.addItem("CO5")
            self.layout.addWidget(self.copcq5,5,7)
            self.copcq5.currentTextChanged.connect(self.coseec)
            
            self.copcq6= QComboBox()
            #self.copcq6.addItem("CO")
            self.copcq6.addItem("CO1")
            self.copcq6.addItem("CO2")
            self.copcq6.addItem("CO3")
            self.copcq6.addItem("CO4")
            self.copcq6.addItem("CO5")
            self.layout.addWidget(self.copcq6,6,7)
            self.copcq6.currentTextChanged.connect(self.coseec)
            
            self.copcq7= QComboBox()
            #self.copcq7.addItem("CO")
            self.copcq7.addItem("CO1")
            self.copcq7.addItem("CO2")
            self.copcq7.addItem("CO3")
            self.copcq7.addItem("CO4")
            self.copcq7.addItem("CO5")
            self.layout.addWidget(self.copcq7,7,7)
            self.copcq7.currentTextChanged.connect(self.coseec)
            
            self.copcq8= QComboBox()
            #self.copcq8.addItem("CO")
            self.copcq8.addItem("CO1")
            self.copcq8.addItem("CO2")
            self.copcq8.addItem("CO3")
            self.copcq8.addItem("CO4")
            self.copcq8.addItem("CO5")
            self.layout.addWidget(self.copcq8,8,7)
            self.copcq8.currentTextChanged.connect(self.coseec)

            self.copcq9= QComboBox()
            #self.copcq9.addItem("CO")
            self.copcq9.addItem("CO1")
            self.copcq9.addItem("CO2")
            self.copcq9.addItem("CO3")
            self.copcq9.addItem("CO4")
            self.copcq9.addItem("CO5")
            self.layout.addWidget(self.copcq9,9,7)
            self.copcq9.currentTextChanged.connect(self.coseec)

            
        elif (self.partc==10):
            self.copcq1= QComboBox()
            #self.copcq1.addItem("CO")

            self.copcq1.addItem("CO1")
            self.copcq1.addItem("CO2")
            self.copcq1.addItem("CO3")
            self.copcq1.addItem("CO4")
            self.copcq1.addItem("CO5")
            self.layout.addWidget(self.copcq1,1,7)
            self.copcq1.currentTextChanged.connect(self.coseec)
            self.copcq2= QComboBox()
            #self.copcq2.addItem("CO")

            self.copcq2.addItem("CO1")
            self.copcq2.addItem("CO2")
            self.copcq2.addItem("CO3")
            self.copcq2.addItem("CO4")
            self.copcq2.addItem("CO5")
            self.layout.addWidget(self.copcq2,2,7)
            self.copcq2.currentTextChanged.connect(self.coseec)
            
            self.copcq3= QComboBox()
            #self.copcq3.addItem("CO")

            self.copcq3.addItem("CO1")
            self.copcq3.addItem("CO2")
            self.copcq3.addItem("CO3")
            self.copcq3.addItem("CO4")
            self.copcq3.addItem("CO5")
            self.layout.addWidget(self.copcq3,3,7)
            
            self.copcq3.currentTextChanged.connect(self.coseec)
            
            self.copcq4= QComboBox()
            #self.copcq4.addItem("CO")
            self.copcq4.addItem("CO1")
            self.copcq4.addItem("CO2")
            self.copcq4.addItem("CO3")
            self.copcq4.addItem("CO4")
            self.copcq4.addItem("CO5")
            self.layout.addWidget(self.copcq4,4,7)
            self.copcq4.currentTextChanged.connect(self.coseec)
            
            self.copcq5= QComboBox()
            #self.copcq5.addItem("CO")
            self.copcq5.addItem("CO1")
            self.copcq5.addItem("CO2")
            self.copcq5.addItem("CO3")
            self.copcq5.addItem("CO4")
            self.copcq5.addItem("CO5")
            self.layout.addWidget(self.copcq5,5,7)
            self.copcq5.currentTextChanged.connect(self.coseec)
            
            self.copcq6= QComboBox()
            #self.copcq6.addItem("CO")
            self.copcq6.addItem("CO1")
            self.copcq6.addItem("CO2")
            self.copcq6.addItem("CO3")
            self.copcq6.addItem("CO4")
            self.copcq6.addItem("CO5")
            self.layout.addWidget(self.copcq6,6,7)
            self.copcq6.currentTextChanged.connect(self.coseec)
            
            self.copcq7= QComboBox()
            #self.copcq7.addItem("CO")
            self.copcq7.addItem("CO1")
            self.copcq7.addItem("CO2")
            self.copcq7.addItem("CO3")
            self.copcq7.addItem("CO4")
            self.copcq7.addItem("CO5")
            self.layout.addWidget(self.copcq7,7,7)
            self.copcq7.currentTextChanged.connect(self.coseec)
            
            self.copcq8= QComboBox()
            #self.copcq8.addItem("CO")
            self.copcq8.addItem("CO1")
            self.copcq8.addItem("CO2")
            self.copcq8.addItem("CO3")
            self.copcq8.addItem("CO4")
            self.copcq8.addItem("CO5")
            self.layout.addWidget(self.copcq8,8,7)
            self.copcq8.currentTextChanged.connect(self.coseec)

            self.copcq9= QComboBox()
            #self.copcq9.addItem("CO")
            self.copcq9.addItem("CO1")
            self.copcq9.addItem("CO2")
            self.copcq9.addItem("CO3")
            self.copcq9.addItem("CO4")
            self.copcq9.addItem("CO5")
            self.layout.addWidget(self.copcq9,9,7)
            self.copcq9.currentTextChanged.connect(self.coseec)

            self.copcq10= QComboBox()
            #self.copcq10.addItem("CO")
            self.copcq10.addItem("CO1")
            self.copcq10.addItem("CO2")
            self.copcq10.addItem("CO3")
            self.copcq10.addItem("CO4")
            self.copcq10.addItem("CO5")
            self.layout.addWidget(self.copcq10,10,7)
            self.copcq10.currentTextChanged.connect(self.coseec)

    def coseec(self):
        self.partccos=[]
        
        if (self.partc==1):
            q1 = self.copcq1.currentText()
            
            self.partccos.append(q1)
             
            print(self.partccos)

        elif (self.partc==2):
            q1 = self.copcq1.currentText()
            q2 = self.copcq2.currentText()
            
            self.partccos.append(q1)
            self.partccos.append(q2)
             
            print(self.partccos)
        elif (self.partc==3):
            q1 = self.copcq1.currentText()
            q2 = self.copcq2.currentText()
            q3 = self.copcq3.currentText()
            
            self.partccos.append(q1)
            self.partccos.append(q2)
            self.partccos.append(q3)
             
            print(self.partccos)
        elif (self.partc==4):
            q1 = self.copcq1.currentText()
            q2 = self.copcq2.currentText()
            q3 = self.copcq3.currentText()
            q4 = self.copcq4.currentText()
            
            self.partccos.append(q1)
            self.partccos.append(q2)
            self.partccos.append(q3)
            self.partccos.append(q4)
             
            print(self.partccos)
        elif (self.partc==4):
            q1 = self.copcq1.currentText()
            q2 = self.copcq2.currentText()
            q3 = self.copcq3.currentText()
            q4 = self.copcq4.currentText()
            
            self.partccos.append(q1)
            self.partccos.append(q2)
            self.partccos.append(q3)
            self.partccos.append(q4)
             
            print(self.partccos)
        elif (self.partc==5):
            q1 = self.copcq1.currentText()
            q2 = self.copcq2.currentText()
            q3 = self.copcq3.currentText()
            q4 = self.copcq4.currentText()
            q5 = self.copcq5.currentText()
            
            self.partccos.append(q1)
            self.partccos.append(q2)
            self.partccos.append(q3)
            self.partccos.append(q4)
            self.partccos.append(q5)
             
            print(self.partccos)
        elif (self.partc==6):
            q1 = self.copcq1.currentText()
            q2 = self.copcq2.currentText()
            q3 = self.copcq3.currentText()
            q4 = self.copcq4.currentText()
            q5 = self.copcq5.currentText()
            q6 = self.copcq6.currentText()
            
            self.partccos.append(q1)
            self.partccos.append(q2)
            self.partccos.append(q3)
            self.partccos.append(q4)
            self.partccos.append(q5)
            self.partccos.append(q6)
             
            print(self.partccos)
        elif (self.partc==7):
            q1 = self.copcq1.currentText()
            q2 = self.copcq2.currentText()
            q3 = self.copcq3.currentText()
            q4 = self.copcq4.currentText()
            q5 = self.copcq5.currentText()
            q6 = self.copcq6.currentText()
            q7 = self.copcq7.currentText()
            
            self.partccos.append(q1)
            self.partccos.append(q2)
            self.partccos.append(q3)
            self.partccos.append(q4)
            self.partccos.append(q5)
            self.partccos.append(q6)
            self.partccos.append(q7)
             
            print(self.partccos)
        elif (self.partc==8):
            q1 = self.copcq1.currentText()
            q2 = self.copcq2.currentText()
            q3 = self.copcq3.currentText()
            q4 = self.copcq4.currentText()
            q5 = self.copcq5.currentText()
            q6 = self.copcq6.currentText()
            q7 = self.copcq7.currentText()
            q8 = self.copcq8.currentText()
            
            self.partccos.append(q1)
            self.partccos.append(q2)
            self.partccos.append(q3)
            self.partccos.append(q4)
            self.partccos.append(q5)
            self.partccos.append(q6)
            self.partccos.append(q7)
            self.partccos.append(q8)
             
            print(self.partccos)
        elif (self.partc==9):
            q1 = self.copcq1.currentText()
            q2 = self.copcq2.currentText()
            q3 = self.copcq3.currentText()
            q4 = self.copcq4.currentText()
            q5 = self.copcq5.currentText()
            q6 = self.copcq6.currentText()
            q7 = self.copcq7.currentText()
            q8 = self.copcq8.currentText()
            q9 = self.copcq9.currentText()
            
            self.partccos.append(q1)
            self.partccos.append(q2)
            self.partccos.append(q3)
            self.partccos.append(q4)
            self.partccos.append(q5)
            self.partccos.append(q6)
            self.partccos.append(q7)
            self.partccos.append(q8)
            self.partccos.append(q9)
             
            print(self.partccos)
        elif (self.partc==10):
            q1 = self.copcq1.currentText()
            q2 = self.copcq2.currentText()
            q3 = self.copcq3.currentText()
            q4 = self.copcq4.currentText()
            q5 = self.copcq5.currentText()
            q6 = self.copcq6.currentText()
            q7 = self.copcq7.currentText()
            q8 = self.copcq8.currentText()
            q9 = self.copcq9.currentText()
            q10 = self.copcq10.currentText()
            
            self.partccos.append(q1)
            self.partccos.append(q2)
            self.partccos.append(q3)
            self.partccos.append(q4)
            self.partccos.append(q5)
            self.partccos.append(q6)
            self.partccos.append(q7)
            self.partccos.append(q8)
            self.partccos.append(q9)
            self.partccos.append(q10)
             
            print(self.partccos)
        global partccos
        partccos=self.partccos

    def switch(self):
        #print("signal emitted.....>>>")
        self.coseeb()
        self.cosee()
        self.coseec()
        self.spinbox_toggled2()

        self.spinbox_toggledc2()
        self.switch_window.emit()


class Login(QtWidgets.QWidget):

    switch_window = QtCore.pyqtSignal()
    switch_window1 = QtCore.pyqtSignal()

    def __init__(self):
        QtWidgets.QWidget.__init__(self)
        self.setWindowTitle('Mark Easy(20%)')
        #self.resize(1300, 700)


        layout = QtWidgets.QGridLayout()

        self.button = QtWidgets.QPushButton('create new co')
        self.button.setFixedWidth(100)

        self.button.clicked.connect(self.login)
        self.button.setGeometry(QtCore.QRect(140, 240, 281, 31))
        self.button1 = QtWidgets.QPushButton('Administator')
        self.button1.setFixedWidth(100)

        self.button1.clicked.connect(self.admin)
        self.button1.setGeometry(QtCore.QRect(140, 240, 281, 31))
        layout.addWidget(self.button,4,4)

        layout.addWidget(self.button1,5,4)

        self.setLayout(layout)

    def login(self):
        self.switch_window.emit()
        
    def admin(self, fileName):
        self.switch_window1.emit()
        '''fileName, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Open excel",
               (QtCore.QDir.homePath()), "xls (*.xls )")''' 
class WindowTwo(QWidget):
    def __init__(self):
        QWidget.__init__(self)
        self.setWindowTitle("Mark Easy(100%)")
        layout = QGridLayout()
        self.setLayout(layout)

        dd1 = "Percentage of students obtainted more than 50% marks in relevant CO"
        self.com=QLabel(dd1)
        self.com.setFont(QFont("Times New Roman",10))
        self.com.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com,2,0)

        dd2 = "No. of students obtained more than 50% marks in relevant CO "
        self.com1=QLabel(dd2)
        self.com1.setFont(QFont("Times New Roman",10))
        self.com1.setAlignment(Qt.AlignRight)

        layout.addWidget(self.com1,1,0)
        dd3 = "Correlation level achieved "
        com2=QLabel(dd3)
        com2.setAlignment(Qt.AlignRight)
        com2.setFont(QFont("Times New Roman",10))
        layout.addWidget(com2,3,0)
        self.com01=QLabel(" ")
        self.com01.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com01,0,1)
        self.com02=QLabel(" ")
        self.com02.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com02,0,2)
        self.com03=QLabel(" ")
        self.com03.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com03,0,3)
        self.com04=QLabel(" ")
        self.com04.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com04,0,4)
        self.com05=QLabel(" ")
        self.com05.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com05,0,5)

        self.com11=QLabel(" ")
        self.com11.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com11,1,1)
        self.com12=QLabel(" ")
        self.com12.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com12,2,1)
        self.com13=QLabel(" ")
        self.com13.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com13,3,1)

        self.com21=QLabel(" ")
        self.com21.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com21,1,2)
        self.com22=QLabel(" ")
        self.com22.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com22,2,2)
        self.com23=QLabel(" ")
        self.com23.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com23,3,2)

        self.com31=QLabel(" ")
        self.com31.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com31,1,3)
        self.com32=QLabel(" ")
        self.com32.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com32,2,3)
        self.com33=QLabel(" ")
        self.com33.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com33,3,3)

        self.com41=QLabel(" ")
        self.com41.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com41,1,4)
        self.com42=QLabel(" ")
        self.com42.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com42,2,4)
        self.com43=QLabel(" ")
        self.com43.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com43,3,4)

        self.com51=QLabel(" ")
        self.com51.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com51,1,5)
        self.com52=QLabel(" ")
        self.com52.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com52,2,5)
        self.com53=QLabel(" ")
        self.com53.setAlignment(Qt.AlignRight)
        layout.addWidget(self.com53,3,5)
        #self.toolbutton1.clicked.connect(self.butclick)
        self.cop= QComboBox()
        #self.copcq6.addItem()
        self.cop.addItem("50%")
        self.cop.addItem("45%")
        self.cop.addItem("40%")
        self.cop.addItem("35%")
        self.cop.addItem("30%")
        self.cop.addItem("25%")
        self.cop.addItem("20%")
        self.cop.addItem("15%")
        self.cop.addItem("10%")
        self.cop.addItem("5%")
        self.cop.addItem("55%")
        self.cop.addItem("60%")
        self.cop.addItem("65%")
        self.cop.addItem("70%")
        self.cop.addItem("75%")
        self.cop.addItem("80%")
        self.cop.addItem("85%")
        self.cop.addItem("90%")
        self.cop.addItem("95%")
        self.cop.addItem("100%")

        layout.addWidget(self.cop,4,0)
        self.cop.currentTextChanged.connect(self.persee)
    
        
        self.tablewidget2 = QTableWidget()
        
        self.parta=partasp

        self.partb=partbsp
        self.partc=partcsp

        self.partbsuba=partbsub
        if (self.partb>0):
            self.val=0
            for i in range(len(self.partbsuba)):
                add=(self.partbsuba[i])
                partb=self.val+add
                self.val=partb
                #print(self.val,self.partbsuba)
        else:
            self.val=0
        self.partcsuba=partcsub
        #print(self.partcsuba)
        if (self.partc>0):
            valc=0
            for i in range(len(self.partcsuba)):
                add=(self.partcsuba[i])
                partb=valc+add
                valc=partb
                #print(valc,self.partcsuba,)
        else:
            valc=0
            #print(valc,"2")
        pbs=self.parta+2
        pcs=self.parta+2+self.val
        
        self.r=3+trows
        self.c=self.parta+2+self.val+valc
        #print(self.parta,self.c,"valuues")
        self.tablewidget2.setRowCount(self.r)
        self.tablewidget2.setColumnCount(self.c)
        self.tablewidget2.itemChanged.connect(self.getData)
        layout.addWidget(self.tablewidget2, 5, 0,15,15)
        self.tablewidget2.setSpan(0,2,1,self.parta)
        if (self.partb>0):
            self.tablewidget2.setSpan(0,pbs,1,self.val)
        if (self.partc>0):
            self.tablewidget2.setSpan(0,pcs,1,valc)
        #print("passed level 1")
        if (self.partb>0):
            
            val=pbs
            for i in range(len(self.partbsuba)):
                #print("passed level B")
                self.tablewidget2.setSpan(1,val,1,self.partbsuba[i])
                add=(self.partbsuba[i])
                partb=val+add
            
                val=partb
                #print(val)

                #print(val,self.partbsuba[i],"3")
        else:
            val=0

        if (self.partc>0):

            val=pcs
            #print("passed level 2")

            for i in range(len(self.partcsuba)):
                #print("passed level C")
                self.tablewidget2.setSpan(1,val,1,self.partcsuba[i])
                add=(self.partcsuba[i])
                partb=val+add
            
                val=partb
                #print(val,self.partcsuba[i],"4")
        else:
            val=0
        #print("passed level 3")

        if (self.partb>0):
            col=2+self.parta
            que=self.parta+1
            for i in range(0,self.partb):
               dr=" A "
               de=" B "
               df=str(que)
               mm=i%2
               nn=i%2
               if mm==0:
               
                   dd=" Question  "+df+dr
               else:
                
                   dd=" Question  "+df+de
               
               com=QLabel(dd)
               com.setFont(QFont("Stencil",10))
               
               self.tablewidget2.setCellWidget(1, col, com)
               if i==1 or i==3 or i==5 or i==7 or i==9 or i==11 or i==13 or i==15 or i==17 or i==19:
                  que +=1
               col +=self.partbsuba[i]
        #print("passed level 4")
               
        #print(col)
        if (self.partc>0):
            col=2+self.parta+self.val
            inter=self.partb/2
            chinter=int(inter)
            que=self.parta+chinter+1
            #print(que)
            for i in range(0,self.partc):
               dr=" A "
               de=" B "
               df=str(que)
               mm=i%2
               nn=i%2
           
               if mm==0:
               
                   dd=" Question  "+df+dr
               else:
                
                   dd=" Question  "+df+de 
               com=QLabel(dd)
               com.setFont(QFont("Stencil",10))

               self.tablewidget2.setCellWidget(1, col, com)
               if i==1 or i==3 or i==5 or i==7 or i==9 :
                  que +=1
               col +=self.partcsuba[i]
        col=2
        que=1
        for i in range(0,self.parta):
           df=str(que)
           dd=" Question  "+df 
           com=QLabel(dd)
           com.setFont(QFont("Stencil",10))

           self.tablewidget2.setCellWidget(1, col, com)
           col +=1
           que +=1
        
        self.tablewidget2.setAlternatingRowColors(True)
        self.tablewidget2.setEditTriggers(QAbstractItemView.AnyKeyPressed)
        self.tablewidget2.setSelectionMode(QAbstractItemView.ExtendedSelection)

        #self.tablewidget2.setSelectionBehavior(QAbstractItemView.SelectColumns)

        self.tablewidget2.verticalHeader().setVisible(True)

        self.tablewidget2.horizontalHeader().setVisible(True)
        self.tablewidget2.setSortingEnabled(False)

        com=QLabel("Register no")
        com.setFont(QFont("Showcard Gothic",10))
        com.setStyleSheet("color:green")

        self.tablewidget2.setCellWidget(0, 0, com)
        com=QLabel("Name")
        com.setFont(QFont("Showcard Gothic",10))
        com.setStyleSheet("color:brown")
        #print("passed level 4")

        self.tablewidget2.setCellWidget(0, 1, com)
        if (self.parta>0):
            com=QLabel("  PART A ")
            com.setFont(QFont("Arial Black ",10))
            com.setStyleSheet("color:orange")

            self.tablewidget2.setCellWidget(0, 2, com)
        if (self.partb>0):
            com=QLabel("  PART B ")

            com.setFont(QFont("Arial Black",10))

            com.setStyleSheet("color:turquoise")

            self.tablewidget2.setCellWidget(0,pbs, com)
        if (self.partc>0):
            
            co=QLabel("PART C")
            co.setFont(QFont("Arial Black",10))

            co.setStyleSheet("color:yellow")
            self.tablewidget2.setCellWidget(0, pcs, co)
        
        self.button2 =  QPushButton()
        self.button2.setText("calculate / save")
        self.button2.setFixedWidth(100)
        self.button2.setToolTip('calulate student marks and save file')
        self.button2.clicked.connect(self.calculate)
        self.tablewidget2.setCellWidget(2, 0, self.button2)
        self.tablewidget2.itemChanged.connect(self.getData)
        self.data()
        #print("passed level 3")

    def data(self):

        # To open Workbook 
        wb = xlrd.open_workbook(loc) 
        sheet = wb.sheet_by_index(0) 
        #print( sheet.name)
        #print (sheet.nrows)
        #print (sheet.ncols)
        # For row 0 and column 0

        
        global trows
        trows=sheet.nrows
        cols=sheet.ncols

        studentdata=[]
        for row in range(trows):
            datainrow=[]
            for col in range(cols):
                    text=sheet.cell_value(row,col)
                    datainrow.append((text))
            studentdata.append(datainrow)
        #print(studentdata)
        
        for row1 in range(trows):
            datainrow=studentdata[row1]
            #print(datainrow)
            for col1 in range(cols):
                text=datainrow[col1]
                if type(text) is str:
                    newitem = QTableWidgetItem(text)
                    self.tablewidget2.setItem(row1+3, col1, (newitem))
                    #print('string',text)
                else :
                    text1=int(text)
                    #print('integer',text1)
                    text2=str(text1)
                    newitem = QTableWidgetItem(text2)
                
                    #print(text,row1+3,col1)
                    self.tablewidget2.setItem(row1+3, col1, (newitem))
    
    def getData(self):
        self.rtr=self.tablewidget2.rowCount()
        rowgiven=self.rtr
        #print("line")
        data=[]
        for row in range(rowgiven):
            datainrow=[]
            for col in range(self.tablewidget2.columnCount()):
                if self.tablewidget2.item(row,col):
                    text=self.tablewidget2.item(row,col).text()
                    #print(text)
                    typ=(type(text))
                    englishletter="abcdefghijklnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ"
                    integer="1234567890"
                    empty=" "
                    if (typ==str):
                        a=text
                        if (len(text)>0):
                            string=a[0]
                            if englishletter.count(string):
                                datainrow.append((text))
                            elif integer.count(string):
                                datainrow.append(int(text))
                            else:
                                datainrow.append(0)
                        else:
                            datainrow.append(0)
                    else:
                        datainrow.append(0)
                else:
                    datainrow.append(0)
                
            data.append(datainrow)
        
        self.data=data 
        #print(self.data)
        self.tablewidget2.resizeColumnsToContents()
    def persee(self):
        print(self.cop.currentText())
        val=self.cop.currentText()
        print(type(val))
        if (val=='50%'):
            divd=2
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='45%'):
            divd=2.22
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='40%'):
            divd=2.5
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='35%'):
            divd=2.85
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='30%'):
            divd=3.33
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='25%'):
            divd=4
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='20%'):
            divd=5
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='15%'):
            divd=6.66
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='10%'):
            divd=10
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='5%'):
            divd=20
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='55%'):
            divd=1.81
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='60%'):
            divd=1.66
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='65%'):
            divd=1.533
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='70%'):
            divd=1.422
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='75%'):
            divd=1.33
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='80%'):
            divd=1.25
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='85%'):
            divd=1.176
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='90%'):
            divd=1.11
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='95%'):
            divd=1.052
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        elif (val=='100%'):
            divd=1
            self.com.setText("Percentage of students obtainted more than  "+str(val)+"  marks in relevant CO")
            self.com1.setText("No. of students obtained more than " +str(val)+ "  marks in relevant CO ")
        global divdval
        divdval=divd
    def calculate(self):
        self.persee()
        cor=self.tablewidget2.columnCount()
        if (self.c<cor):
            self.tablewidget2.setColumnCount(self.c)
        if (self.partb>0):
            pbco=partbcos
            partbsubd=partbsub
            partbcorrectco=[]
            for i in range (len(partbsubd)):
                no=partbsubd[i]
                co=pbco[i]
                for j in range(no):
                    partbcorrectco.append(co)
        else:
            partbcorrectco=[]
        if (self.partc>0):
            pcco=partccos
            partcsubd=partcsub
            partccorrectco=[]
            for i in range (len(partcsubd)):
                no=partcsubd[i]
                co=pcco[i]
                for j in range(no):
                    partccorrectco.append(co)
        else:
            partccorrectco=[]
        #print(partacos)
        #print(partccos)
        #print(self.data)
        markaget=marka
        markbget=markb
        markcget=markc
        
        paco=partacos
        pbco=partbcos
        pcco=partccos
        pccl=paco+partbcorrectco+partccorrectco
        #print(self.parta)
        #print(self.partb)
        #print(self.partc)
        
        #print(pccl)
        co1index=[]
        co2index=[]
        co3index=[]
        co4index=[]
        co5index=[]
        for i in range (len(pccl)):
            co=pccl[i]
            if (co=="CO1"):
                co1index.append(i)
            elif (co=="CO2"):
                co2index.append(i)
            elif (co=="CO3"):
                co3index.append(i)
            elif (co=="CO4"):
                co4index.append(i)
            elif (co=="CO5"):
                co5index.append(i)
        
       
        #print(co1index)
        #print(co2index)
        #print(co3index)
        #print(co4index)
        #print(co5index)
        studentdata=self.data

        comark1=[]
        for row1 in range(3,self.tablewidget2.rowCount()):
                datainrow=studentdata[row1]
                #print("2fds")
                val1=0
                for j in range(len(co1index)):
                    b1=co1index[j]
                    col1=b1+2
                    text=datainrow[col1]
                    
                    a1=text
                    add=a1+val1
                    val1=add
                comark1.append(val1)
        #print(comark1)
        comark2=[]
        for row1 in range(3,self.tablewidget2.rowCount()):
                datainrow=studentdata[row1]
                #print("2fds")
                val1=0
                for j in range(len(co2index)):
                    b1=co2index[j]
                    col1=b1+2
                    text=datainrow[col1]
                    
                    a1=text
                    add=a1+val1
                    val1=add
                comark2.append(val1)
        #print(comark2)
        comark3=[]
        for row1 in range(3,self.tablewidget2.rowCount()):
                datainrow=studentdata[row1]
                #print("2fds")
                val1=0
                for j in range(len(co3index)):
                    b1=co3index[j]
                    col1=b1+2
                    text=datainrow[col1]
                    
                    a1=text
                    add=a1+val1
                    val1=add
                comark3.append(val1)
        #print(comark3)
        comark4=[]
        for row1 in range(3,self.tablewidget2.rowCount()):
                datainrow=studentdata[row1]
                #print("2fds")
                val1=0
                for j in range(len(co4index)):
                    b1=co4index[j]
                    col1=b1+2
                    text=datainrow[col1]
                    
                    a1=text
                    add=a1+val1
                    val1=add
                comark4.append(val1)
        #print(comark4)
        comark5=[]
        for row1 in range(3,self.tablewidget2.rowCount()):
                datainrow=studentdata[row1]
                #print("2fds")
                val1=0
                for j in range(len(co5index)):
                    b1=co5index[j]
                    col1=b1+2
                    text=datainrow[col1]
                    
                    a1=text
                    add=a1+val1
                    val1=add
                comark5.append(val1)
        #print(comark5)
        #print("1dsgh")
        total=[]
        for row1 in range(3,self.tablewidget2.rowCount()):
                datainrow=studentdata[row1]
                #print("2fds")
                val=0
                for col1 in range(2,self.tablewidget2.columnCount()):
                    
                    text=datainrow[col1]
                    a1=text
                    add=a1+val
                    val=add
                total.append(val)
        #print(total)
        if (len(co1index)>0):
            column=self.tablewidget2.columnCount()
            col=column
            #print(col,column)
            self.tablewidget2.setColumnCount(column+1)
            dd=" co1 "
            com=QLabel(dd)
            com.setFont(QFont("Stencil",10))

            self.tablewidget2.setCellWidget(1, col, com)
            datainrow=comark1
            #print(datainrow)
            for row in range(len(datainrow)):
                text=datainrow[row]
                row1=row+3
                if type(text) is str:
                    newitem = QTableWidgetItem(text)
                    self.tablewidget2.setItem(row1, col, (newitem))
                    #print('string',text)
                else :
                    text1=int(text)
                    #print('integer',text1)
                    text2=str(text1)
                    newitem = QTableWidgetItem(text2)
                
                    #print(text,row1+3,col1)
                    self.tablewidget2.setItem(row1, col, (newitem))
        if (len(co2index)>0):
            column=self.tablewidget2.columnCount()
            col=column
            #print(col,column)
            self.tablewidget2.setColumnCount(column+1)
            dd=" co2 "
            com=QLabel(dd)
            com.setFont(QFont("Stencil",10))

            self.tablewidget2.setCellWidget(1, col, com)
            datainrow=comark2
            for row in range(len(datainrow)):
                text=datainrow[row]
                row1=row+3
                if type(text) is str:
                    newitem = QTableWidgetItem(text)
                    self.tablewidget2.setItem(row1, col, (newitem))
                    #print('string',text)
                else :
                    text1=int(text)
                    #print('integer',text1)
                    text2=str(text1)
                    newitem = QTableWidgetItem(text2)
                
                    #print(text,row1+3,col1)
                    self.tablewidget2.setItem(row1, col, (newitem))
        if (len(co3index)>0):
            column=self.tablewidget2.columnCount()
            col=column
            #print(col,column)
            self.tablewidget2.setColumnCount(column+1)
            dd=" co3 "
            com=QLabel(dd)
            com.setFont(QFont("Stencil",10))

            self.tablewidget2.setCellWidget(1, col, com)
            datainrow=comark3
            for row in range(len(datainrow)):
                text=datainrow[row]
                row1=row+3
                if type(text) is str:
                    newitem = QTableWidgetItem(text)
                    self.tablewidget2.setItem(row1, col, (newitem))
                    #print('string',text)
                else :
                    text1=int(text)
                    #print('integer',text1)
                    text2=str(text1)
                    newitem = QTableWidgetItem(text2)
                
                    #print(text,row1+3,col1)
                    self.tablewidget2.setItem(row1, col, (newitem))
        if (len(co4index)>0):
            column=self.tablewidget2.columnCount()
            col=column
            #print(col,column)
            self.tablewidget2.setColumnCount(column+1)
            dd=" co4 "
            com=QLabel(dd)
            com.setFont(QFont("Stencil",10))

            self.tablewidget2.setCellWidget(1, col, com)
            datainrow=comark4
            for row in range(len(datainrow)):
                text=datainrow[row]
                row1=row+3
                if type(text) is str:
                    newitem = QTableWidgetItem(text)
                    self.tablewidget2.setItem(row1, col, (newitem))
                    #print('string',text)
                else :
                    text1=int(text)
                    #print('integer',text1)
                    text2=str(text1)
                    newitem = QTableWidgetItem(text2)
                
                    #print(text,row1+3,col1)
                    self.tablewidget2.setItem(row1, col, (newitem))
        if (len(co5index)>0):
            column=self.tablewidget2.columnCount()
            col=column
            #print(col,column)
            self.tablewidget2.setColumnCount(column+1)
            dd=" co5 "
            com=QLabel(dd)
            com.setFont(QFont("Stencil",10))

            self.tablewidget2.setCellWidget(1, col, com)
            datainrow=comark5
            for row in range(len(datainrow)):
                text=datainrow[row]
                row1=row+3
                if type(text) is str:
                    newitem = QTableWidgetItem(text)
                    self.tablewidget2.setItem(row1, col, (newitem))
                    #print('string',text)
                else :
                    text1=int(text)
                    #print('integer',text1)
                    text2=str(text1)
                    newitem = QTableWidgetItem(text2)
                
                    #print(text,row1+3,col1)
                    self.tablewidget2.setItem(row1, col, (newitem))
        if (len(total)>0):
            column=self.tablewidget2.columnCount()
            col=column
            #print(col,column)
            self.tablewidget2.setColumnCount(column+1)
            dd=" Total Marks "
            com=QLabel(dd)
            com.setFont(QFont("Stencil",10))

            self.tablewidget2.setCellWidget(1, col, com)
            datainrow=total
            for row in range(len(datainrow)):
                text=datainrow[row]
                row1=row+3
                if type(text) is str:
                    newitem = QTableWidgetItem(text)
                    self.tablewidget2.setItem(row1, col, (newitem))
                    #print('string',text)
                else :
                    text1=int(text)
                    #print('integer',text1)
                    text2=str(text1)
                    newitem = QTableWidgetItem(text2)
                    #print(text,row1+3,col1)
                    self.tablewidget2.setItem(row1, col, (newitem))
        val=self.parta+2
        if (self.partb>0):
            for i in range(len(partbsub)):
                index=partbsub[i]
                add=val+index
                val=add
            partbvalb=val
        if (self.partc>0):
            val=partbvalb
            for i in range(len(partcsub)):
                index=partbsub[i]
                add=val+index
                val=add
            partcvalc=val
        studentdata=self.data
        if (self.parta>0):
            partamark=[]
            for row1 in range(3,self.tablewidget2.rowCount()):
                datainrow=studentdata[row1]
                #print("2fds")
                val1=0
                for j in range(2,self.parta):
                    text=datainrow[j]
                    a1=text
                    add=a1+val1
                    val1=add
                partamark.append(val1)
        partbmark=[]
        if (self.partb>0):
            for row1 in range(3,self.tablewidget2.rowCount()):
                datainrow=studentdata[row1]
                #print("2fds")
                val1=0
                for j in range((2+self.parta),partbvalb):
                    text=datainrow[j]
                    a1=text
                    add=a1+val1
                    val1=add
                partbmark.append(val1)
        partcmark=[]
        if (self.partc>0):
            for row1 in range(3,self.tablewidget2.rowCount()):
                datainrow=studentdata[row1]
                #print("2fds")
                val1=0
                for j in range(partbvalb,partcvalc):
                    text=datainrow[j]
                    a1=text
                    add=a1+val1
                    val1=add
                partcmark.append(val1)
        if (self.parta>0):
            column=self.tablewidget2.columnCount()
            col=column
            #print(col,column)
            self.tablewidget2.setColumnCount(column+1)
            dd=" Part A Marks "
            com=QLabel(dd)
            com.setFont(QFont("Stencil",10))

            self.tablewidget2.setCellWidget(1, col, com)
            datainrow=partamark
            for row in range(len(datainrow)):
                text=datainrow[row]
                row1=row+3
                if type(text) is str:
                    newitem = QTableWidgetItem(text)
                    self.tablewidget2.setItem(row1, col, (newitem))
                    #print('string',text)
                else :
                    text1=int(text)
                    #print('integer',text1)
                    text2=str(text1)
                    newitem = QTableWidgetItem(text2)
                    #print(text,row1+3,col1)
                    self.tablewidget2.setItem(row1, col, (newitem))
        if (self.partb>0):
            column=self.tablewidget2.columnCount()
            col=column
            #print(col,column)
            self.tablewidget2.setColumnCount(column+1)
            dd=" Part B Marks "
            com=QLabel(dd)
            com.setFont(QFont("Stencil",10))

            self.tablewidget2.setCellWidget(1, col, com)
            datainrow=partbmark
            for row in range(len(datainrow)):
                text=datainrow[row]
                row1=row+3
                if type(text) is str:
                    newitem = QTableWidgetItem(text)
                    self.tablewidget2.setItem(row1, col, (newitem))
                    #print('string',text)
                else :
                    text1=int(text)
                    #print('integer',text1)
                    text2=str(text1)
                    newitem = QTableWidgetItem(text2)
                    #print(text,row1+3,col1)
                    self.tablewidget2.setItem(row1, col, (newitem))
        if (self.partc>0):
            column=self.tablewidget2.columnCount()
            col=column
            #print(col,column)
            self.tablewidget2.setColumnCount(column+1)
            dd=" Part C Marks "
            com=QLabel(dd)
            com.setFont(QFont("Stencil",10))

            self.tablewidget2.setCellWidget(1, col, com)
            datainrow=partcmark
            for row in range(len(datainrow)):
                text=datainrow[row]
                row1=row+3
                if type(text) is str:
                    newitem = QTableWidgetItem(text)
                    self.tablewidget2.setItem(row1, col, (newitem))
                    #print('string',text)
                else :
                    text1=int(text)
                    #print('integer',text1)
                    text2=str(text1)
                    newitem = QTableWidgetItem(text2)
                    #print(text,row1+3,col1)
                    self.tablewidget2.setItem(row1, col, (newitem))
        markamulti=len(paco)
        markbmulti=len(pbco)
        markcmulti=len(pcco)
        val1=0
        if (len(paco)>0):
            for i in range(markamulti):
                add=marka+val1
                val1=add
        val2=0
        if (len(pbco)>0):
            for i in range(markbmulti):
                add=markb+val2
                val2=add
        val3=0
        if (len(pcco)>0):
            for i in range(markcmulti):
                add=markc+val3
                val3=add
        tot=val1+val2+val3
        #print("dsgfghjdf",tot)
        global totl
        totl=tot
        cotota1=[]
        cotota2=[]
        cotota3=[]
        cotota4=[]
        cotota5=[]
        if (len(paco)>0):
            for i in range (len(paco)):
                co=paco[i]
                if (co=="CO1"):
                    cotota1.append(i)
                elif (co=="CO2"):
                    cotota2.append(i)
                elif (co=="CO3"):
                    cotota3.append(i)
                elif (co=="CO4"):
                    cotota4.append(i)
                elif (co=="CO5"):
                    cotota5.append(i)
        coa1=0
        for i in range(len(cotota1)):
            vala=marka
            mark=vala+coa1
            coa1=mark
        coa2=0
        for i in range(len(cotota2)):
            vala=marka
            mark=vala+coa2
            coa2=mark
        coa3=0
        for i in range(len(cotota3)):
            vala=marka
            mark=vala+coa3
            coa3=mark
        coa4=0
        for i in range(len(cotota4)):
            vala=marka
            mark=vala+coa4
            coa4=mark
        coa5=0
        for i in range(len(cotota5)):
            vala=marka
            mark=vala+coa5
            coa5=mark
            
        cototb1=[]
        cototb2=[]
        cototb3=[]
        cototb4=[]
        cototb5=[]
        if (len(pbco)>0):
            for i in range (len(pbco)):
                co=pbco[i]
                if (co=="CO1"):
                    cototb1.append(i)
                elif (co=="CO2"):
                    cototb2.append(i)
                elif (co=="CO3"):
                    cototb3.append(i)
                elif (co=="CO4"):
                    cototb4.append(i)
                elif (co=="CO5"):
                    cototb5.append(i)
        cob1=0
        for i in range(len(cototb1)):
            vala=markb
            mark=vala+cob1
            cob1=mark
        cob2=0
        for i in range(len(cototb2)):
            vala=markb
            mark=vala+cob2
            cob2=mark
        cob3=0
        for i in range(len(cototb3)):
            vala=markb
            mark=vala+cob3
            cob3=mark
        cob4=0
        for i in range(len(cototb4)):
            vala=markb
            mark=vala+cob4
            cob4=mark
        cob5=0
        for i in range(len(cototb5)):
            vala=markb
            mark=vala+cob5
            cob5=mark
        
        cototc1=[]
        cototc2=[]
        cototc3=[]
        cototc4=[]
        cototc5=[]
        if (len(pcco)>0):
            for i in range (len(pcco)):
                co=pcco[i]
                if (co=="CO1"):
                    cototc1.append(i)
                elif (co=="CO2"):
                    cototc2.append(i)
                elif (co=="CO3"):
                    cototc3.append(i)
                elif (co=="CO4"):
                    cototc4.append(i)
                elif (co=="CO5"):
                    cototc5.append(i)
        coc1=0
        for i in range(len(cototc1)):
            vala=markc
            mark=vala+coc1
            coc1=mark
        coc2=0
        for i in range(len(cototc2)):
            vala=markc
            mark=vala+coc2
            coc2=mark
        coc3=0
        for i in range(len(cototc3)):
            vala=markc
            mark=vala+coc3
            coc3=mark
        coc4=0
        for i in range(len(cototc4)):
            vala=markc
            mark=vala+coc4
            coc4=mark
        coc5=0
        for i in range(len(cototc5)):
            vala=markc
            mark=vala+coc5
            coc5=mark
        
        cot1=coa1+cob1+coc1
        cot2=coa2+cob2+coc2
        cot3=coa3+cob3+coc3
        cot4=coa4+cob4+coc4
        cot5=coa5+cob5+coc5
        #print(cot1,cot2,cot3,cot4,cot5)
        if (len(co1index)>0):
            co501=int(cot1/divdval)
            #print(co501)
            datainrow=comark1
            #print(datainrow)
            count1=0
            for row in range(len(datainrow)):
                text=datainrow[row]
                if (text>=co501):
                    add=count1+1
                    count1=add
                #print(count1)
        if (len(co2index)>0):
            co502=int(cot2/divdval)
            #print(co502)
            datainrow=comark2
            #print(datainrow)
            count2=0
            for row in range(len(datainrow)):
                text=datainrow[row]
                if (text>=co502):
                    add=count2+1
                    count2=add
                #print(count2)

        if (len(co3index)>0):
            co503=(cot3/divdval)
            datainrow=comark3
            #print(datainrow)
            count3=0
            for row in range(len(datainrow)):
                text=datainrow[row]
                if (text>=co503):
                    add=count3+1
                    count3=add
        if (len(co4index)>0):
            co504=int(cot4/divdval)
            datainrow=comark4
            #print(datainrow)
            count4=0
            for row in range(len(datainrow)):
                text=datainrow[row]
                if (text>=co504):
                    add=count4+1
                    count4=add
        if (len(co5index)>0):
            co505=int(cot5/divdval)
            datainrow=comark5
            #print(datainrow)
            count5=0
            for row in range(len(datainrow)):
                text=datainrow[row]
                if (text>=co505):
                    add=count5+1
                    count5=add
        if (len(co1index)>0):
            corcount1=(count1/trows)*100
            percount1=(corcount1)
        if (len(co2index)>0):
            corcount2=(count2/trows)*100
            percount2=(corcount2)
        if (len(co3index)>0):
            corcount3=(count3/trows)*100
            percount3=(corcount3)
        if (len(co4index)>0):
            corcount4=(count4/trows)*100
            percount4=(corcount4)
        if (len(co5index)>0):
            corcount5=(count5/trows)*100
            percount5=(corcount5)
        if (len(co1index)>0):
            if (percount1>=60):
                atlco1=1
            elif(percount1>=70):
                atlco1=2
            elif(percount1>=80):
                atlco1=3
            else:
                atlco1=0
        if (len(co2index)>0):
            if (percount2>=60):
                atlco2=1
            elif(percount2>=70):
                atlco2=2
            elif(percount2>=80):
                atlco2=3
            else:
                atlco2=0
        if (len(co3index)>0):
            if (percount3>=60):
                atlco3=1
            elif(percount3>=70):
                atlco3=2
            elif(percount3>=80):
                atlco3=3
            else:
                atlco3=0
        if (len(co4index)>0):
            if (percount4>=60):
                atlco4=1
            elif(percount4>=70):
                atlco4=2
            elif(percount4>=80):
                atlco4=3
            else:
                atlco4=0
        if (len(co5index)>0):
            if (percount5>=60):
                atlco5=1
            elif(percount5>=70):
                atlco5=2
            elif(percount5>=80):
                atlco5=3
            else:
                atlco5=0

        currentco=[]
        if (len(co1index)>0):
            currentco.append("co1")
        if (len(co2index)>0):
            currentco.append("co2")
        if (len(co3index)>0):
            currentco.append("co3")
        if (len(co4index)>0):
            currentco.append("co4")
        if (len(co5index)>0):
            currentco.append("co5")
        #print(currentco)
        #print("yes level achieved")
        self.rtr=self.tablewidget2.rowCount()
        rowgiven=self.rtr
        data=[]
        for row in range(rowgiven):
            datainrow=[]
            for col in range(self.tablewidget2.columnCount()):
                if self.tablewidget2.item(row,col):
                    text=self.tablewidget2.item(row,col).text()
                    #print(text)
                    typ=(type(text))
                    englishletter="abcdefghijklnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ"
                    integer="1234567890"
                    empty=" "
                    if (typ==str):
                        a=text
                        if (len(text)>0):
                            string=a[0]
                            if englishletter.count(string):
                                datainrow.append((text))
                            elif integer.count(string):
                                datainrow.append(int(text))
                            else:
                                datainrow.append(0)
                        else:
                            datainrow.append(0)
                    else:
                        datainrow.append(0)
                else:
                    datainrow.append(0)
                
            data.append(datainrow)
        
        studentdata1=data
        wdataindex=[0,1]
        dcolumn=self.c
        addcolumn=self.tablewidget2.columnCount()
        value=addcolumn-dcolumn
        atn=0
        #print("value",value,"dcolumn",dcolumn,"addcolumn",addcolumn)
        for i in range(value):
            ind=dcolumn+atn
            wdataindex.append(ind)
            dcolumn=ind
            atn=1
        #print(wdataindex)
        worddata=[]
        for row1 in range(3,self.tablewidget2.rowCount()):
                datainrow=studentdata1[row1]
                #print("2fds")
                windata=[]
                for j in range(len(wdataindex)):
                    #print(j)
                    #print(len(datainrow))
                    col=wdataindex[j]
                    #print(col)
                    text=datainrow[col]
                    #print(text)
                    windata.append(str(text))
                worddata.append(windata)
        #print(worddata)            
        name=staffname
        desig=designation
        subname=subjectname
        subcode=subjectcode
        yearof=year
        branch=department
        sec=section
        document = Document()
        document.add_heading('CANDIDATES RESULT',0)
        p = document.add_paragraph('Cycle Test Results for ')
        p.add_run(branch).bold = True
        p.add_run(str(yearof)).bold=True
        p.add_run(' and  ')
        p.add_run(str(sec)+'  section.').italic = True
        document.add_heading('Staff Name and Designation',level=1)

        document.add_heading(name,level=4)

        document.add_paragraph(desig).bold=True
        document.add_heading('Class Details',level=1)
        document.add_paragraph('Department  : ' +str(branch)).bold=True
        document.add_paragraph('Year : ' +str(yearof)).bold=True
        document.add_paragraph('Section : ' +str(sec)).bold=True
        document.add_paragraph('Subject Name :' +str(subname)).bold=True
        document.add_paragraph('Subject Code : '+ str(subcode)).bold=True
        document.add_heading('Students Details',level=1)
        document.add_paragraph('No. of Students : '+str(trows)).bold=True
        document.add_paragraph('No. of Absents : ').bold=True
        document.add_paragraph('No. of Presents : ').bold=True
        document.add_heading('Question Paper Pattern',level=1)
        if (self.parta>0):
            document.add_paragraph('No of  Part A Questions : '+str(self.parta),style='List Bullet').bold=True
            if (len(cotota1)>0):
                document.add_paragraph('Questions from Co1: '+str(len(cotota1))+'      Marks:'+str(coa1),style='List Number').bold=True
            if (len(cotota2)>0):
                document.add_paragraph('Questions from Co2: '+str(len(cotota2))+'      Marks:'+str(coa2),style='List Number').bold=True
            if (len(cotota3)>0):
                document.add_paragraph('Questions from Co3: '+str(len(cotota3))+'      Marks:'+str(coa3),style='List Number').bold=True
            if (len(cotota4)>0):
                document.add_paragraph('Questions from Co4: '+str(len(cotota4))+'      Marks:'+str(coa4),style='List Number').bold=True
            if (len(cotota5)>0):
                document.add_paragraph('Questions from Co5: '+str(len(cotota5))+'      Marks:'+str(coa5),style='List Number').bold=True
        if (self.partb>0):
            #document.add_paragraph('From Co2',style='List Number').bold=True
            document.add_paragraph('No of  Part B Questions : '+str(int(self.partb/2))+'    Either Or type Questions'+'   Total Questions:  '+str(self.partb),style='List Bullet').bold=True
            if (len(cototb1)>0):
                document.add_paragraph('Questions from Co1: '+str(len(cototb1))+'      Marks:'+str(cob1),style='List Number').bold=True
            if (len(cototb2)>0):
                document.add_paragraph('Questions from Co2: '+str(len(cototb2))+'      Marks:'+str(cob2),style='List Number').bold=True
            if (len(cototb3)>0):
                document.add_paragraph('Questions from Co3: '+str(len(cototb3))+'      Marks:'+str(cob3),style='List Number').bold=True
            if (len(cototb4)>0):
                document.add_paragraph('Questions from Co4: '+str(len(cototb4))+'      Marks:'+str(cob4),style='List Number').bold=True
            if (len(cototb5)>0):
                document.add_paragraph('Questions from Co5: '+str(len(cototb5))+'      Marks:'+str(cob5),style='List Number').bold=True
        #document.add_paragraph('From Co1',style='List Number').bold=True
        #document.add_paragraph('From Co2',style='List Number').bold=True
        if (self.partc>0):
            document.add_paragraph('No of  Part C Questions : '+str(int(self.partc/2))+'     Either Or type Questions'+'   Total Questions:  '+str(self.partc),style='List Bullet').bold=True
            #document.add_paragraph('Total Marks'+str(int(totl)),style='List Bullet').bold=True
            if (len(cototc1)>0):
                document.add_paragraph('Questions from Co1: '+str(len(cototc1))+'      Marks:'+str(coc1),style='List Number').bold=True
            if (len(cototc2)>0):
                document.add_paragraph('Questions from Co2: '+str(len(cototc2))+'      Marks:'+str(coc2),style='List Number').bold=True
            if (len(cototc3)>0):
                document.add_paragraph('Questions from Co3: '+str(len(cototc3))+'      Marks:'+str(coc3),style='List Number').bold=True
            if (len(cototc4)>0):
                document.add_paragraph('Questions from Co4: '+str(len(cototc4))+'      Marks:'+str(coc4),style='List Number').bold=True
            if (len(cototc5)>0):
                document.add_paragraph('Questions from Co5: '+str(len(cototc5))+'      Marks:'+str(coc5),style='List Number').bold=True

        #document.add_paragraph('From Co1',style='List Number').bold=True
        #document.add_paragraph('From Co2',style='List Number').bold=True
              
        num1=len(currentco)
        num=currentco
        #print(num,num1)
        for i in range(len(currentco)):
            
            #print(i,num[i])
            if (i==0):
                if (num[0]=="co1"):
                    text="CO1"
                    self.com01.setText(" "+str(text))
                    self.com11.setText(" "+str(count1))
                    self.com12.setText(" "+str(percount1)+"%")
                    self.com13.setText(" "+str(atlco1))
                    
                    #print("2")
                elif (num[0]=="co2"):
                    text="CO2"
                    self.com01.setText(" "+str(text))
                    self.com11.setText(" "+str(count2))
                    self.com12.setText(" "+str(percount2)+"%")
                    self.com13.setText(" "+str(atlco2))
                    #print("3")
                elif (num[0]=="co3"):
                    text="CO3"
                    self.com01.setText(" "+str(text))
                    self.com11.setText(" "+str(count3))
                    self.com12.setText(" "+str(percount3)+"%")
                    self.com13.setText(" "+str(atlco3))
                    #print("4")
                elif (num[0]=="co4"):
                    text="CO4"
                    self.com01.setText(" "+str(text))
                    self.com11.setText(" "+str(count4))
                    self.com12.setText(" "+str(percount4)+"%")
                    self.com13.setText(" "+str(atlco4))
                    #print("5")
                elif (num[0]=="co5"):
                    text="CO5"
                    self.com01.setText(" "+str(text))
                    self.com11.setText(" "+str(count5))
                    self.com12.setText(" "+str(percount5)+"%")
                    self.com13.setText(" "+str(atlco5))
                    #print("6")
            if (i==1):
                if (num[1]=="co2"):
                    text="CO2"
                    self.com02.setText(" "+str(text))
                    self.com21.setText(" "+str(count2))
                    self.com22.setText(" "+str(percount2)+"%")
                    self.com23.setText(" "+str(atlco2))
                    #print("7")
                elif (num[1]=="co3"):
                    text="CO3"
                    self.com02.setText(" "+str(text))
                    self.com21.setText(" "+str(count3))
                    self.com22.setText(" "+str(percount3)+"%")
                    self.com23.setText(" "+str(atlco3))
                    #print("8")
                elif (num[1]=="co4"):
                    text="CO4"
                    self.com02.setText(" "+str(text))
                    self.com21.setText(" "+str(count4))
                    self.com22.setText(" "+str(percount4)+"%")
                    self.com23.setText(" "+str(atlco4))
                    #print("9")
                elif (num[1]=="co5"):
                    text="CO5"
                    self.com02.setText(" "+str(text))
                    self.com21.setText(" "+str(count5))
                    self.com22.setText(" "+str(percount5)+"%")
                    self.com23.setText(" "+str(atlco5))
                    #print("10")
            if (i==2):
                if (num[2]=="co3"):
                    text="CO3"
                    self.com03.setText(" "+str(text))
                    self.com31.setText(" "+str(count3))
                    self.com32.setText(" "+str(percount3)+"%")
                    self.com33.setText(" "+str(atlco3))
                    #print("11")
                elif (num[2]=="co4"):
                    text="CO4"
                    self.com03.setText(" "+str(text))
                    self.com31.setText(" "+str(count4))
                    self.com32.setText(" "+str(percount4)+"%")
                    self.com33.setText(" "+str(atlco4))
                    #print("12")
                elif (num[2]=="co5"):
                    text="CO5"
                    self.com03.setText(" "+str(text))
                    self.com31.setText(" "+str(count5))
                    self.com32.setText(" "+str(percount5)+"%")
                    self.com33.setText(" "+str(atlco5))
                    #print("13")
            if (i==3):
                if (num[3]=="co4"):
                    text="CO4"
                    self.com04.setText(" "+str(text))
                    self.com41.setText(" "+str(count4))
                    self.com42.setText(" "+str(percount4)+"%")
                    self.com43.setText(" "+str(atlco4))
                    #print("14")
                elif (num[3]=="co5"):
                    text="CO5"
                    self.com04.setText(" "+str(text))
                    self.com41.setText(" "+str(count5))
                    self.com42.setText(" "+str(percount5)+"%")
                    self.com43.setText(" "+str(atlco5))
                    #print("15")
            if (i==4):
                if (num[4]=="co5"):
                    text="CO5"
                    self.com05.setText(" "+str(text))
                    self.com51.setText(" "+str(count5))
                    self.com52.setText(" "+str(percount5)+"%")
                    self.com53.setText(" "+str(atlco5))
        #print("proceed")
        #document.add_paragraph('From Co2',style='List Number').bold=True


        #document.add_paragraph('Intense quote',style='Title')
        #document.add_picture('monty-truth.png',width=Inches(1.25))
        # get table data -------------
        headcells=['REGISTER NO','NAME']
        data=[]
        studentdata=self.data
        if (len(co1index)>0):
            headcells.append('CO1 MARKS')
        #print("proceed")
        if (len(co2index)>0):
            headcells.append('CO2 MARKS')
        #print("proceed")
        if (len(co3index)>0):
            headcells.append('CO3 MARKS')
        #print("proceed")
        if (len(co4index)>0):
            headcells.append('CO4 MARKS')
        #print("proceed")
        if (len(co5index)>0):
            headcells.append('CO5 MARKS')

        headcells.append('TOTAL MARKS')
        #print("proeed3")
        if (self.parta>0):
            headcells.append('PART A MARKS')
        #print("proceed")
        if (self.partb>0):
            headcells.append('PART B MARKS')
        #print("proceed")
        if (self.partc>0):
            headcells.append('PART C MARKS')
        windata=worddata[0]
        table = document.add_table(rows=len(worddata)+1,cols=len(windata))
        # populate header row --------

        # add a data row for each item
        for i in range(len(worddata)):
            item=worddata[i]
            for j in range(len(item)):
                #print(item[j]) 
                cell = table.cell(i+1, j)
                cell.text=item[j]
        heading_cells = table.rows[0].cells
        for i in range(len(headcells)):
            heading_cells[i].text = headcells[i]
        val=self.cop.currentText()
        document.add_heading('No. of students obtained more than ' +str(val)+ '  marks in relevant CO ',level=1)
        for i in range(len(currentco)):
            if (currentco[i]=="co1"):
                document.add_paragraph('In Co1  :'+str(count1),style='List Bullet').bold=True
            if (currentco[i]=="co2"):
                document.add_paragraph('In Co2  :'+str(count2),style='List Bullet').bold=True
            if (currentco[i]=="co3"):
                document.add_paragraph('In Co3  :'+str(count3),style='List Bullet').bold=True
            if (currentco[i]=="co4"):
                document.add_paragraph('In Co4  :'+str(count4),style='List Bullet').bold=True
            if (currentco[i]=="co5"):
                document.add_paragraph('In Co5  :'+str(count5),style='List Bullet').bold=True
        document.add_heading('Percentage of students obtainted more than  '+str(val)+'  marks in relevant CO',level=1)
        for i in range(len(currentco)):
            if (currentco[i]=="co1"):
                document.add_paragraph('In Co1  :'+str(percount1)+'%',style='List Bullet').bold=True
            if (currentco[i]=="co2"):
                document.add_paragraph('In Co2  :'+str(percount2)+'%',style='List Bullet').bold=True
            if (currentco[i]=="co3"):
                document.add_paragraph('In Co3  :'+str(percount3)+'%',style='List Bullet').bold=True
            if (currentco[i]=="co4"):
                document.add_paragraph('In Co4  :'+str(percount4)+'%',style='List Bullet').bold=True
            if (currentco[i]=="co5"):
                document.add_paragraph('In Co5  :'+str(percount5)+'%',style='List Bullet').bold=True
        document.add_heading('Correlation level achieved ',level=1)
        for i in range(len(currentco)):
            if (currentco[i]=="co1"):
                document.add_paragraph('In Co1  :'+str(atlco1),style='List Bullet').bold=True
            if (currentco[i]=="co2"):
                document.add_paragraph('In Co2  :'+str(atlco2),style='List Bullet').bold=True
            if (currentco[i]=="co3"):
                document.add_paragraph('In Co3  :'+str(atlco3),style='List Bullet').bold=True
            if (currentco[i]=="co4"):
                document.add_paragraph('In Co4  :'+str(atlco4),style='List Bullet').bold=True
            if (currentco[i]=="co5"):
                document.add_paragraph('In Co5  :'+str(atlco5),style='List Bullet').bold=True

        # add a data row for each item
        table.style = 'Table Grid'
        self.filename = QFileDialog.getSaveFileName(self, 'Save This File  '+str(staffname),filter='Word Files (*.docx)')
        if self.filename[0]:
            #print(self.filename[0])
            result=self.filename[0]
        if result:
            document.save(result)                    

            
    def keyPressEvent(self, ev):
        
        if (ev.key() == Qt.Key_V) and (ev.modifiers() & Qt.ControlModifier): 
            self.pasteSelection()
        if (ev.key() == Qt.Key_C) and (ev.modifiers() & Qt.ControlModifier): 
            self.copySelection()
    def copySelection(self):
        selection = self.tablewidget2.selectedIndexes()
        data=self.data
        if selection:
            rows = sorted(index.row() for index in selection)
            columns = sorted(index.column() for index in selection)
            rowcount = rows[-1] - rows[0] + 1
            colcount = columns[-1] - columns[0] + 1
            table = [[''] * colcount for _ in range(rowcount)]
            for index in selection:
                row = index.row() - rows[0]
                column = index.column() - columns[0]
                table[row][column] = index.data()
                
            stream = io.StringIO()
            csv.writer(stream, delimiter='\t').writerows(table)
            QApplication.clipboard().setText(stream.getvalue())
        
    def pasteSelection(self):
        selection = self.tablewidget2.selectedIndexes()
        if selection:
            model = self.tablewidget2.model()
            buffer = QApplication.clipboard().text()
            print(buffer)

            rows = sorted(index.row() for index in selection)
            columns = sorted(index.column() for index in selection)
            
            reader = csv.reader(io.StringIO(buffer), delimiter='\t')
            if len(rows) == 1 and len(columns) == 1:
                for i, line in enumerate(reader):
                    for j, cell in enumerate(line):
                        model.setData(model.index(rows[0]+i,columns[0]+j), cell)
            else:
                arr = [ [ cell for cell in row ] for row in reader]
                #print(arr)
                for index in selection:
                    row = index.row() - rows[0]
                    column = index.column() - columns[0]
                    model.setData(model.index(index.row(), index.column()), arr[row][column])
    
    
            
class adminWindow(QtWidgets.QWidget):
    def __init__(self):
        QWidget.__init__(self)
        self.setWindowTitle("HELLO")
        layout = QGridLayout()
        self.setLayout(layout)
        self.tablewidget = QTableWidget()
        self.tablewidget.setRowCount(1)
        self.tablewidget.setColumnCount(2)
        self.tablewidget.setHorizontalHeaderLabels(["Register NO", "Name"])
        self.tablewidget.setGeometry(10, 50, 780, 645)
        self.tablewidget.itemChanged.connect(self.getData)

        self.tablewidget.setColumnWidth(3,4)
        self.tablewidget.setRowHeight(2,3)
        #self.tablewidget.gridStyle(Qt.PenStyle)
        #self.tablewidget.setViewport()

        #self.tablewidget.isIndexHidden(True)
        #self.tablewidget.setFrameStyle(Raised)
        #self.tablewidget.frameWidth(50)
        layout.addWidget(self.tablewidget, 5, 0,4,15)
        menubar = QMenuBar()
        layout.addWidget(menubar, 0, 0)

        menu = menubar.addMenu("&File")
        open_action = menu.addAction("&Open")
        def open_file():
            #print("work")
            fileName, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Open Excel",
                   (QtCore.QDir.homePath()), "Excel (*.xlsx *.xls)")
            if fileName:
                admin_file = fileName 

            # To open Workbook 
            wb = xlrd.open_workbook(admin_file) 
            sheet = wb.sheet_by_index(0) 
            print( sheet.name)
            print (sheet.nrows)
            print (sheet.ncols)
            # For row 0 and column 0

            
            trows=sheet.nrows
            cols=sheet.ncols
            self.tablewidget.setRowCount(trows)

            studentdata=[]
            for row in range(trows):
                datainrow=[]
                for col in range(cols):
                        text=sheet.cell_value(row,col)
                        datainrow.append((text))
                studentdata.append(datainrow)
            #print(studentdata)
            
            for row1 in range(trows):
                datainrow=studentdata[row1]
                #print(datainrow)
                for col1 in range(cols):
                    text=datainrow[col1]
                    if type(text) is str:
                        newitem = QTableWidgetItem(text)
                        self.tablewidget.setItem(row1, col1, (newitem))
                        #print('string',text)
                    else :
                        text1=int(text)
                        #print('integer',text1)
                        text2=str(text1)
                        newitem = QTableWidgetItem(text2)
                    
                        #print(text,row1+3,col1)
                        self.tablewidget.setItem(row1, col1, (newitem))
        open_action.triggered.connect(open_file)
        open_action.setShortcut(QKeySequence.Open)
        #menu.addAction(open_action)
        save_action = menu.addAction("save")
        def save_file():
            self.getData()
            self.filename = QFileDialog.getSaveFileName(self, 'Save File',filter='Excel Files (*.xls)')
            if self.filename[0]:
                #print(self.filename[0])
                result=self.filename[0]
            if result:
                wb = Workbook() 

                # add_sheet is used to create sheet. 
                sheet1 = wb.add_sheet('Sheet 1')
               
                row=self.tablewidget.rowCount()
                col=self.tablewidget.columnCount()
                print(row,col)
                for row1 in range(row):
                    datainrow=tabledata[row1]
                    for col1 in range(col):
                        text=datainrow[col1]
                        print(text)
                        if col1==0:
                            text1=int(text)
                            sheet1.write(row1, col1, text1) 
                        else:
                            sheet1.write(row1, col1, text) 
                        
                
                wb.save(result) 
            
        save_action.triggered.connect(save_file)
        save_action.setShortcut(QKeySequence.Save)
        self.label=QLabel("year :")
        layout.addWidget(self.label,1,1)
        self.lineedit = QLineEdit()
        #self.lineedit.returnPressed.connect(self.return_pressed)
        layout.addWidget(self.lineedit, 1, 2)
        self.label=QLabel("Department :")
        layout.addWidget(self.label,2,1)
        self.lineedit = QLineEdit()
        #self.lineedit.returnPressed.connect(self.return_pressed)
        layout.addWidget(self.lineedit, 2, 2)
        self.label=QLabel("Section :")
        layout.addWidget(self.label,3,1)
        self.lineedit = QLineEdit()
        #self.lineedit.returnPressed.connect(self.return_pressed)
        layout.addWidget(self.lineedit, 3, 2)
        
    def openfile(self):
        fileName, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Open Excel",
               (QtCore.QDir.homePath()), "Excel (*.xlsx *.xls)")
        if fileName:
            admin_file = fileName 

        # To open Workbook 
        wb = xlrd.open_workbook(admin_file) 
        sheet = wb.sheet_by_index(0) 
        print( sheet.name)
        print (sheet.nrows)
        print (sheet.ncols)
        # For row 0 and column 0

        
        global trows
        trows=sheet.nrows
        cols=sheet.ncols

        studentdata=[]
        for row in range(trows):
            datainrow=[]
            for col in range(cols):
                    text=sheet.cell_value(row,col)
                    datainrow.append((text))
            studentdata.append(datainrow)
        #print(studentdata)
        
        for row1 in range(trows):
            datainrow=studentdata[row1]
            #print(datainrow)
            for col1 in range(cols):
                text=datainrow[col1]
                if type(text) is str:
                    newitem = QTableWidgetItem(text)
                    self.tablewidget.setItem(row1, col1, (newitem))
                    print('string',text)
                else :
                    text1=int(text)
                    #print('integer',text1)
                    text2=str(text1)
                    newitem = QTableWidgetItem(text2)
                
                    #print(text,row1+3,col1)
                    self.tablewidget.setItem(row1, col1, (newitem))        
        
    def registerinput(self):
        text,result=QInputDialog.getInt(self,'inputdialog',"enter rows")
        if result==True:
            self.ddd=text
        self.row=self.tablewidget.rowCount()
    
        self.tablewidget.setRowCount(self.row+self.ddd)
        #print(self.ddd)
    
    def getData(self):
        self.rtr=self.tablewidget.rowCount()
        rowgiven=self.rtr
        data=[]
        for row in range(rowgiven):
            datainrow=[]
            for col in range(2):
                if self.tablewidget.item(row,col):
                    text=self.tablewidget.item(row,col).text()
                    datainrow.append((text))
                else:
                    datainrow.append("")
            data.append(datainrow)
        global tabledata 
        self.data=data
        tabledata=data
        #print(self.data)
        self.tablewidget.resizeColumnsToContents()

    def keyPressEvent(self, ev):
        
        if (ev.key() == Qt.Key_V) and (ev.modifiers() & Qt.ControlModifier): 
            self.pasteSelection()
        if (ev.key() == Qt.Key_C) and (ev.modifiers() & Qt.ControlModifier): 
            self.copySelection()
    def copySelection(self):
        selection = self.tablewidget.selectedIndexes()
        data=self.data
        if selection:
            rows = sorted(index.row() for index in selection)
            columns = sorted(index.column() for index in selection)
            rowcount = rows[-1] - rows[0] + 1
            colcount = columns[-1] - columns[0] + 1
            table = [[''] * colcount for _ in range(rowcount)]
            for index in selection:
                row = index.row() - rows[0]
                column = index.column() - columns[0]
                table[row][column] = index.data()
                
            stream = io.StringIO()
            csv.writer(stream, delimiter='\t').writerows(table)
            QApplication.clipboard().setText(stream.getvalue())
        
    def pasteSelection(self):
        selection = self.tablewidget.selectedIndexes()
        if selection:
            model = self.tablewidget.model()
            buffer = QApplication.clipboard().text()
            print(buffer)

            rows = sorted(index.row() for index in selection)
            columns = sorted(index.column() for index in selection)
            
            reader = csv.reader(io.StringIO(buffer), delimiter='\t')
            if len(rows) == 1 and len(columns) == 1:
                for i, line in enumerate(reader):
                    for j, cell in enumerate(line):
                        model.setData(model.index(rows[0]+i,columns[0]+j), cell)
            else:
                arr = [ [ cell for cell in row ] for row in reader]
                for index in selection:
                    row = index.row() - rows[0]
                    column = index.column() - columns[0]
                    model.setData(model.index(index.row(), index.column()), arr[row][column])
    
class Controller:

    def __init__(self):
        pass

    def show_login(self):
        self.login = Login()
        self.login.switch_window.connect(self.show_main)
        self.login.switch_window1.connect(self.show_admin)

        self.login.show()
    def show_admin(self):
        self.adwindow = adminWindow()
        self.login.hide()
        self.adwindow.show()
    def show_main(self):
        self.window = MainWindow()
        self.window.switch_window.connect(self.show_main1)
        self.window.switch_window1.connect(self.re_login)
        self.login.hide()
        self.window.show()
    def show_main1(self):
        self.window1 = MainWindow1()
        self.window1.switch_window.connect(self.show_window_three)
        self.window1.switch_window1.connect(self.re_login1)

        self.window.hide()
        self.window1.show()
    def show_window_three(self):
        self.window_three = Windowthree()
        self.window_three.switch_window.connect(self.show_window_two)

        self.window1.hide()
        self.window_three.show()
    def show_window_two(self):
        self.window_two = WindowTwo()
        self.window_three.hide()
        self.window_two.show()
    
    def re_login(self):
        self.login = Login()
        self.login.switch_window.connect(self.show_main)
        self.window.hide()
        self.login.show()
    def re_login1(self):
        self.window = MainWindow()
        self.window.switch_window.connect(self.show_main1)
        self.window.switch_window1.connect(self.re_login)
        self.window1.hide()
        self.window.show()

def main():
    app = QtWidgets.QApplication(sys.argv)
    app.setStyle('Fusion')#Windows,Fusion,Plastique,Breeze,Oxygen
    
    palette = QPalette()
    palette.setColor(QPalette.Window, QColor(53,53,53,225))
    palette.setColor(QPalette.WindowText, Qt.white)
    palette.setColor(QPalette.Base, QColor(25, 25, 25,255))
    palette.setColor(QPalette.AlternateBase, QColor(53, 53, 53,100))
    palette.setColor(QPalette.ToolTipBase, Qt.white)
    palette.setColor(QPalette.ToolTipText, Qt.white)
    palette.setColor(QPalette.Text, Qt.white)
    palette.setColor(QPalette.Button, QColor(53, 53, 53,225))
    palette.setColor(QPalette.ButtonText, Qt.white)
    palette.setColor(QPalette.BrightText, Qt.red)
    palette.setColor(QPalette.Link, QColor(42, 130, 218))
    palette.setColor(QPalette.Highlight, QColor(254, 254, 218,255))
    palette.setColor(QPalette.HighlightedText, Qt.blue)
    app.setPalette(palette)
    controller = Controller()
    controller.show_login()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()
